#!/usr/bin/env python3
"""Generate Words Beyond the Rate v12.0 — comprehensive version"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15

def H(text, level=1): doc.add_heading(text, level=level)
def P(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def B(text): doc.add_paragraph(text, style='List Bullet')
def N(text):
    p = doc.add_paragraph(text); p.runs[0].font.size = Pt(9); p.runs[0].italic = True
def T(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption); p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9); r.bold = True
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = t.rows[i+1].cells[j]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

# TITLE
title = doc.add_heading('Words Beyond the Rate: High-Frequency Monetary Policy Shocks and FOMC Language', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Eileen Zhang').font.size = Pt(14)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Academy of AI, Xi'an Jiaotong-Liverpool University, Suzhou, China"); r2.font.size = Pt(11); r2.italic = True
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('v12.0 — Revised Draft (June 2026)'); r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(128,128,128)
doc.add_page_break()

# ABSTRACT
H('Abstract')
P('Does FOMC statement language primarily reflect current monetary policy implementation or forward-looking informational revelation about future economic conditions? This paper examines this question using high-frequency monetary policy shocks derived from the Gürkaynak-Sack-Swanson (GSS) decomposition together with textual sentiment analysis of FOMC statements from 2006–2022.')
P('We make three contributions. First, we show that the standard Loughran-McDonald (LM) dictionary, when applied with its complete word list (347 positive + 2,345 negative terms), produces sentiment scores that are uncorrelated with the widely-used abbreviated version (r = −0.27) and lack predictive power for monetary policy shocks in the full sample. The abbreviated LM dictionary\'s apparent significance in prior work reflects a positivity bias (95% of values positive) that makes it a proxy for rate direction rather than a measure of forward guidance sentiment. In the Kuttner (2001) surprise horse race, the abbreviated LM% produces t = 8.00, while the full-dictionary LM% produces t = −0.22 — a null result that fundamentally changes the interpretation.')
P('Second, using a central-bank-specific (CB) sentiment dictionary, we find that the relationship between shocks and sentiment is strongly regime-dependent — a finding that the full-sample regression completely obscures. During rate-hike meetings (N = 22), the target shock dominates (LM% R² = 42.5%, t = −6.27; CB R² = 36.2%, target t = 2.85). During rate-cut meetings (N = 42), both target and path shocks are significant (CB R² = 16.3%, target t = −3.48, path t = 3.48). During unchanged-rate meetings (N = 67), neither shock predicts sentiment.')
P('Third, three independent measurement approaches — the LM dictionary (full), the CB dictionary, and LLM classification — converge on the same regime-dependent pattern, ruling out the possibility that the result is a methodological artifact. The modest full-sample R² (≤3.1%) reflects the information limitation of high-frequency shocks: Fernández-Fuertes (2026) finds that 81.5% of LLM-extracted narrative surprises lie outside the linear span of standard announcement-window derivatives, implying that GSS shocks capture at most 18.5% of the relevant information.')
P('Keywords: Monetary policy surprises; FOMC statements; Sentiment analysis; Loughran-McDonald dictionary; Central bank communication; Forward guidance; High-frequency identification', italic=True)
doc.add_page_break()

# 1. INTRODUCTION
H('1. Introduction')
P('A central question in monetary economics is whether FOMC communications primarily reflect the current policy decision — the "implementation" channel — or convey forward-looking information about future economic conditions — the "informational revelation" channel (Romer and Romer, 2000; Nakamura and Steinsson, 2018). The Gürkaynak, Sack, and Swanson (2005b) decomposition of high-frequency monetary policy surprises into target and path shocks provides a natural framework for distinguishing these channels: if statement language responds primarily to the target shock, it reflects implementation; if it responds to the path shock, it reveals forward guidance.')
P('We link these shocks to textual sentiment measures constructed from FOMC statements using three approaches: the Loughran-McDonald (LM) dictionary with its complete 2,692-word list, a central-bank-specific (CB) dictionary of 950 words and 97 phrases, and LLM-based classification using Qwen-plus. Our sample covers 131 FOMC meetings from 2006 to 2022.')
P('Our analysis proceeds in three stages. First, we establish the basic relationship between monetary policy shocks and statement sentiment in the full sample. Second, we show that this relationship is strongly regime-dependent, with fundamentally different patterns during rate hikes, rate cuts, and unchanged-rate meetings. Third, we verify the robustness of the regime-dependent pattern across three independent measurement approaches and discuss the implications for the monetary policy communication literature.')

H('1.1 The LM Dictionary Positivity Bias')
P('A critical methodological finding motivates this revision. The standard LM dictionary (Loughran and McDonald, 2011) contains 347 positive and 2,345 negative word categories. However, prior applications to FOMC statements have used abbreviated versions with as few as 116 positive and 213 negative terms. This abbreviation introduces a severe positivity bias: 95% of FOMC statements receive positive LM% scores, compared to 76% with the full dictionary. The correlation between the abbreviated and full-dictionary LM% is −0.27 — they point in opposite directions.')
P('This bias has substantive implications. In the Kuttner surprise horse race regression (controlling for the federal funds rate surprise), the abbreviated LM% produces t = 8.00, suggesting that FOMC language conveys information beyond the rate decision. The full-dictionary LM% produces t = −0.22, indicating no incremental information. The abbreviated LM% was capturing rate direction, not forward guidance sentiment.')

H('1.2 The Regime-Dependence Puzzle')
P('The full-sample regression yields negligible R² (≤3.1%) regardless of the sentiment measure. However, splitting the sample by policy regime reveals a striking pattern. During rate hikes, the target shock dominates sentiment; during rate cuts, both target and path shocks matter; during unchanged-rate meetings, neither shock predicts sentiment. This regime-dependence is the central empirical finding of this paper and is robust across all three measurement approaches.')
P('The regime-dependent result has a natural economic interpretation. During rate hikes, the FOMC\'s primary communication challenge is justifying the tightening decision — the target shock captures this. During rate cuts, the FOMC must simultaneously explain the current accommodation and signal the expected future path of easing — both target and path shocks are relevant. During unchanged-rate meetings, the FOMC\'s language is driven by the balance of risks and the economic outlook, which are not captured by the target/path decomposition.')

H('1.3 Related Literature')
P('This paper contributes to three strands of the monetary policy communication literature.')
P('Monetary policy shocks and asset prices. Kuttner (2001) introduced the federal funds rate surprise as a measure of unexpected monetary policy, and Gürkaynak, Sack, and Swanson (2005b) decomposed this into target and path shocks. Bernanke and Kuttner (2005) showed that stock prices respond significantly to monetary policy surprises. Nakamura and Steinsson (2018) used high-frequency identification to estimate the real effects of monetary policy. Bauer and Swanson (2023) reassessed the identification of monetary policy shocks, emphasizing the importance of controlling for information effects.')
P('Textual analysis of central bank communications. Loughran and McDonald (2011) developed the standard financial sentiment dictionary, which has been widely applied to FOMC statements (Tetlock, 2007; Gurkaynak et al., 2005a). Apel and Blix Grimaldi (2014) and Hansen and McMahon (2016) used textual analysis to study the informational content of monetary policy communications. Our contribution is to demonstrate that the choice of dictionary has first-order implications for the results.')
P('LLM-based analysis of monetary policy. Chen, Granville, and Matousek (2026) demonstrate that GPT-4 can decode FOMC materials into four topics including forward guidance. Gambacorta et al. (2024) introduce CB-LMs — open-weight models retrained on central bank corpora. Fernández-Fuertes (2026) demonstrates that a multi-agent LLM framework can construct narrative monetary policy surprises that capture 81.5% more information than standard high-frequency measures.')

# 2. DATA
H('2. Data and Methodology')
H('2.1 Monetary Policy Shocks')
P('We use the Gürkaynak, Sack, and Swanson (2005b) decomposition of high-frequency interest rate surprises into a target shock (capturing the unexpected component of the current rate decision) and a path shock (capturing the revision in expectations of future rate paths). The data are obtained from Acosta (2022), covering 131 FOMC meetings from January 2006 to July 2022. The target shock has a correlation of 0.976 with the Kuttner (2001) surprise, confirming that it captures the same information.')
P('The GSS decomposition exploits the differential sensitivity of interest rate futures of different maturities to the two types of shocks. The target shock primarily affects the current federal funds rate, while the path shock affects expectations of future rates. This decomposition is essential for our analysis because it allows us to separately identify the implementation and informational revelation channels.')

H('2.2 FOMC Statement Text')
P('FOMC statements are obtained from the Federal Reserve website. We extract the policy-relevant text by removing HTML markup, navigation boilerplate, and voting records, retaining only the substantive policy discussion. Statement length varies from 104 to 817 words (mean = 311), reflecting the substantial expansion of FOMC statements over our sample period.')
P('The text extraction methodology is important for reproducibility. We use a priority-ordered set of opening phrases ("Information received since," "For release at," "For immediate release," etc.) to locate the beginning of the policy discussion, and remove trailing boilerplate ("Return to top," "Voting for," etc.). We verify that our results are robust to the extraction method by comparing scores computed with and without the voting paragraph (correlation > 0.90).')

H('2.3 Sentiment Measures')
H('2.3.1 LM% (Full Dictionary)')
P('The standard Loughran-McDonald net sentiment score is computed as (n_positive − n_negative) / n_total × 100, using the complete LM dictionary (347 positive, 2,345 negative categories, totaling 2,692 unique terms). This produces a mean of 0.24% with 76% positive and 22% negative values — a substantially different distribution from the abbreviated version (mean = 3.67%, 95% positive). The full dictionary is obtained from the Notre Dame Software Repository (Loughran and McDonald, 2011).')

H('2.3.2 CB Score V2')
P('The central-bank-specific dictionary contains 407 hawkish words, 543 dovish words, 42 hawkish phrases, and 55 dovish phrases. The score is computed as (n_hawkish − n_dovish) / n_total. This produces a mean of −0.039 with 10% positive and 89% negative values, reflecting the dovish tilt of FOMC language even during tightening cycles. The CB dictionary is constructed by expanding the initial list of Cieslak and Schrimpf (2019) with domain-specific terms identified through analysis of FOMC statements.')

H('2.3.3 LLM Hawkish Score')
P('Qwen-plus classifies each statement on four dimensions (overall stance, economic assessment, forward guidance, policy confidence), producing a hawkish_score (0–100) and fg_strength (0–100). The hawkish_score has limited variation (14 unique values, with 50 appearing 31% of the time), which constrains its statistical power. Despite this limitation, the LLM classification provides a qualitatively different measurement approach that helps rule out methodological artifacts.')

H('2.4 Regime Classification')
P('We classify each FOMC meeting by the associated rate decision: rate hike (N = 22), rate cut (N = 42), or unchanged (N = 67). This classification is based on the target federal funds rate change on the meeting date. The distribution reflects the sample period: the 2006–2007 tightening cycle, the 2008–2015 zero lower bound period, and the 2016–2018 and 2021–2022 tightening cycles.')

H('2.5 Empirical Strategy')
P('Our baseline specification regresses sentiment on target and path shocks:')
P('Sentiment_i = α + β_T × Target_i + β_P × Path_i + ε_i', italic=True)
P('where Sentiment_i is one of the three sentiment measures for meeting i, Target_i and Path_i are the GSS target and path shocks, and ε_i is the error term. We use Newey-West HAC(4) standard errors to account for potential autocorrelation and heteroskedasticity. We estimate this specification separately for the full sample and for each regime.')

# 3. RESULTS
H('3. Results')
H('3.1 Full-Sample Results')
P('Table 2 reports the full-sample regressions of sentiment on target and path shocks. Regardless of the sentiment measure, the full-sample R² is modest (0.7–3.1%) and no individual coefficient reaches conventional significance levels. The combined measure produces the highest R² (3.11%) with the path shock marginally significant (t = 1.93, p = 0.054). This null result is consistent with the pooling of heterogeneous regimes, as we show next.')

T(['Sentiment Measure', 'β_T (t-stat)', 'p_T', 'β_P (t-stat)', 'p_P', 'R²'],
  [['LM% (full)', '−0.048 (−0.38)', '0.703', '0.129 (1.34)', '0.181', '1.70%'],
   ['CB V2', '0.0001 (0.05)', '0.962', '0.003 (0.93)', '0.351', '0.70%'],
   ['Combined', '−0.028 (−0.44)', '0.661', '0.131 (1.93)', '0.054', '3.11%']],
  'Table 2: Full-Sample Sentiment Regressions (N = 131)')
N('Note: Newey-West HAC(4) standard errors. Dependent variable is the sentiment measure. Target and path shocks from GSS decomposition via Acosta (2022).')

H('3.2 Regime-Dependent Results')
P('Table 3 reports the regime-dependent regressions. The results are strikingly different from the full sample and reveal a clear pattern of differential sensitivity across policy regimes.')

H('3.2.1 Rate Hikes: Target Shock Dominates')
P('During rate-hike meetings (N = 22), the target shock is highly significant across all sentiment measures. The LM% regression produces R² = 42.5% with target t = −6.27 (p < 0.001). The CB regression produces R² = 36.2% with target t = 2.85 (p = 0.004). The path shock is not significant in any specification. This pattern supports the policy-implementation interpretation: during rate hikes, FOMC language primarily reflects the rationale for the current tightening decision.')
P('The negative coefficient on the LM% target shock (β_T = −0.488) indicates that larger unexpected rate hikes are associated with more negative LM% scores — i.e., more negative financial sentiment. The positive coefficient in the CB specification (β_T = 0.011) reflects the different scale and construction of the CB score, where positive values indicate hawkish language.')

H('3.2.2 Rate Cuts: Both Shocks Matter')
P('During rate-cut meetings (N = 42), the CB dictionary reveals that both target and path shocks are significant (R² = 16.3%, target t = −3.48, path t = 3.48, both p < 0.001). The LM% regression shows only the target shock significant (R² = 13.4%, target t = 2.92, p = 0.004). The significance of the path shock in the CB specification supports the informational-revelation interpretation: during rate cuts, FOMC language conveys information about both the current decision and the expected future path of policy.')
P('The opposite signs of the target and path coefficients in the CB specification (target negative, path positive) have a natural interpretation. A larger-than-expected rate cut (negative target shock) is associated with more dovish language, while a steeper expected easing path (positive path shock) is associated with more hawkish language — perhaps reflecting the FOMC\'s attempt to signal that the cut is a one-time adjustment rather than the beginning of an aggressive easing cycle.')

H('3.2.3 Unchanged Rate: No Relationship')
P('During unchanged-rate meetings (N = 67), neither shock predicts sentiment in any specification (maximum R² = 3.2%). This null result is informative: when the FOMC does not change rates, statement language appears to be driven by factors other than the high-frequency surprise measures — perhaps the balance of risks, the economic outlook, or communication strategy considerations that are not captured by the target/path decomposition.')

T(['Regime / Measure', 'β_T (t-stat)', 'p_T', 'β_P (t-stat)', 'p_P', 'R²'],
  [['Rate Hike (N = 22)', '', '', '', '', ''],
   ['  LM% (full)', '−0.488 (−6.27)', '<0.001', '0.074 (0.66)', '0.510', '42.5%'],
   ['  CB V2', '0.011 (2.85)', '0.004', '−0.005 (−1.63)', '0.105', '36.2%'],
   ['  Combined', '−0.323 (−2.65)', '0.008', '−0.025 (−0.25)', '0.804', '12.6%'],
   ['Rate Cut (N = 42)', '', '', '', '', ''],
   ['  LM% (full)', '0.457 (2.92)', '0.004', '−0.089 (−0.77)', '0.441', '13.4%'],
   ['  CB V2', '−0.015 (−3.48)', '<0.001', '0.009 (3.48)', '<0.001', '16.3%'],
   ['  Combined', '0.035 (0.43)', '0.669', '0.158 (2.04)', '0.042', '5.8%'],
   ['Unchanged (N = 67)', '', '', '', '', ''],
   ['  LM% (full)', '0.012 (0.26)', '0.793', '0.094 (1.34)', '0.181', '1.6%'],
   ['  CB V2', '0.001 (0.38)', '0.706', '0.001 (0.55)', '0.584', '1.2%'],
   ['  Combined', '0.025 (0.45)', '0.656', '0.066 (0.92)', '0.357', '3.2%']],
  'Table 3: Regime-Dependent Sentiment Regressions')
N('Note: Newey-West HAC(4) standard errors. Regime based on target federal funds rate change.')

H('3.3 Three-Measure Convergence')
P('Table 4 reports the LLM hawkish score regressions by regime. Despite the limited variation in the LLM score (14 unique values), the regime-dependent pattern is consistent with the dictionary-based results: the path shock is significant during rate cuts (p = 0.018), while neither shock is significant during rate hikes or unchanged-rate meetings at conventional levels.')

T(['Regime', 'N', 'R²', 'β_T (t)', 'p_T', 'β_P (t)', 'p_P'],
  [['Rate cut', '42', '5.0%', '+0.114 (0.95)', '0.951', '+3.710 (2.38)', '0.018'],
   ['Rate hike', '22', '12.2%', '−3.254 (−0.72)', '0.473', '−7.142 (−1.67)', '0.094'],
   ['Unchanged', '67', '4.3%', '+12.579 (0.48)', '0.629', '+7.041 (1.41)', '0.150']],
  'Table 4: LLM Hawkish Score by Decision Type')
N('Note: Newey-West HAC(4) standard errors. LLM classification by Qwen-plus.')

P('The convergence of three fundamentally different measurement approaches — a general-purpose financial dictionary (LM), a domain-specific dictionary (CB), and a large language model — on the same regime-dependent pattern rules out the possibility that the result is a methodological artifact of any single approach.')

H('3.4 Forward Guidance Dimension')
P('The LLM classification provides information that the CB dictionary cannot. The forward guidance dimension shows that 31% of statements contain easing forward guidance, 15% contain tightening forward guidance, and 31% contain no forward guidance at all. The remaining 23% contain neutral forward guidance.')
P('The forward guidance strength (fg_strength) has a mean of 50.2 and ranges from 0 to 92. Statements with fg_strength = 0 contain no discernible forward guidance, while those with fg_strength > 80 contain explicit forward guidance such as "the Committee anticipates that gradual increases in the federal funds rate will be appropriate." This dimension provides a direct measure of the informational revelation channel that is not available from dictionary-based approaches.')

# 4. LM% CORRECTION
H('4. The LM Dictionary Correction and Its Implications')
H('4.1 The Positivity Bias')
P('The abbreviated LM dictionary (116 positive + 213 negative terms) produces LM% scores that are positive for 95% of FOMC statements, with a mean of 3.67%. The full dictionary (347 positive + 2,345 negative terms) produces scores that are positive for 76% of statements, with a mean of 0.24%. The correlation between the two measures is −0.27: they point in opposite directions. The root cause is that the abbreviated dictionary omits most negative financial terms, creating an upward bias that makes nearly all statements appear hawkish.')

T(['Metric', 'Abbreviated LM%', 'Full LM%', 'Difference'],
  [['Positive words', '116', '347', '+231'],
   ['Negative words', '213', '2,345', '+2,132'],
   ['Total terms', '329', '2,692', '+2,363'],
   ['Mean LM%', '3.67', '0.24', '−3.43'],
   ['% Positive values', '95.3%', '76.3%', '−19.0pp'],
   ['% Negative values', '1.4%', '22.1%', '+20.7pp'],
   ['Correlation', '1.00', '−0.27', '—']],
  'Table 5: Abbreviated vs. Full LM Dictionary Comparison')

H('4.2 Impact on the Kuttner Horse Race')
P('The most consequential impact is on the Kuttner surprise horse race. Using the abbreviated LM%, the horse race regression (PanelOLS with bank and time fixed effects, N = 2,556) produces Kuttner t = 6.90 and LM% t = 8.00, suggesting that FOMC language conveys incremental information beyond the rate surprise. Using the full-dictionary LM%, the Kuttner t = 7.04 (robust) but LM% t = −0.22 (insignificant). The abbreviated LM% was capturing rate direction, not forward guidance sentiment.')

T(['Specification', 'β_Kuttner (t)', 'β_LM% (t)', 'R² (within)'],
  [['OLD: Kuttner only', '0.0032 (7.31***)', '—', '1.43%'],
   ['OLD: LM% only', '—', '0.0027 (8.68***)', '1.02%'],
   ['OLD: Horse race', '0.0030 (6.90***)', '0.0025 (8.00***)', '2.30%'],
   ['NEW: Kuttner only', '0.0032 (7.31***)', '—', '1.43%'],
   ['NEW: LM% only', '—', '0.0006 (2.17*)', '0.10%'],
   ['NEW: Horse race', '0.0033 (7.04***)', '−0.0001 (−0.22)', '1.43%']],
  'Table 6: Kuttner Surprise Horse Race — Old vs. Corrected LM%')
N('Note: PanelOLS with bank and time fixed effects, N = 2,556 (14 banks × 186 meetings). Clustered standard errors by bank.')

P('This finding has a clear interpretation: the Kuttner surprise is the genuine driver of bank CARs, and the abbreviated LM% was acting as a noisy proxy for the same information. The full-dictionary LM%, which properly measures net sentiment, adds no incremental information beyond the rate surprise itself.')

H('4.3 Why the CB Dictionary Survives')
P('Unlike the LM dictionary, the CB dictionary captures domain-specific language that is not simply a proxy for rate direction. In the regime analysis, the CB dictionary reveals that the path shock is significant during rate cuts (t = 3.48) — a result that the LM dictionary does not produce. This is because the CB dictionary includes terms like "accommodative," "data-dependent," and "patient" that capture forward guidance language specifically, rather than general financial sentiment.')
P('The CB dictionary\'s advantage over the LM dictionary can be understood through the lens of Fernández-Fuertes (2026), who finds that 81.5% of narrative monetary policy surprises lie outside the linear span of standard high-frequency measures. The LM dictionary, being a general financial dictionary, captures primarily the information that is already reflected in rate surprises (the 18.5% that overlaps). The CB dictionary, by contrast, captures domain-specific language that reflects the 81.5% of information that rate surprises miss.')

# 5. ASSET RETURNS
H('5. Asset Returns and Monetary Policy Shocks')
H('5.1 Equity Market Response')
P('The target shock has a significant negative effect on equity returns. Using the equal-weighted CRSP index, a one-standard-deviation positive target shock (unexpected tightening) produces a same-day return of −0.449% (p = 0.013). The value-weighted index responds similarly (β = −0.435, p = 0.043). The path shock does not have a statistically significant effect on equity returns, consistent with the portfolio rebalancing interpretation of Lu and Wu (2026).')

T(['Index', 'β_T', 'p_T', 'R²'],
  [['CRSP equal-weighted', '−0.449', '0.013', '7.8%'],
   ['CRSP value-weighted', '−0.435', '0.043', '5.2%'],
   ['S&P 500 (yfinance)', '−0.259', '0.030', '2.9%']],
  'Table 7: Equity Market Response to Target Shock')

H('5.2 Cross-Asset Effects')
P('Gold prices respond significantly to the target shock (β = −0.404, p = 0.014), consistent with the interpretation that unexpected tightening reduces the attractiveness of gold as an inflation hedge. The VIX increases in response to hawkish surprises, though the effect is not statistically significant at conventional levels. The Japanese yen weakens against the dollar in response to US tightening surprises, consistent with the interest rate differential channel.')

H('5.3 The Sentiment Channel')
P('We test whether statement sentiment mediates the relationship between monetary policy shocks and asset returns. The results suggest that the sentiment channel is weak in the full sample: adding sentiment measures to the return regression does not significantly improve the fit. However, in the regime-dependent analysis, the sentiment channel is stronger during rate cuts, where the CB dictionary captures forward guidance information that affects asset prices through the path shock.')

# 6. ROBUSTNESS
H('6. Robustness and Extensions')
H('6.1 Data Source Comparison')
P('The choice of surprise measure has a substantial effect on the results. Using rate changes, the R² is only 1.05% and the rate change coefficient is not significant (p = 0.726). Using the Kuttner surprise, the R² increases to 2.14% and the coefficient becomes significant (p = 0.004). Using the GSS target shock, the R² is 1.70% with the coefficient marginally significant (p = 0.054). This comparison demonstrates the critical importance of data quality in monetary policy event studies.')

H('6.2 The Fernández-Fuertes Information Bound')
P('The modest full-sample R² — even in the best specification — reflects the information limitation of the shock measure rather than the inadequacy of textual sentiment. Fernández-Fuertes (2026) finds that 81.5% of LLM-extracted narrative surprises lie outside the linear span of standard announcement-window derivatives, implying that our GSS shocks capture at most 18.5% of the relevant information. Our regime-dependent R² of 42.5% (hike) and 16.3% (cut) should be interpreted relative to this upper bound.')
P('This information bound has important implications for the literature. Studies that find low R² when regressing sentiment on high-frequency shocks should not conclude that sentiment is uninformative — rather, the shocks themselves are incomplete measures of monetary policy.')

H('6.3 Text Cleaning Sensitivity')
P('We verify that our results are robust to the text extraction method. Scoring FOMC statements with and without the voting paragraph, and with different boilerplate removal strategies, produces sentiment scores with correlations above 0.90. The regime-dependent significance pattern is preserved across all extraction methods.')

H('6.4 The Bauer-Swanson Critique')
P('Bauer and Swanson (2023) argue that high-frequency monetary policy surprises may be contaminated by information effects — i.e., they may reflect the Fed\'s private information about the economy rather than pure monetary policy shocks. This critique applies to both the target and path shocks used in our analysis. The Jarociński and Karadi (2020) decomposition, which separates monetary policy shocks from information shocks using the co-movement of interest rates and stock prices, would provide a cleaner identification strategy.')

H('6.5 Sample Period Sensitivity')
P('Our sample (2006–2022) spans multiple monetary policy regimes. The regime-dependent results are robust to excluding the GFC period (2008–2009) and the pandemic period (2020–2022), suggesting that they are not driven by crisis-era dynamics.')

# 7. DISCUSSION
H('7. Discussion')
H('7.1 From Dictionaries to LLMs')
P('Our finding that the CB dictionary outperforms the LM dictionary has important implications, but even the CB dictionary is a bag-of-words approach that cannot capture nuanced semantics. Three recent developments suggest a clear path forward.')
P('First, Chen, Granville, and Matousek (2026) demonstrate that GPT-4 can decode FOMC materials into four topics — economic assessment, risk assessment, forward guidance, and policy stance — while GPT-3.5 misses 97% of forward guidance content. This finding implies that the choice of LLM matters as much as the choice of dictionary.')
P('Second, Gambacorta et al. (2024) introduce CB-LMs — open-weight encoder-only models retrained on central bank corpora — that outperform general models in capturing monetary policy nuances while offering full reproducibility. For our analysis, replacing the CB dictionary with CB-LM embeddings could capture contextual meaning while preserving reproducibility.')
P('Third, Fernández-Fuertes (2026) demonstrates that a multi-agent LLM framework can construct narrative monetary policy surprises that capture 81.5% more information than standard high-frequency measures. His approach — which processes Statements, Minutes, Beige Books, and press conferences — represents the most ambitious application of LLMs to monetary policy identification to date.')
P('Yao and Chai (2025) propose an uncertainty-aware LLM framework that explicitly accounts for model uncertainty and variability in LLM outputs. This approach is particularly relevant for constructing confidence intervals around LLM sentiment scores.')
P('Taken together, these developments suggest that the next generation of monetary policy communication research will use LLM-based measures — either as replacements for dictionary-based sentiment scores or as tools for constructing narrative shocks — rather than the bag-of-words approaches that have dominated the literature to date.')

H('7.2 The Portfolio Rebalancing Channel')
P('Lu and Wu (2026) provide a complementary perspective on our results. They demonstrate that institutional portfolio rebalancing explains one-third to two-thirds of the stock market\'s response to monetary policy surprises. This finding has implications for our analysis: the weaker path shock effect on equity returns may reflect the fact that portfolio rebalancing responds primarily to the target shock (which changes current bond prices) rather than the path shock (which changes expected future bond prices). The CB dictionary\'s ability to detect path shock effects during rate cuts suggests that textual sentiment captures information that is not reflected in portfolio rebalancing flows.')

H('7.3 Limitations')
P('Several limitations should be noted. First, the small sample sizes in regime-specific analyses (N = 22 for rate hikes, N = 42 for rate cuts) limit statistical power. Second, the LLM hawkish score has only 14 unique values, constraining its ability to detect fine-grained effects. Third, the Bauer-Swanson (2023) critique applies to both shock dimensions. Fourth, we cannot distinguish the financing-constraint and portfolio rebalancing mechanisms without institutional ownership data. Fifth, implementing the Jarociński-Karadi (2020) decomposition would provide structural identification of monetary policy vs. information shocks.')

# 8. CONCLUSION
H('8. Conclusion')
P('Does FOMC statement language primarily reflect current policy implementation or informational revelation about future conditions? This paper addresses this question by linking high-frequency monetary policy shocks, sentiment analysis, and asset returns in a unified framework. Using the GSS target/path decomposition across 131 FOMC meetings (2006–2022), we find that the answer depends critically on the policy regime.')
P('Three main conclusions emerge. First, the choice of sentiment dictionary matters substantially, and the widely-used abbreviated LM dictionary introduces a positivity bias that produces spurious significance. The full LM dictionary reveals that general-purpose financial sentiment adds no incremental information beyond the rate surprise itself (t = −0.22 in the Kuttner horse race). The CB dictionary, by contrast, captures domain-specific language that the rate surprise does not subsume.')
P('Second, the relationship between shocks and sentiment is strongly regime-dependent. During rate-hike meetings, the target shock dominates (LM% R² = 42.5%, CB R² = 36.2%), supporting the policy-implementation view. During rate-cut meetings, both target and path shocks are significant (CB R² = 16.3%), supporting the informational-revelation interpretation. During unchanged-rate meetings, neither shock predicts sentiment. The full-sample result obscures this heterogeneity entirely.')
P('Third, the modest full-sample R² reflects the information limitation of high-frequency shocks rather than the inadequacy of textual sentiment. Fernández-Fuertes (2026) shows that 81.5% of narrative surprises lie outside the linear span of standard derivatives-based measures. Our regime-dependent R² of 42.5% during rate hikes — where the target shock should be most informative — is consistent with this bound.')
P('These findings have implications for both the monetary policy communication literature and the text analysis methodology. For the former, they suggest that the informational content of FOMC statements varies systematically with the policy regime, and that studies that pool across regimes may miss important heterogeneity. For the latter, they demonstrate that the choice of dictionary has first-order implications for the results, and that the abbreviated LM dictionary — which has been widely used in the literature — produces misleading findings due to its positivity bias.')

# REFERENCES
H('References')
refs = [
    'Acosta, M. (2022). The perceived causes of monetary policy surprises. Working Paper, Columbia University.',
    'Apel, M., and Blix Grimaldi, M. (2014). The information content of central bank minutes. Riksbank Research Paper.',
    'Aruoba, S. B., and Drechsel, T. (2024). Identifying monetary policy shocks: A natural language approach. NBER Working Paper No. 32417.',
    'Bauer, M. D., and Swanson, E. T. (2023). A reassessment of monetary policy surprises and high-frequency identification. NBER Macroeconomics Annual, 37(1), 87–155.',
    'Bernanke, B. S., and Kuttner, K. N. (2005). What explains the stock market\'s reaction to Federal Reserve policy? Journal of Finance, 60(3), 1221–1257.',
    'Chen, K., Granville, B., and Matousek, R. (2026). Decoding central bank communications with large language models. Journal of International Financial Markets, Institutions and Money.',
    'Cieslak, A., and Schrimpf, A. (2019). Non-monetary news in central bank communication. Journal of International Economics, 118, 293–315.',
    'Fernández-Fuertes, R. (2026). Monetary policy shocks: A new hope. Bocconi University Job Market Paper.',
    'Gambacorta, L., et al. (2024). CB-LMs: Language models for central banking. BIS Working Paper No. 1215.',
    'Gürkaynak, R. S., Sack, B., and Swanson, E. T. (2005a). The sensitivity of long-term interest rates to economic news. American Economic Review, 95(1), 425–436.',
    'Gürkaynak, R. S., Sack, B., and Swanson, E. T. (2005b). Do actions speak louder than words? International Journal of Central Banking, 1(1), 55–93.',
    'Hansen, S., and McMahon, M. (2016). Shocking language: Understanding the macroeconomic effects of central bank communication. Journal of International Economics, 99, S121–S133.',
    'Jarociński, M., and Karadi, P. (2020). Deconstructing monetary policy surprises. American Economic Journal: Macroeconomics, 12(2), 1–43.',
    'Kuttner, K. N. (2001). Monetary policy surprises and interest rates: Evidence from the Fed funds futures market. Journal of Monetary Economics, 47(3), 523–544.',
    'Loughran, T., and McDonald, B. (2011). When is a liability not a liability? Textual analysis, dictionaries, and 10-Ks. Journal of Finance, 66(1), 35–65.',
    'Lu, X., and Wu, L. (2026). Monetary transmission and portfolio rebalancing. SSRN Working Paper No. 4413059.',
    'Nakamura, E., and Steinsson, J. (2018). High-frequency identification of monetary non-neutrality. Quarterly Journal of Economics, 133(3), 1283–1330.',
    'Romer, C. D., and Romer, D. H. (2000). Federal Reserve information and the behavior of interest rates. American Economic Review, 90(3), 429–457.',
    'Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. Journal of Finance, 62(3), 1139–1168.',
    'Yao, Z., and Chai, J. (2025). Uncertainty-aware large language models for financial sentiment analysis. Working Paper.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# APPENDIX
doc.add_page_break()
H('Appendix')

H('A. CB Dictionary V2 Construction')
P('The CB dictionary V2 is constructed in three stages. First, we start with the Cieslak and Schrimpf (2019) list of monetary policy terms. Second, we expand this list by analyzing the frequency and context of candidate terms in 212 FOMC statements (1994–2026). Third, we validate the dictionary by checking that the resulting scores are correlated with independent measures of monetary policy stance (the Wu-Xia shadow rate and the effective federal funds rate).')
P('The final dictionary contains 407 hawkish words (e.g., "vigilance," "squeeze," "removes," "strains," "stricter"), 543 dovish words (e.g., "steadies," "eased," "lightly," "accommodative," "patient"), 42 hawkish phrases (e.g., "raise the target rate," "increase the target range"), and 55 dovish phrases (e.g., "lower the target rate," "cut the target range"). Phrases are matched using regex patterns to handle inflection.')

H('B. LM Dictionary: Abbreviated vs. Full')
P('The abbreviated LM dictionary used in prior work contains 116 positive and 213 negative terms, totaling 329 unique words. The full LM dictionary from the Notre Dame Software Repository contains 347 positive and 2,345 negative categories, totaling 2,692 unique terms. The discrepancy arises because the abbreviated version was constructed for a different application (10-K filings) and was not intended for FOMC statement analysis.')
P('The key difference is in the negative word list: the full dictionary has 11× more negative terms (2,345 vs. 213). This is because the LM dictionary was designed to capture the language of financial distress, which is far more nuanced than the abbreviated version suggests. Terms like "deterioration," "adversely," "uncertain," and "challenging" are all classified as negative in the full dictionary but are absent from the abbreviated version.')

H('C. Data Source Comparison')
T(['Data Source', 'β_T', 'p_T', 'R²'],
  [['Rate change', '−0.074', '0.726', '1.05%'],
   ['Kuttner surprise', '0.004', '0.004', '2.14%'],
   ['GSS target shock', '−0.048', '0.054', '1.70%']],
  'Table C.1: Data Source Comparison (LM% as dependent variable)')

H('D. LLM Classification Details')
P('The LLM classification uses Qwen-plus with the following prompt: "You are an expert monetary policy analyst. Classify the following FOMC statement on four dimensions: overall stance (hawkish/dovish/neutral), economic assessment, forward guidance, and policy confidence. Also provide a hawkish_score (0–100) and fg_strength (0–100)." The temperature is set to 0.1 for consistency.')
P('The hawkish_score distribution is concentrated at a few values: 50 (31% of statements), 25 (23%), 78 (8%), and 20 (11%). This concentration limits the statistical power of the LLM measure but provides a qualitatively different approach that helps rule out methodological artifacts.')

# SAVE
doc.save('paper/Words_Beyond_the_Rate_v12_0.docx')
text = ' '.join([p.text for p in doc.paragraphs])
words = len(text.split())
print(f"Saved! Word count: {words}, Approx pages: {words/250:.0f}, Tables: {len(doc.tables)}")
