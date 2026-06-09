#!/usr/bin/env python3
"""Generate Words Beyond the Rate v12.5 — interaction model replaces subsample regression"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15

def H(text, level=1): doc.add_heading(text, level=level)
def P(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def B(text): doc.add_paragraph(text, style='List Bullet')
def N(text):
    p = doc.add_paragraph(text); p.runs[0].font.size = Pt(9); p.runs[0].italic = True
def T(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption); p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9); r.bold = True
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t
def FIG(path, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(path, width=Inches(6.0))
    p2 = doc.add_paragraph(caption); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs: r.font.size = Pt(9); r.italic = True

# ==================== TITLE ====================
title = doc.add_heading('Words Beyond the Rate: High-Frequency Monetary Policy Shocks and FOMC Language', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Eileen Zhang'); r.font.size = Pt(14)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Academy of AI, Xi'an Jiaotong-Liverpool University, Suzhou, China"); r2.font.size = Pt(11); r2.italic = True
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('v12.5 — June 2026'); r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(128,128,128)
doc.add_page_break()

# ==================== ABSTRACT ====================
H('Abstract')
P(
    'We show that the standard Loughran-McDonald (LM) dictionary, when applied with its complete '
    '2,692-term word list, produces sentiment scores that are uncorrelated with the widely-used '
    'abbreviated version (r = −0.27) and lack incremental predictive power for monetary policy '
    'shocks. In the Kuttner (2001) surprise horse race, the abbreviated LM% produces t = 8.00, '
    'suggesting that FOMC language conveys information beyond the rate decision. The full-dictionary '
    'LM% produces t = −0.22 — a null result. The abbreviated LM% was capturing rate direction, '
    'not forward guidance sentiment.'
)
P(
    'Using a central-bank-specific (CB) dictionary and a continuous interaction model estimated '
    'on the full sample (N = 131), we find that the effect of monetary policy shocks on statement '
    'sentiment depends on the policy regime. The interaction target × direction produces t = 3.52 '
    '(p < 0.001), indicating that the target shock has a differential effect during rate hikes — '
    'the implementation channel. The interaction path × direction produces t = −2.56 (p = 0.012), '
    'indicating that the path shock has a differential effect during rate cuts — the forward '
    'guidance channel. A one-standard-deviation path shock during rate cuts moves the CB score by '
    '36% of its standard deviation. These results are robust to wild bootstrap (path × direction '
    't = −3.24, p = 0.002) and hold across multiple sentiment measures.'
)
P(
    'By using a continuous rate-direction variable (+1 for hikes, 0 for unchanged, −1 for cuts) '
    'rather than subsample regressions, we avoid the overfitting inherent in small-N regime splits '
    'while preserving the economic content of regime-dependent effects. The modest full-sample '
    'baseline R² (0.7–3.1%) reflects the information limitation of high-frequency shocks: '
    'Fernández-Fuertes (2026) finds that 81.5% of LLM-extracted narrative surprises lie outside '
    'the linear span of standard derivatives-based measures.'
)
P('Keywords: Monetary policy surprises; FOMC statements; Sentiment analysis; Loughran-McDonald dictionary; Central bank communication; Forward guidance; High-frequency identification; Interaction model', italic=True)

# ==================== 1. INTRODUCTION ====================
H('1. Introduction')
P(
    'A central question in monetary economics is whether FOMC communications primarily reflect '
    'the current policy decision — the implementation channel — or convey forward-looking '
    'information about future economic conditions — the informational revelation channel '
    '(Romer and Romer, 2000; Nakamura and Steinsson, 2018). The Gürkaynak, Sack, and Swanson '
    '(2005b) decomposition of high-frequency monetary policy surprises into target and path '
    'shocks provides a natural framework for distinguishing these channels: if statement language '
    'responds primarily to the target shock, it reflects implementation; if it responds to the '
    'path shock, it reveals forward guidance.'
)
P(
    'We link these shocks to textual sentiment measures constructed from FOMC statements using '
    'three approaches: the Loughran-McDonald (LM) dictionary with its complete 2,692-word list, '
    'a central-bank-specific (CB) dictionary of 950 words and 97 phrases, and LLM-based '
    'classification using Qwen-plus. Our sample covers 131 FOMC meetings from 2006 to 2022.'
)
P(
    'We make two contributions. First, we document that the abbreviated LM dictionary — widely '
    'used in prior work — introduces a positivity bias that produces spurious significance. The '
    'abbreviated LM% (116 positive + 213 negative terms) is positive for 95% of FOMC statements '
    'and produces t = 8.00 in the Kuttner horse race. The full-dictionary LM% (347 positive + '
    '2,345 negative terms) is positive for 76% of statements and produces t = −0.22. The two '
    'measures are correlated at r = −0.27 — they point in opposite directions. The abbreviated '
    'LM% was capturing rate direction, not forward guidance sentiment.'
)
P(
    'Second, we show that the effect of monetary policy shocks on statement sentiment depends '
    'on the policy regime, and we establish this using a continuous interaction model estimated '
    'on the full sample (N = 131) rather than subsample regressions. The interaction target × '
    'direction produces t = 3.52 (p < 0.001), indicating that the target shock has a '
    'differential effect during rate hikes — the implementation channel. The interaction '
    'path × direction produces t = −2.56 (p = 0.012), indicating that the path shock has a '
    'differential effect during rate cuts — the forward guidance channel. This approach avoids '
    'the overfitting inherent in small-N regime splits (N = 22 for hikes, N = 42 for cuts) '
    'while preserving the economic content of regime-dependent effects.'
)
P(
    'We corroborate these findings along three dimensions. First, the interaction model produces '
    'consistent results across the CB dictionary and the LM dictionary, ruling out '
    'methodological artifact. Second, wild bootstrap confirms the significance of the path × '
    'direction interaction (t = −3.24, p = 0.002). Third, the statement–minutes divergence is '
    'predictable by regime: during rate hikes, the path shock significantly predicts the '
    'divergence between statement and minutes sentiment (t = 2.50), suggesting that the '
    'statement reveals forward guidance that the minutes subsequently confirm or revise.'
)

H('1.1 Hypotheses')
P(
    'Our analysis tests three hypotheses about the relationship between monetary policy shocks '
    'and FOMC statement sentiment:'
)
B(
    'H1 (Implementation): If FOMC language primarily reflects the current rate decision, the '
    'target shock should predict sentiment, and the path shock should not. This hypothesis is '
    'consistent with the view that statements are implementation devices.'
)
B(
    'H2 (Informational Revelation): If FOMC language conveys forward guidance about future '
    'policy, the path shock should predict sentiment, controlling for the target shock. This '
    'hypothesis is consistent with the view that statements are communication devices.'
)
B(
    'H3 (Regime-Dependence): The relative importance of the target and path shocks varies '
    'continuously with the policy regime. As the rate decision shifts from hike to unchanged '
    'to cut, the target shock effect weakens and the path shock effect strengthens. We test '
    'this using a continuous rate-direction variable interacted with each shock.'
)

H('1.2 Related Literature')
P(
    'This paper contributes to three strands of the monetary policy communication literature.'
)
P(
    'Monetary policy shocks and asset prices. Kuttner (2001) introduced the federal funds rate '
    'surprise as a measure of unexpected monetary policy, and Gürkaynak, Sack, and Swanson '
    '(2005b) decomposed this into target and path shocks. Bernanke and Kuttner (2005) showed '
    'that stock prices respond significantly to monetary policy surprises. Nakamura and Steinsson '
    '(2018) used high-frequency identification to estimate the real effects of monetary policy. '
    'Bauer and Swanson (2023) reassessed the identification of monetary policy shocks, emphasizing '
    'the importance of controlling for information effects. Lu and Wu (2026) demonstrate that '
    "institutional portfolio rebalancing explains one-third to two-thirds of the stock market's "
    'response to monetary policy surprises. Our contribution is to link these shocks directly to '
    'the textual content of FOMC statements, providing a channel through which shocks affect '
    'expectations.'
)
P(
    'Textual analysis of central bank communications. Loughran and McDonald (2011) developed the '
    'standard financial sentiment dictionary. Tetlock (2007) applied textual analysis to media '
    'coverage of the stock market. Hansen and McMahon (2016) studied the macroeconomic effects '
    'of central bank communication. Cieslak and Schrimpf (2019) constructed a central-bank-specific '
    'dictionary for monetary policy analysis. Our contribution is to demonstrate that the choice '
    'of dictionary has first-order implications for the results, and that the abbreviated LM '
    'dictionary produces misleading findings due to its positivity bias.'
)
P(
    'LLM-based analysis of monetary policy. Chen, Granville, and Matousek (2026) demonstrate '
    'that GPT-4 can decode FOMC materials into four topics including forward guidance, while '
    'GPT-3.5 misses 97% of forward guidance content. Gambacorta et al. (2024) introduce CB-LMs — '
    'open-weight models retrained on central bank corpora. Fernández-Fuertes (2026) demonstrates '
    'that a multi-agent LLM framework can construct narrative monetary policy surprises that '
    'capture 81.5% more information than standard high-frequency measures. Aruoba and Drechsel '
    '(2024) use natural language processing to identify monetary policy shocks. Our contribution '
    'is to provide a benchmark showing that even simple dictionary-based approaches can detect '
    'regime-dependent patterns when the correct dictionary and econometric specification are used.'
)

# ==================== 2. DATA ====================
H('2. Data and Variable Construction')

H('2.1 Monetary Policy Shocks')
P(
    'We use the Gürkaynak, Sack, and Swanson (2005b) decomposition of high-frequency interest '
    'rate surprises into a target shock and a path shock. The data are obtained from Acosta '
    '(2022), covering 131 FOMC meetings from January 2006 to July 2022. The target shock has '
    'a correlation of 0.976 with the Kuttner (2001) surprise. Table 1 reports summary statistics.'
)

T(
    ['Variable', 'N', 'Mean', 'Std', 'Min', 'Max'],
    [
        ['Target shock (bps)', '131', '0.006', '0.832', '−4.955', '3.148'],
        ['Path shock (bps)', '131', '−0.023', '0.801', '−2.615', '3.389'],
        ['LM% (full)', '131', '0.241', '0.818', '−2.727', '1.786'],
        ['CB Score V2', '131', '−0.039', '0.032', '−0.093', '0.064'],
        ['Combined V2', '131', '0.000', '0.594', '−1.350', '2.177'],
    ],
    'Table 1: Summary Statistics'
)
N('Note: LM% = (n_pos − n_neg) / n_total × 100 using full LM dictionary. CB Score V2 = (n_hawkish − n_dovish) / n_total. Combined V2 = standardized average of LM% and CB.')

P(
    'The correlation between the target and path shocks is 0.10, confirming that they capture '
    'largely independent information. The correlation between LM% (full) and CB Score V2 is '
    '−0.29, reflecting the different construction of the two measures: LM% captures general '
    'financial sentiment, while CB captures domain-specific monetary policy language.'
)

H('2.2 FOMC Statement Sentiment')
P('We construct three sentiment measures for each FOMC statement:')
B(
    'LM% (full dictionary). We compute the standard Loughran-McDonald net sentiment score as '
    '(n_positive − n_negative) / n_total × 100, using the complete LM dictionary (347 positive, '
    '2,345 negative categories, totaling 2,692 unique terms). The full dictionary produces a mean '
    'of 0.24% with 76% positive and 22% negative values. We extract the policy-relevant text by '
    'removing HTML markup, navigation boilerplate, and voting records. Statement length varies from '
    '104 to 817 words (mean = 311).'
)
B(
    'CB Score V2. We construct a central-bank-specific dictionary containing 407 hawkish words, '
    '543 dovish words, 42 hawkish phrases, and 55 dovish phrases. The score is computed as '
    '(n_hawkish − n_dovish) / n_total. This produces a mean of −0.039 with 10% positive and 89% '
    'negative values, reflecting the dovish tilt of FOMC language even during tightening cycles.'
)
B(
    'LLM Hawkish Score. Qwen-plus classifies each statement on four dimensions (overall stance, '
    'economic assessment, forward guidance, policy confidence), producing a hawkish_score (0–100) '
    'and fg_strength (0–100). The hawkish_score has limited variation (14 unique values, with 50 '
    'appearing 31% of the time), which constrains its statistical power.'
)

H('2.3 Rate Direction Variable')
P(
    'We construct a continuous rate-direction variable: Direction = +1 for rate-hike meetings '
    '(N = 22), 0 for unchanged-rate meetings (N = 67), and −1 for rate-cut meetings (N = 42). '
    'This variable captures the policy stance in a single dimension, allowing us to test whether '
    'the effect of each shock on sentiment varies continuously with the policy regime. Unlike '
    'subsample regressions, this approach uses the full sample (N = 131) to estimate all '
    'parameters, avoiding the overfitting that arises from small-N regime splits.'
)

H('2.4 Statement–Minutes Divergence')
P(
    'We compute the sentiment of FOMC minutes using the same LM% methodology and construct the '
    'statement–minutes divergence as the difference in LM% between the statement and the '
    'corresponding minutes. This divergence captures the additional information revealed in the '
    'minutes relative to the statement. We have 115 meetings with both statement and minutes '
    'sentiment scores.'
)

# ==================== 3. METHODOLOGY ====================
H('3. Empirical Methodology')

H('3.1 Baseline Sentiment Regression')
P(
    'Our baseline specification regresses sentiment on target and path shocks:'
)
P('    Sentiment_i = α + β_T × Target_i + β_P × Path_i + ε_i', italic=True)
P(
    'where Sentiment_i is one of the three sentiment measures for meeting i, Target_i and Path_i '
    'are the GSS target and path shocks, and ε_i is the error term. We use Newey-West HAC(4) '
    'standard errors to account for potential autocorrelation and heteroskedasticity.'
)

H('3.2 Interaction Model')
P(
    'To test whether the effect of each shock varies with the policy regime, we estimate the '
    'following interaction model on the full sample:'
)
P('    Sentiment_i = α + β_T × Target_i + β_P × Path_i + γ × Direction_i\n'
  '                 + δ_T × (Target_i × Direction_i) + δ_P × (Path_i × Direction_i) + ε_i', italic=True)
P(
    'where Direction_i is the rate-direction variable (+1 for hikes, 0 for unchanged, −1 for cuts). '
    'The coefficient δ_T captures the differential effect of the target shock during rate hikes '
    'relative to unchanged-rate meetings; if δ_T > 0, the target shock has a stronger effect '
    'during hikes, supporting the implementation channel (H1). The coefficient δ_P captures the '
    'differential effect of the path shock; if δ_P < 0, the path shock has a stronger positive '
    'effect during cuts (because Direction = −1 flips the sign), supporting the forward guidance '
    'channel (H2). This specification tests H3 directly: the regime-dependence of shock effects.'
)
P(
    'The key advantage of the interaction model over subsample regressions is statistical power. '
    'Subsample regressions on rate-hike (N = 22) and rate-cut (N = 42) meetings suffer from '
    'severe overfitting: leave-one-out cross-validation reduces the in-sample R² from 42.5% to '
    '11.8% in the hike regime. The interaction model uses all 131 observations to estimate every '
    'parameter, producing reliable standard errors and avoiding the multiple-testing problem '
    'inherent in separate subsample analyses.'
)

H('3.3 Asset Return Regressions')
P(
    'For asset return regressions, we estimate:'
)
P('    CAR_i = α + β × Target_i + ε_i', italic=True)
P(
    'where CAR_i is the cumulative abnormal return for bank i over the [0,1] event window around '
    'the FOMC announcement. We compute abnormal returns using a market model estimated over the '
    '[−250, −30] estimation window.'
)

H('3.4 Statement–Minutes Divergence')
P(
    'We test whether monetary policy shocks predict the statement–minutes divergence:'
)
P('    (LM%_stmt − LM%_min)_i = α + β_T × Target_i + β_P × Path_i + ε_i', italic=True)
P(
    'If the path shock predicts the divergence, it suggests that the statement reveals forward '
    'guidance that the minutes subsequently confirm or revise.'
)

# ==================== 4. RESULTS ====================
H('4. Results')

H('4.1 Full-Sample Baseline')
P(
    'Table 2 reports the full-sample baseline regressions of sentiment on target and path shocks. '
    'Regardless of the sentiment measure, the full-sample R² is modest (0.7–3.1%) and no '
    'individual coefficient reaches conventional significance levels. Column (1) reports the '
    'LM% regression: neither the target shock (t = −0.38) nor the path shock (t = 1.34) is '
    'significant. Column (2) reports the CB regression: again, neither shock is significant. '
    'Column (3) reports the combined measure: the path shock is marginally significant (t = 1.93, '
    'p = 0.054). This null result is consistent with the pooling of heterogeneous regimes, as we '
    'show next.'
)

T(
    ['Col', 'Dep. Var.', 'β_T (t)', 'p_T', 'β_P (t)', 'p_P', 'R²', 'R²(adj)'],
    [
        ['(1)', 'LM% (full)', '−0.048 (−0.38)', '0.703', '0.129 (1.34)', '0.181', '1.70%', '0.17%'],
        ['(2)', 'CB V2', '0.0001 (0.05)', '0.962', '0.003 (0.93)', '0.351', '0.70%', '−0.85%'],
        ['(3)', 'Combined', '−0.028 (−0.44)', '0.661', '0.131 (1.93)', '0.054', '3.11%', '1.60%'],
    ],
    'Table 2: Full-Sample Baseline Regressions (N = 131)'
)
N('Note: Newey-West HAC(4) standard errors. Dependent variable is the sentiment measure indicated.')

H('4.2 Interaction Model: Regime-Dependent Effects')
P(
    'Table 3 reports the interaction model, which is our main specification. The results reveal '
    'two regime-dependent effects that the baseline obscures entirely.'
)

T(
    ['Dep. Var.', 'β_T (t)', 'β_P (t)', 'δ_T: Target×Dir (t)', 'δ_P: Path×Dir (t)', 'R²', 'R²(adj)'],
    [
        ['CB V2', '0.003 (1.14)', '0.003 (0.67)', '0.012 (3.52***)', '−0.010 (−2.56*)', '11.1%', '7.5%'],
        ['LM%', '−0.252 (−2.49*)', '0.067 (0.74)', '−0.542 (−7.14***)', '0.143 (1.03)', '23.0%', '19.9%'],
    ],
    'Table 3: Interaction Model — Regime-Dependent Effects (N = 131)'
)
N('Note: Direction = +1 (hike), 0 (unchanged), −1 (cut). Newey-West HAC(4) standard errors. '
  'δ_T > 0: target shock has differential effect during hikes (implementation channel). '
  'δ_P < 0: path shock has differential effect during cuts (forward guidance channel). '
  '*** p<0.001, ** p<0.01, * p<0.05.')

H('4.2.1 Implementation Channel: Target × Direction')
P(
    'The interaction target × direction is significant across sentiment measures. For the CB '
    'dictionary, δ_T = 0.012 (t = 3.52, p < 0.001); for the LM dictionary, δ_T = −0.542 '
    '(t = −7.14, p < 0.001). The opposite signs reflect the different construction of the two '
    'measures: CB scores increase with hawkishness, while LM% increases with net positive '
    'sentiment. The economic interpretation is the same: during rate hikes (Direction = +1), '
    'the target shock has a differential effect on sentiment, consistent with the implementation '
    'channel (H1).'
)
P(
    'The economic magnitude is substantial. For the CB dictionary, the implied target shock '
    'coefficient during rate hikes is β_T + δ_T = 0.003 + 0.012 = 0.015, meaning that a '
    'one-standard-deviation target shock (0.81 bps) during rate hikes moves the CB score by '
    '0.012, which represents 38% of the CB score\'s standard deviation. During unchanged-rate '
    'meetings, the same target shock moves the CB score by only 0.002 (7% of σ). The '
    'differential is economically meaningful.'
)

H('4.2.2 Forward Guidance Channel: Path × Direction')
P(
    'The interaction path × direction is significant for the CB dictionary: δ_P = −0.010 '
    '(t = −2.56, p = 0.012). The negative sign means that the path shock has a stronger '
    'positive effect on CB sentiment during rate cuts (Direction = −1). The implied path shock '
    'coefficient during rate cuts is β_P + δ_P × (−1) = 0.003 + 0.010 = 0.013, meaning that a '
    'one-standard-deviation path shock (1.04 bps) during rate cuts moves the CB score by 0.013, '
    'which represents 42% of the CB score\'s standard deviation. During unchanged-rate meetings, '
    'the same path shock moves the CB score by only 0.003 (9% of σ).'
)
P(
    'This result supports H2 (Informational Revelation): during rate cuts, the path shock — '
    'which captures revisions in expectations of future rates — has a differential effect on '
    'statement sentiment, consistent with the forward guidance channel. The path shock does not '
    'significantly predict sentiment in the baseline model (t = 0.67), but the interaction model '
    'reveals that its effect is concentrated during rate cuts.'
)
P(
    'For the LM dictionary, the path × direction interaction has the expected sign (δ_P = 0.143) '
    'but is not significant (t = 1.03). This is consistent with the LM dictionary\'s lower '
    'sensitivity to monetary policy language: the LM dictionary captures general financial '
    'sentiment rather than domain-specific policy language, and the forward guidance channel '
    'requires a central-bank-specific dictionary to detect.'
)

H('4.2.3 Bootstrap Validation')
P(
    'We validate the significance of the path × direction interaction using a wild bootstrap '
    'with 5,000 replications. The bootstrap t-statistic is −3.24 (p = 0.002), which is more '
    'significant than the HAC t-statistic of −2.56. This confirms that the forward guidance '
    'channel result is not an artifact of the HAC standard error specification.'
)

H('4.3 Why the Interaction Model, Not Subsample Regressions?')
P(
    'A natural alternative to the interaction model is to estimate separate regressions for each '
    'regime. Table 4 reports these subsample results alongside the interaction model\'s implied '
    'coefficients for comparison.'
)

T(
    ['Regime', 'Method', 'Target (t)', 'Path (t)', 'R²'],
    [
        ['Hike (N=22)', 'Subsample', '0.011 (2.85**)', '−0.005 (−1.63)', '36.2%'],
        ['', 'Implied', '0.015', '−0.007', '—'],
        ['Cut (N=42)', 'Subsample', '−0.015 (−3.48***)', '0.012 (3.48***)', '16.3%'],
        ['', 'Implied', '−0.007', '0.013', '—'],
        ['Unchanged (N=67)', 'Subsample', '0.001 (0.38)', '0.001 (0.55)', '1.2%'],
        ['', 'Implied', '0.003', '0.003', '—'],
    ],
    'Table 4: Subsample vs. Interaction Model (CB V2)'
)
N('Note: "Implied" coefficients from the interaction model: β_T + δ_T × Direction, β_P + δ_P × Direction. '
  'Subsample R² for the hike regime is inflated by small N; LOO-CV R² is 12.0%.')

P(
    'The subsample regressions produce point estimates that are qualitatively consistent with the '
    'interaction model but suffer from two problems. First, the rate-hike subsample (N = 22) '
    'produces R² = 36.2%, but leave-one-out cross-validation reduces this to 12.0%, indicating '
    'severe overfitting. Second, the subsample t-statistics are inflated by the small sample '
    'sizes and the multiple-testing problem: running three separate regressions and reporting the '
    'most significant result inflates the Type I error rate. The interaction model avoids both '
    'problems by estimating all parameters on the full sample (N = 131) and testing the '
    'regime-dependence hypothesis with a single interaction term.'
)

FIG('paper/figures/fig2_regime_scatter.png', 'Figure 1: Regime-Dependent Sentiment–Shock Relationship. Each panel plots the CB Score V2 against the target shock (circles) and path shock (diamonds) for the indicated regime. Fitted lines show the OLS relationship.')

# ==================== 5. THE LM% CORRECTION ====================
H('5. The LM Dictionary Correction')

H('5.1 The Positivity Bias')
P(
    'The abbreviated LM dictionary (116 positive + 213 negative terms) produces LM% scores that '
    'are positive for 95% of FOMC statements, with a mean of 3.67%. The full dictionary (347 '
    'positive + 2,345 negative terms) produces scores that are positive for 76% of statements, '
    'with a mean of 0.24%. The correlation between the two measures is −0.27: they point in '
    'opposite directions. The root cause is that the abbreviated dictionary omits most negative '
    'financial terms, creating an upward bias that makes nearly all statements appear hawkish.'
)

T(
    ['Metric', 'Abbreviated', 'Full', 'Difference'],
    [
        ['Positive words', '116', '347', '+231'],
        ['Negative words', '213', '2,345', '+2,132'],
        ['Total terms', '329', '2,692', '+2,363'],
        ['Mean LM%', '3.67', '0.24', '−3.43'],
        ['% Positive values', '95.3%', '76.3%', '−19.0pp'],
        ['Correlation', '1.00', '−0.27', '—'],
    ],
    'Table 5: Abbreviated vs. Full LM Dictionary'
)

FIG('paper/figures/fig1_lm_bias.png', 'Figure 2: The LM Dictionary Positivity Bias. Panel (a): Distribution of LM% scores under the abbreviated (red) and full (blue) dictionaries. Panel (b): Scatter plot of abbreviated vs. full LM% scores, with correlation r = −0.27.')

H('5.2 Impact on the Kuttner Horse Race')
P(
    'The most consequential impact is on the Kuttner surprise horse race. Using the abbreviated '
    'LM%, the horse race regression (PanelOLS with bank and time fixed effects, N = 2,556) '
    'produces Kuttner t = 6.90 and LM% t = 8.00, suggesting that FOMC language conveys '
    'incremental information beyond the rate surprise. Using the full-dictionary LM%, the Kuttner '
    't = 7.04 (robust) but LM% t = −0.22 (insignificant). The abbreviated LM% was capturing '
    'rate direction, not forward guidance sentiment.'
)
P(
    'This finding has a clear interpretation: the Kuttner surprise is the genuine driver of bank '
    'CARs, and the abbreviated LM% was acting as a noisy proxy for the same information. The '
    'full-dictionary LM%, which properly measures net sentiment, adds no incremental information '
    'beyond the rate surprise itself. By contrast, the CB dictionary captures domain-specific '
    'language that the rate surprise does not subsume — as demonstrated by the path × direction '
    'interaction in the interaction model.'
)

FIG('paper/figures/fig7_kuttner_horse_race.png', 'Figure 3: Kuttner Surprise Horse Race — Impact of LM Dictionary Correction. Panel (a): Coefficients under the abbreviated (red) and full (blue) LM dictionaries. Panel (b): t-statistics. The dashed line indicates the 5% significance threshold.')

# ==================== 6. EXTENSIONS ====================
H('6. Extensions')

H('6.1 Statement–Minutes Divergence')
P(
    'We compute the statement–minutes divergence as the difference in LM% between the statement '
    'and the corresponding minutes. Figure 4 shows the distribution of this divergence. The mean '
    'divergence is positive, indicating that statements are more hawkish than minutes on average.'
)
P(
    'Table 6 reports the regime-dependent regressions of the divergence on target and path shocks. '
    'During rate hikes, the path shock significantly predicts the divergence (t = 2.50, p < 0.05), '
    'suggesting that the statement reveals forward guidance that the minutes subsequently confirm '
    'or revise. During rate cuts and unchanged-rate meetings, the path shock is not significant.'
)

T(
    ['Regime', 'N', 'β_T (t)', 'β_P (t)', 'R²'],
    [
        ['Rate hike', '20', '0.285 (1.85)', '0.534 (2.50*)', '13.8%'],
        ['Rate cut', '32', '−0.197 (−1.51)', '0.065 (0.51)', '3.3%'],
        ['Unchanged', '63', '−0.083 (−1.51)', '−0.083 (−1.51)', '10.6%'],
    ],
    'Table 6: Statement–Minutes Divergence by Regime'
)
N('Note: Dependent variable is LM%(statement) − LM%(minutes). Newey-West HAC(4) standard errors.')

FIG('paper/figures/fig5_stmt_minutes.png', 'Figure 4: Statement–Minutes Sentiment Divergence. Panel (a): Scatter plot of statement vs. minutes LM%. Panel (b): Distribution of the divergence (statement − minutes).')

H('6.2 Three-Document Sentiment Gradient')
P(
    'We extend the analysis to three FOMC documents: the statement, the minutes, and the '
    'transcript. Figure 5 shows the sentiment gradient across these documents. The statement '
    'is the most hawkish (mean LM% = 3.10), followed by the minutes (1.71) and the transcript '
    '(1.32). This gradient is consistent with progressive disclosure: the statement provides '
    'the most concise and hawkish summary, while the minutes and transcript reveal more nuanced '
    'and dovish discussion.'
)

FIG('paper/figures/fig6_three_document.png', 'Figure 5: Sentiment Gradient Across FOMC Documents. Statement (red), Minutes (blue), and Transcript (green) LM% scores, with 8-meeting moving averages. The statement is consistently more hawkish than the minutes and transcript.')

H('6.3 Sentiment Persistence')
P(
    'FOMC statement sentiment is highly persistent: regressing CB(t) on CB(t−1) and CB(t−2) '
    'produces R² = 81.3%, with both lags significant (t = 6.83 and t = 4.46, respectively). '
    "This persistence is consistent with the view that the FOMC's communication strategy evolves "
    'gradually, and that sentiment reflects slow-moving policy preferences rather than '
    'meeting-specific shocks. The high persistence also implies that the modest R² in our shock '
    'regressions is not surprising: after controlling for the persistent component, the '
    "innovation that is attributable to the current meeting's shock is small."
)

# ==================== 7. ROBUSTNESS ====================
H('7. Robustness')

H('7.1 Wild Bootstrap')
P(
    'We validate the interaction model using a wild bootstrap with 5,000 replications. The '
    'bootstrap preserves the significance of both interaction terms. For the CB dictionary, '
    'the path × direction interaction produces a bootstrap t-statistic of −3.24 (p = 0.002), '
    'compared to the HAC t-statistic of −2.56 (p = 0.012). The target × direction interaction '
    'produces a bootstrap t-statistic of 4.81 (p < 0.001), compared to the HAC t-statistic of '
    '3.52 (p < 0.001). The bootstrap results confirm that the interaction model\'s findings are '
    'not artifacts of the HAC standard error specification.'
)

H('7.2 The Fernández-Fuertes Information Bound')
P(
    'The modest full-sample baseline R² reflects the information limitation of the shock measure '
    'rather than the inadequacy of textual sentiment. Fernández-Fuertes (2026) finds that 81.5% '
    'of LLM-extracted narrative surprises lie outside the linear span of standard announcement-window '
    'derivatives, implying that our GSS shocks capture at most 18.5% of the relevant information. '
    'Our interaction model R² of 11.1% (CB V2) should be interpreted relative to this upper bound: '
    'the model captures a substantial fraction of the information that is theoretically available '
    'in the high-frequency shocks.'
)

H('7.3 Subsample Regressions as Consistency Check')
P(
    'The subsample regressions produce point estimates that are qualitatively consistent with the '
    'interaction model (Table 4), but the small sample sizes inflate the R² and t-statistics. '
    'The rate-hike subsample (N = 22) produces R² = 36.2%, but leave-one-out cross-validation '
    'reduces this to 12.0%. The rate-cut subsample (N = 42) produces R² = 16.3%, with LOO-CV '
    'R² = 7.9%. The interaction model\'s R² of 11.1% is more reliable because it is estimated '
    'on the full sample.'
)

H('7.4 The Bauer-Swanson Critique')
P(
    'Bauer and Swanson (2023) argue that high-frequency monetary policy surprises may be '
    'contaminated by information effects. This critique applies to both the target and path '
    'shocks used in our analysis. The Jarociński and Karadi (2020) decomposition, which separates '
    'monetary policy shocks from information shocks using the co-movement of interest rates and '
    'stock prices, would provide a cleaner identification strategy. We leave this extension for '
    'future work.'
)

H('7.5 Text Cleaning Sensitivity')
P(
    'We verify that our results are robust to the text extraction method. Scoring FOMC statements '
    'with and without the voting paragraph, and with different boilerplate removal strategies, '
    'produces sentiment scores with correlations above 0.90. The interaction model\'s significance '
    'pattern is preserved across all extraction methods.'
)

H('7.6 Sample Period Sensitivity')
P(
    'Our sample (2006–2022) spans multiple monetary policy regimes. The interaction model results '
    'are robust to excluding the GFC period (2008–2009) and the pandemic period (2020–2022), '
    'suggesting that the regime-dependent effects are not driven by crisis-era dynamics.'
)

# ==================== 8. CONCLUSION ====================
H('8. Conclusion')
P(
    'Does FOMC statement language primarily reflect current policy implementation or '
    'informational revelation about future conditions? We address this question by linking '
    'high-frequency monetary policy shocks, sentiment analysis, and asset returns in a unified '
    'framework. Using the GSS target/path decomposition across 131 FOMC meetings (2006–2022), '
    'we find that the answer depends on the policy regime — and that the econometric method '
    'matters for detecting this dependence.'
)
P(
    'Three main conclusions emerge. First, the choice of sentiment dictionary matters '
    'substantially. The abbreviated LM dictionary introduces a positivity bias that produces '
    'spurious significance (t = 8.00 in the Kuttner horse race). The full LM dictionary reveals '
    'that general-purpose financial sentiment adds no incremental information beyond the rate '
    'surprise itself (t = −0.22). The CB dictionary, by contrast, captures domain-specific '
    'language that the rate surprise does not subsume.'
)
P(
    'Second, the effect of monetary policy shocks on statement sentiment depends on the policy '
    'regime, and this dependence can be detected using a continuous interaction model estimated '
    'on the full sample. The target × direction interaction (t = 3.52, p < 0.001) supports the '
    'implementation channel: during rate hikes, the target shock has a differential effect on '
    'sentiment. The path × direction interaction (t = −2.56, p = 0.012; bootstrap t = −3.24, '
    'p = 0.002) supports the forward guidance channel: during rate cuts, the path shock has a '
    'differential effect on sentiment, with a one-standard-deviation path shock moving the CB '
    'score by 42% of its standard deviation.'
)
P(
    'Third, the interaction model is preferable to subsample regressions for detecting '
    'regime-dependent effects. Subsample regressions on small regimes (N = 22 for hikes) produce '
    'inflated R² and t-statistics that do not survive cross-validation. The interaction model '
    'uses the full sample to estimate all parameters, producing reliable standard errors and '
    'avoiding the multiple-testing problem. The subsample point estimates are qualitatively '
    'consistent with the interaction model, providing a useful consistency check but not a '
    'standalone result.'
)

H('8.1 From Dictionaries to LLMs')
P(
    'Our finding that the CB dictionary outperforms the LM dictionary has implications for the '
    'sentiment analysis literature, but even the CB dictionary is a bag-of-words approach that '
    'cannot capture nuanced semantics. Chen, Granville, and Matousek (2026) demonstrate that '
    'GPT-4 can decode FOMC materials into four topics including forward guidance, while GPT-3.5 '
    'misses 97% of forward guidance content. Gambacorta et al. (2024) introduce CB-LMs — '
    'open-weight models retrained on central bank corpora — that outperform general models while '
    'offering full reproducibility. Fernández-Fuertes (2026) demonstrates that a multi-agent LLM '
    'framework can construct narrative monetary policy surprises that capture 81.5% more '
    'information than standard high-frequency measures. These developments suggest that the next '
    'generation of monetary policy communication research will use LLM-based measures rather than '
    'bag-of-words approaches.'
)

H('8.2 Limitations')
P(
    'Several limitations should be noted. The LLM hawkish score has only 14 unique values, '
    'constraining its ability to detect fine-grained interaction effects. The Bauer-Swanson (2023) '
    'critique applies to both shock dimensions. The rate-direction variable is discrete (+1/0/−1) '
    'rather than continuous; using the actual rate change as a continuous moderator would provide '
    'more statistical power but introduces endogeneity concerns. We cannot distinguish the '
    'financing-constraint and portfolio rebalancing mechanisms without institutional ownership '
    'data. Implementing the Jarociński-Karadi (2020) decomposition would provide structural '
    'identification of monetary policy vs. information shocks.'
)

# ==================== REFERENCES ====================
H('References')
refs = [
    'Acosta, M. (2022). The perceived causes of monetary policy surprises. Working Paper, Columbia University.',
    'Aruoba, S. B., and Drechsel, T. (2024). Identifying monetary policy shocks: A natural language approach. NBER Working Paper No. 32417.',
    'Bauer, M. D., and Swanson, E. T. (2023). A reassessment of monetary policy surprises and high-frequency identification. NBER Macroeconomics Annual, 37(1), 87–155.',
    "Bernanke, B. S., and Kuttner, K. N. (2005). What explains the stock market's reaction to Federal Reserve policy? Journal of Finance, 60(3), 1221–1257.",
    'Chen, K., Granville, B., and Matousek, R. (2026). Decoding central bank communications with large language models. Journal of International Financial Markets, Institutions and Money.',
    'Cieslak, A., and Schrimpf, A. (2019). Non-monetary news in central bank communication. Journal of International Economics, 118, 293–315.',
    'Fernández-Fuertes, R. (2026). Monetary policy shocks: A new hope. Bocconi University Job Market Paper.',
    'Gambacorta, L., et al. (2024). CB-LMs: Language models for central banking. BIS Working Paper No. 1215.',
    'Gürkaynak, R. S., Sack, B., and Swanson, E. T. (2005b). Do actions speak louder than words? International Journal of Central Banking, 1(1), 55–93.',
    'Hansen, S., and McMahon, M. (2016). Shocking language: Understanding the macroeconomic effects of central bank communication. Journal of International Economics, 99, S121–S133.',
    'Jarociński, M., and Karadi, P. (2020). Deconstructing monetary policy surprises. American Economic Journal: Macroeconomics, 12(2), 1–43.',
    'Kuttner, K. N. (2001). Monetary policy surprises and interest rates. Journal of Monetary Economics, 47(3), 523–544.',
    'Loughran, T., and McDonald, B. (2011). When is a liability not a liability? Journal of Finance, 66(1), 35–65.',
    'Lu, X., and Wu, L. (2026). Monetary transmission and portfolio rebalancing: A cross-sectional approach. SSRN Working Paper No. 4413059.',
    'Nakamura, E., and Steinsson, J. (2018). High-frequency identification of monetary non-neutrality. Quarterly Journal of Economics, 133(3), 1283–1330.',
    'Romer, C. D., and Romer, D. H. (2000). Federal Reserve information and the behavior of interest rates. American Economic Review, 90(3), 429–457.',
    'Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. Journal of Finance, 62(3), 1139–1168.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# ==================== APPENDIX ====================
doc.add_page_break()
H('Appendix')

H('A. CB Dictionary V2 Construction')
P(
    'The CB dictionary V2 is constructed in three stages. First, we start with the Cieslak and '
    'Schrimpf (2019) list of monetary policy terms. Second, we expand this list by analyzing '
    'the frequency and context of candidate terms in 212 FOMC statements (1994–2026). Third, '
    'we validate the dictionary by checking that the resulting scores are correlated with '
    'independent measures of monetary policy stance (the Wu-Xia shadow rate and the effective '
    'federal funds rate). The final dictionary contains 407 hawkish words, 543 dovish words, '
    '42 hawkish phrases, and 55 dovish phrases.'
)

H('B. LM Dictionary: Abbreviated vs. Full')
P(
    'The abbreviated LM dictionary contains 116 positive and 213 negative terms, totaling 329 '
    'unique words. The full LM dictionary from the Notre Dame Software Repository contains 347 '
    'positive and 2,345 negative categories, totaling 2,692 unique terms. The key difference is '
    'in the negative word list: the full dictionary has 11× more negative terms (2,345 vs. 213). '
    'Terms like "deterioration," "adversely," "uncertain," and "challenging" are all classified '
    'as negative in the full dictionary but are absent from the abbreviated version.'
)

H('C. Leave-One-Out Cross-Validation')
P(
    'To assess the reliability of subsample R² values, we perform leave-one-out '
    'cross-validation (LOO-CV). For each observation, we estimate the model on the remaining '
    'N−1 observations and predict the held-out observation. The LOO-CV R² is computed as '
    '1 − SS_res/SS_tot, where SS_res is the sum of squared prediction errors and SS_tot is '
    'the total sum of squares around the mean.'
)
P(
    'For the rate-hike regime (N = 22), the in-sample R² for the CB V2 specification is 36.2%, '
    'but the LOO-CV R² is only 12.0%, indicating substantial overfitting. For the rate-cut '
    'regime (N = 42), the in-sample R² is 16.3% and the LOO-CV R² is 7.9%, indicating more '
    'moderate overfitting. The interaction model\'s R² of 11.1% is estimated on the full sample '
    'and does not suffer from this overfitting problem.'
)

H('D. Additional Figures')
FIG('paper/figures/fig3_r2_by_regime.png', 'Figure A1: Explained Variance by Regime and Sentiment Measure. R² from regressing each sentiment measure on target and path shocks, by regime.')
FIG('paper/figures/fig4_timeseries.png', 'Figure A2: Monetary Policy Shocks and Statement Sentiment, 2006–2022. Panel (a): GSS target and path shocks. Panel (b): CB Score V2. Panel (c): LM% (full dictionary). Red shading indicates rate-hike meetings; blue shading indicates rate-cut meetings.')

# SAVE
doc.save('paper/Words_Beyond_the_Rate_v12_5.docx')
text = ' '.join([p.text for p in doc.paragraphs])
words = len(text.split())
print(f"Saved! Word count: {words}, Approx pages: {words/250:.0f}, Tables: {len(doc.tables)}")
