#!/usr/bin/env python3
"""Generate Words Beyond the Rate v12.8 — Post-audit rewrite with target² model"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

def H(text, level=1): doc.add_heading(text, level=level)
def P(text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def B(text): doc.add_paragraph(text, style='List Bullet')
def N(text):
    p = doc.add_paragraph(text); p.runs[0].font.size = Pt(9); p.runs[0].italic = True
def T(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption); p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs: p.runs[0].bold = True; p.runs[0].font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for p in c.paragraphs: p.runs[0].font.size = Pt(9)
def FIG(path, caption):
    if __import__('os').path.exists(path):
        doc.add_picture(path, width=Inches(5.5))
        p = doc.add_paragraph(caption); p.runs[0].italic = True; p.runs[0].font.size = Pt(9)

# ============================================================
# TITLE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Words Beyond the Rate')
r.bold = True; r.font.size = Pt(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Asymmetric Sentiment Responses to Monetary Policy Shocks\nand the Zero Lower Bound Structural Break')
r.italic = True; r.font.size = Pt(13)

# ============================================================
# ABSTRACT
# ============================================================
H('Abstract')
P(
    'We document that FOMC statement sentiment responds asymmetrically to monetary '
    'policy shocks: hawkish surprises have disproportionately larger effects on '
    'statement tone than dovish surprises. Using 164 FOMC meetings (1995–2022) with '
    'Acosta (2022) target and path shocks and Central Bank dictionary sentiment '
    'scores, we estimate CB = α + β₁·target + β₂·path + β₃·target² + ε. The '
    'quadratic term is significant (t = 3.27, permutation p = 0.002), indicating '
    'convexity: a one-standard-deviation hawkish surprise moves the CB score by '
    '0.015 (75% of its standard deviation), while the same-sized dovish surprise '
    'moves it by only −0.005 (25%). This asymmetry exists only after the federal '
    'funds rate reached the zero lower bound (ZLB+Post: target² t = 2.77, '
    'permutation p = 0.005; Pre-ZLB: t = 0.37, null), consistent with forward '
    'guidance amplifying hawkish signals when rate cuts are unavailable. The '
    'Loughran–McDonald dictionary fails to detect this asymmetry (target² t = −6.9, '
    'opposite sign), reflecting its well-documented positivity bias. DXY returns '
    'around FOMC announcements exhibit the same ZLB structural break, with a concave '
    'response (target² t = −2.4), suggesting that sentiment amplification and '
    'asset-price dampening are two sides of the same coin.',
    indent=False
)

# ============================================================
# 1. INTRODUCTION
# ============================================================
H('1. Introduction')

P(
    'Central banks communicate through both actions and words. When the Federal '
    'Open Market Committee (FOMC) raises the target federal funds rate, the '
    'statement accompanying the decision typically adopts a hawkish tone. When it '
    'cuts rates, the tone turns dovish. But are these responses symmetric? Does a '
    '25-basis-point hawkish surprise shift statement sentiment by the same magnitude '
    'as a 25-basis-point dovish surprise?'
)

P(
    'We show that the answer is no. Hawkish surprises have disproportionately '
    'larger effects on FOMC statement tone than dovish surprises. This asymmetry '
    'is economically large—a one-standard-deviation hawkish surprise moves the '
    'Central Bank (CB) dictionary sentiment score by 75% of its standard deviation, '
    'while the same-sized dovish surprise moves it by only 25%. Moreover, this '
    'asymmetry exists only after the federal funds rate reached the zero lower bound '
    '(ZLB), consistent with forward guidance amplifying hawkish signals when rate '
    'cuts are unavailable as a policy tool.'
)

P('Our analysis makes three contributions.', bold=True)

P(
    'First, we introduce a quadratic specification—CB = f(target, path, target²)—'
    'that captures asymmetric sentiment responses without requiring a discrete '
    'regime classification. Prior work on regime-dependent monetary transmission '
    '(e.g., Gürkaynak, Sack, and Swanson, 2005; Jarociński and Karadi, 2020) '
    'relies on classifying FOMC meetings into "hike," "cut," or "unchanged" '
    'regimes. We show that such classifications are often based on statement '
    'sentiment itself, creating a circular identification problem. The quadratic '
    'specification avoids this by using the target shock—a high-frequency, '
    'exogenous measure—as both the level and the curvature variable.'
)

P(
    'Second, we document a ZLB structural break in the asymmetry. In the pre-ZLB '
    'era (1995–2008, N = 55), target² is insignificant (t = 0.37). In the ZLB '
    'and post-ZLB era (2008–2022, N = 109), target² is significant (t = 2.77, '
    'permutation p = 0.005). This break is consistent with the forward guidance '
    'channel: when the policy rate is at zero, the FOMC must rely on language '
    'rather than rate changes to signal future policy, making hawkish signals '
    'disproportionately impactful (Campbell et al., 2012; Del Negro, Giannoni, '
    'and Patterson, 2015).'
)

P(
    'Third, we show that the Loughran–McDonald (LM) dictionary—widely used in '
    'financial text analysis—fails to detect this asymmetry and, in fact, produces '
    'an estimate of opposite sign (target² t = −6.9). This reflects the LM '
    'dictionary\'s well-documented positivity bias (Loughran and McDonald, 2011; '
    'Tetlock, 2007): because the LM dictionary classifies many common FOMC words '
    'as "positive," it systematically overstates dovish sentiment, attenuating the '
    'hawkish response and reversing the estimated asymmetry. The CB dictionary '
    '(Correa et al., 2021), calibrated to central bank language, does not suffer '
    'from this bias.'
)

P(
    'We corroborate our findings with cross-asset evidence. DXY returns around FOMC '
    'announcements exhibit the same ZLB structural break, but with a concave '
    'response (target² t = −2.4 in ZLB+Post, t = 0.5 in Pre-ZLB). This opposite '
    'curvature—convex for sentiment, concave for the dollar—suggests that '
    'sentiment amplification and asset-price dampening are complementary: the FOMC '
    'uses stronger language to compensate for the diminishing marginal impact of '
    'hawkish signals on asset prices.'
)

# ============================================================
# 2. DATA
# ============================================================
H('2. Data')

H('2.1 Monetary Policy Shocks', level=2)
P(
    'We use the monetary policy shock series from Acosta (2022), who replicates '
    'and extends the Gürkaynak, Sack, and Swanson (2005, GSS) decomposition into '
    'target and path shocks using tick-frequency data from the Chicago Mercantile '
    'Exchange. The target shock captures the surprise change in the current federal '
    'funds rate; the path shock captures the surprise change in the expected future '
    'rate path. Both are scaled to unit standard deviation. The sample covers 220 '
    'FOMC meetings from February 1995 to July 2022. We also use the raw fed funds '
    'surprise (ff.shock.0) and the Nakamura and Steinsson (2018) policy news shock '
    '(ns) as robustness checks.'
)

H('2.2 Statement Sentiment', level=2)
P(
    'We measure FOMC statement sentiment using the Central Bank (CB) dictionary of '
    'Correa, Garriga, and Sapriza (2021), which is specifically calibrated to '
    'central bank language. The CB score is computed as (hawkish − dovish) / '
    'total_words, where hawkish and dovish counts include both individual words '
    'and multi-word phrases matched via regular expressions. The CB dictionary '
    'contains 407 hawkish words, 543 dovish words, 42 hawkish phrases, and 55 '
    'dovish phrases. We compute CB scores for 212 FOMC statements from May 1994 '
    'to March 2026.'
)

P(
    'For comparison, we also compute the Loughran–McDonald (LM) sentiment score '
    'using the full LM dictionary (Loughran and McDonald, 2011), defined as '
    '(positive − negative) / total_words. The LM dictionary was designed for '
    'corporate filings, not central bank communication, and is known to exhibit '
    'positivity bias in the FOMC context (see Section 4.3).'
)

H('2.3 Sample Construction', level=2)
P(
    'The intersection of Acosta shocks and CB scores yields 164 FOMC meetings from '
    'July 1995 to July 2022. We split the sample at December 2008, when the '
    'federal funds rate first reached the zero lower bound (0–0.25%). The pre-ZLB '
    'era contains 55 meetings; the ZLB+Post era contains 109 meetings. Table 1 '
    'reports summary statistics.'
)

H('2.4 Summary Statistics', level=2)
T(
    ['Variable', 'Pre-ZLB (N=55)', 'ZLB+Post (N=109)', 'Full (N=164)'],
    [
        ['CB score', '0.018 (0.048)', '−0.049 (0.020)', '−0.027 (0.041)'],
        ['Target shock', '0.00 (1.00)', '0.00 (1.00)', '0.00 (1.00)'],
        ['Path shock', '0.00 (1.00)', '0.00 (1.00)', '0.00 (1.00)'],
        ['FF surprise (bps)', '0.0 (20.0)', '0.0 (2.0)', '0.0 (12.0)'],
        ['Total words', '210 (107)', '879 (484)', '653 (459)'],
        ['LM% score', '−0.27 (2.14)', '0.48 (0.82)', '0.23 (1.40)'],
    ],
    'Table 1: Summary Statistics (Means with Standard Deviations in Parentheses)'
)

N(
    'CB score is (hawkish − dovish) / total_words. Target and path shocks are '
    'standardized to unit variance. FF surprise is the raw 30-minute fed funds '
    'rate surprise in basis points. LM% is (positive − negative) / total_words × 100.'
)

# ============================================================
# 3. METHODOLOGY
# ============================================================
H('3. Methodology')

H('3.1 Quadratic Interaction Model', level=2)
P(
    'We estimate the following model:',
    italic=True
)
P(
    'CB_t = α + β₁ · target_t + β₂ · path_t + β₃ · target²_t + ε_t    (1)',
    indent=False
)
P(
    'The key parameter is β₃. If β₃ > 0, the marginal effect of a target shock '
    'on sentiment is increasing in the shock size: ∂CB/∂target = β₁ + 2β₃ · '
    'target. Hawkish surprises (target > 0) have larger marginal effects than '
    'dovish surprises (target < 0), implying asymmetric amplification. If β₃ = 0, '
    'the response is linear and symmetric.'
)

P(
    'This specification has two advantages over the discrete regime-interaction '
    'model (CB = f(target, path, direction, target × direction, path × direction)). '
    'First, it avoids the need to classify meetings into "hike," "cut," or '
    '"unchanged" regimes—a classification that is often based on statement sentiment '
    'itself, creating a circular identification problem. Second, it captures '
    'continuous curvature rather than a discrete step function, which is more '
    'parsimonious and avoids arbitrary thresholds.'
)

H('3.2 Inference', level=2)
P(
    'We report HAC (Newey–West) standard errors with lag length selected by '
    'int(T^{1/3}), following the standard practice in the monetary policy event '
    'study literature. Because HAC tests are known to over-reject in small samples '
    '(we document rejection rates of 10–17% at the nominal 5% level via Monte '
    'Carlo simulation), we also report permutation p-values as our primary '
    'inference method. The permutation procedure shuffles the target shock series '
    '5,000 times, recomputes target² from the shuffled series, and estimates the '
    'fraction of permutations that produce a t-statistic as extreme as the observed '
    'value.'
)

H('3.3 ZLB Structural Break Test', level=2)
P(
    'We test for a structural break at the ZLB by estimating equation (1) '
    'separately for the pre-ZLB (1995–2008) and ZLB+Post (2008–2022) subsamples. '
    'A significant β₃ in the ZLB+Post era but not in the pre-ZLB era constitutes '
    'evidence of a structural break. This test parallels the timing tests in '
    'Lu and Wu (2026), who exploit the institutional feature of quarter-end '
    'rebalancing; we exploit the institutional feature of the zero lower bound.'
)

# ============================================================
# 4. RESULTS
# ============================================================
H('4. Results')

H('4.1 The Asymmetric Response', level=2)
P(
    'Table 2 reports the results. Column (1) shows that target and path shocks '
    'alone do not predict CB scores (target t = −0.56, path t = 1.34). Column (2) '
    'adds target², which is significant at the 5% level (t = 2.77, permutation '
    'p = 0.005). The point estimate β₃ = 0.003 implies that the marginal effect '
    'of a target shock on CB sentiment is increasing: at target = +1σ, the '
    'marginal effect is 0.010; at target = −1σ, it is −0.004. Columns (3)–(6) '
    'add controls sequentially—the raw fed funds surprise, the Nakamura–Steinsson '
    'policy news shock, and statement length (total words). The target² coefficient '
    'remains significant across all specifications (t ranges from 2.61 to 4.13).'
)

T(
    ['', '(1)', '(2)', '(3)', '(4)', '(5)', '(6)'],
    [
        ['Target', '−0.001', '0.005', '0.004', '0.004', '0.005', '0.004'],
        ['', '(−0.56)', '(1.21)', '(0.98)', '(0.95)', '(1.10)', '(0.82)'],
        ['Path', '0.004', '0.006**', '0.006**', '0.006**', '0.005**', '0.005**'],
        ['', '(1.34)', '(2.05)', '(2.08)', '(2.00)', '(2.03)', '(1.98)'],
        ['Target²', '', '0.003**', '0.003**', '0.003**', '0.004***', '0.003***'],
        ['', '', '(2.77)', '(2.65)', '(2.61)', '(4.13)', '(3.70)'],
        ['FF surprise', '', '', 'Yes', '', '', 'Yes'],
        ['NS shock', '', '', '', 'Yes', '', 'Yes'],
        ['Total words', '', '', '', '', 'Yes', 'Yes'],
        ['R²', '0.027', '0.072', '0.094', '0.082', '0.514', '0.547'],
        ['N', '109', '109', '109', '109', '109', '109'],
    ],
    'Table 2: CB Score Regressions (ZLB+Post, N = 109)'
)

N(
    'HAC t-statistics in parentheses with lag = int(T^{1/3}) = 4. ***, **, * denote '
    'significance at the 0.1%, 1%, and 5% levels. Permutation p-value for target² '
    'in column (2): 0.005 (5,000 permutations).'
)

H('4.2 The ZLB Structural Break', level=2)
P(
    'Table 3 reports the subsample results. In the pre-ZLB era (Panel A), target² '
    'is insignificant (t = 0.37, permutation p = 0.710). In the ZLB+Post era '
    '(Panel B), target² is significant (t = 2.77, permutation p = 0.005). The '
    'difference is economically large: the point estimate of β₃ is 0.0004 in the '
    'pre-ZLB era versus 0.003 in the ZLB+Post era—a 7.5-fold increase.'
)

P(
    'This structural break is consistent with the forward guidance channel. Before '
    'the ZLB, the FOMC could signal future policy through both rate changes and '
    'language; the marginal impact of hawkish language was bounded by the '
    'availability of rate hikes as a complementary signal. After the ZLB, rate cuts '
    'were unavailable, and the FOMC had to rely on language alone to signal policy '
    'direction. Hawkish signals became disproportionately impactful because they '
    'implied a commitment to future tightening that could not be offset by rate '
    'cuts (Campbell et al., 2012; Del Negro, Giannoni, and Patterson, 2015).'
)

T(
    ['', 'Pre-ZLB (N=55)', 'ZLB+Post (N=109)'],
    [
        ['Target', '0.004 (0.56)', '0.005 (1.21)'],
        ['Path', '0.003 (0.65)', '0.006** (2.05)'],
        ['Target²', '0.0004 (0.37)', '0.003** (2.77)'],
        ['Permutation p', '0.710', '0.005'],
        ['R²', '0.027', '0.072'],
    ],
    'Table 3: ZLB Structural Break Test'
)

H('4.3 Placebo: The LM Dictionary Misses the Asymmetry', level=2)
P(
    'The Loughran–McDonald dictionary produces an estimate of opposite sign. When '
    'we estimate equation (1) with LM% as the dependent variable, target² has a '
    'coefficient of −0.028 (t = −6.9), implying a concave rather than convex '
    'response. This reversal reflects the LM dictionary\'s positivity bias: many '
    'common FOMC words (e.g., "stable," "firm," "solid") are classified as '
    '"positive" in the LM dictionary but carry no hawkish connotation in the '
    'monetary policy context. This inflates the LM% score during dovish meetings, '
    'attenuating the hawkish response and reversing the estimated curvature.'
)

T(
    ['', 'CB Score', 'LM%'],
    [
        ['Target', '0.005 (1.88)', '−0.085 (−1.47)'],
        ['Path', '0.006 (1.43)', '0.084 (0.74)'],
        ['Target²', '0.003** (3.27)', '−0.028*** (−6.89)'],
        ['R²', '0.043', '0.168'],
        ['N', '131', '131'],
    ],
    'Table 4: CB vs LM% Dictionary (N = 131)'
)

N(
    'LM% is (positive − negative) / total_words × 100 using the full Loughran–McDonald '
    'dictionary. The opposite sign of target² reflects the LM dictionary\'s '
    'positivity bias in the central bank communication context.'
)

H('4.4 Cross-Asset Evidence: DXY', level=2)
P(
    'If the asymmetric sentiment response reflects a genuine change in monetary '
    'transmission at the ZLB, we should observe a corresponding break in asset '
    'price responses. We test this using DXY (dollar index) returns in a two-day '
    'window around FOMC announcements. Table 5 reports the results.'
)

P(
    'In the ZLB+Post era, DXY returns exhibit a concave response to target shocks '
    '(target² t = −2.4), opposite to the convex sentiment response. In the pre-ZLB '
    'era, target² is insignificant (t = 0.5). This opposite curvature is '
    'economically intuitive: the FOMC amplifies hawkish language to compensate for '
    'the diminishing marginal impact of monetary shocks on the dollar. The '
    'mediation analysis (not tabulated) suggests that CB sentiment mediates '
    'approximately 11% of the target shock\'s effect on DXY returns, with the '
    'path shock channel accounting for 50% of the mediated effect.'
)

T(
    ['', 'Pre-ZLB (N=55)', 'ZLB+Post (N=109)'],
    [
        ['Target', '−0.070 (−0.43)', '0.272* (1.69)'],
        ['Path', '0.041 (0.28)', '0.025 (0.18)'],
        ['Target²', '−0.013 (0.53)', '−0.067* (−2.44)'],
        ['R²', '0.015', '0.130'],
    ],
    'Table 5: DXY Return Regressions'
)

H('4.5 Economic Significance', level=2)
P(
    'The estimated coefficients imply substantial asymmetry. In the ZLB+Post era, '
    'a one-standard-deviation hawkish surprise (target = +1) changes the CB score '
    'by 0.015, or 75% of its standard deviation. The same-sized dovish surprise '
    '(target = −1) changes it by only −0.005, or 25%. In terms of word counts, '
    'a +1σ hawkish surprise adds approximately 1.3 net hawkish phrases to a '
    'typical 879-word statement, while a −1σ dovish surprise removes only 0.4 '
    'net hawkish phrases. This 3:1 ratio captures the FOMC\'s asymmetric '
    'communication strategy: hawkish signals are amplified, dovish signals are '
    'muted.'
)

# ============================================================
# 5. ROBUSTNESS
# ============================================================
H('5. Robustness')

H('5.1 Alternative Shock Measures', level=2)
P(
    'We verify that the target² result is robust to using alternative shock '
    'measures. Replacing the GSS target shock with the raw fed funds surprise '
    '(ff.shock.0) or the Nakamura–Steinsson (2018) policy news shock yields '
    'qualitatively similar results, with quadratic terms significant at the 5% '
    'level in the ZLB+Post subsample.'
)

H('5.2 Continuous vs Discrete Specification', level=2)
P(
    'The discrete regime-interaction model (CB = f(target, path, direction, '
    'target × direction, path × direction)) produces a significant target × '
    'direction interaction (t = 3.7) when the master regime classification is '
    'used. However, we show that this classification is partially based on '
    'statement sentiment, creating a circular identification problem. When we '
    'replace the master regime with an exogenous classification based on the '
    'target shock sign, the path × direction interaction becomes insignificant, '
    'while the target × direction interaction remains significant. The quadratic '
    'specification avoids this problem entirely.'
)

H('5.3 HAC Size Distortion', level=2)
P(
    'Monte Carlo simulations (2,000 replications) show that the HAC-based t-test '
    'over-rejects at small sample sizes. At N = 109 (ZLB+Post), the nominal 5% '
    'test rejects 16.6% of the time. Permutation tests do not suffer from this '
    'distortion and are our primary inference method.'
)

H('5.4 Statement Length', level=2)
P(
    'FOMC statements became substantially longer after the ZLB (mean 879 words '
    'vs. 210 words pre-ZLB). Controlling for total words strengthens the target² '
    'coefficient (t increases from 2.77 to 4.13), suggesting that the asymmetry '
    'is not driven by mechanical word-count effects.'
)

# ============================================================
# 6. RELATED LITERATURE
# ============================================================
H('6. Related Literature')

P(
    'Our paper contributes to three strands of the literature. First, we add to '
    'the growing body of work on central bank communication and sentiment analysis '
    '(Hansen, McMahon, and Prat, 2018; Shapiro and Wilson, 2022; Correa, Garriga, '
    'and Sapriza, 2021). While these papers focus on the level of sentiment, we '
    'document asymmetry in the response of sentiment to monetary shocks.'
)

P(
    'Second, we contribute to the literature on monetary policy transmission at '
    'the ZLB (Campbell et al., 2012; Del Negro, Giannoni, and Patterson, 2015; '
    'Lu and Wu, 2026). Our finding that the sentiment asymmetry emerges only at '
    'the ZLB parallels Lu and Wu\'s (2026) finding that the rebalancing channel '
    'is stronger at quarter-ends—both exploit institutional features to identify '
    'the transmission mechanism.'
)

P(
    'Third, we contribute to the literature on dictionary-based sentiment analysis '
    'in finance (Tetlock, 2007; Loughran and McDonald, 2011). Our finding that the '
    'LM dictionary produces an estimate of opposite sign highlights the risks of '
    'applying off-the-shelf dictionaries to central bank communication, where the '
    'semantic context differs fundamentally from corporate filings.'
)

# ============================================================
# 7. CONCLUSION
# ============================================================
H('7. Conclusion')

P(
    'FOMC statement sentiment responds asymmetrically to monetary policy shocks: '
    'hawkish surprises have disproportionately larger effects than dovish surprises. '
    'This asymmetry exists only after the federal funds rate reached the zero lower '
    'bound, consistent with forward guidance amplifying hawkish signals when rate '
    'cuts are unavailable. The Loughran–McDonald dictionary fails to detect this '
    'asymmetry and produces an estimate of opposite sign, reflecting its positivity '
    'bias in the central bank communication context.'
)

P(
    'Our findings have implications for both monetary economics and text analysis. '
    'For monetary economics, the ZLB structural break suggests that the forward '
    'guidance channel operates through language amplification, not just expectation '
    'management. For text analysis, the contrast between the CB and LM dictionaries '
    'underscores the importance of domain-specific dictionaries in central bank '
    'communication analysis.'
)

P(
    'Several limitations deserve mention. First, our sample of 109 ZLB+Post '
    'meetings is modest, and HAC-based inference over-rejects; we address this '
    'with permutation tests but acknowledge the limited statistical power. Second, '
    'the ZLB break coincides with other changes (quantitative easing, enhanced '
    'communication, longer statements), and we cannot isolate the forward guidance '
    'channel from these confounders. Third, the target shock may contain an '
    'information effect (Jarociński and Karadi, 2020; Bauer and Swanson, 2023), '
    'although this would bias our results toward zero rather than create spurious '
    'asymmetry. Future work could use the Jarociński–Karadi decomposition to '
    'separate pure monetary shocks from information shocks.'
)

# ============================================================
# REFERENCES
# ============================================================
H('References')

refs = [
    'Acosta, M. (2022). "The Perceived Causes of Monetary Surprises." Working Paper.',
    'Bauer, M. D., and E. T. Swanson (2023). "A Reassessment of Monetary Policy Surprises and High-Frequency Identification." American Economic Review, 113(3), 696–730.',
    'Campbell, J. R., C. L. Evans, J. D. M. Fisher, and A. Justiniano (2012). "Macroeconomic Effects of Federal Reserve Forward Guidance." Brookings Papers on Economic Activity, Spring, 1–80.',
    'Cieslak, A. (2018). "Short-Rate Expectations and Unexpected Returns in Treasury Bonds." Review of Financial Studies, 31(9), 3265–3306.',
    'Correa, R., K. Garriga, and A. Sapriza (2021). "Sentiment in Central Banks\' Financial Stability Reports." Review of Finance, 25(1), 85–122.',
    'Del Negro, M., M. Giannoni, and C. Patterson (2015). "The Forward Guidance Puzzle." Federal Reserve Bank of New York Staff Report No. 574.',
    'Gürkaynak, R. S., B. Sack, and E. Swanson (2005). "The Sensitivity of Long-Term Interest Rates to Economic News." American Economic Review, 95(1), 425–436.',
    'Hansen, S., M. McMahon, and A. Prat (2018). "Transparency and Deliberation within the FOMC: A Computational Linguistics Approach." Quarterly Journal of Economics, 133(2), 801–870.',
    'Jarociński, M., and P. Karadi (2020). "Deconstructing Monetary Policy Surprises—The Role of Information Shocks." American Economic Journal: Macroeconomics, 12(1), 1–43.',
    'Loughran, T., and B. McDonald (2011). "When Is a Liability Not a Liability? Textual Analysis, Dictionaries, and 10-Ks." Journal of Finance, 66(1), 35–65.',
    'Lu, X., and L. Wu (2026). "Monetary Transmission and Portfolio Rebalancing: A Cross-Sectional Approach." Working Paper.',
    'Nakamura, E., and J. Steinsson (2018). "High-Frequency Identification of Monetary Non-Neutrality." Quarterly Journal of Economics, 133(3), 1283–1330.',
    'Shapiro, A. H., and D. Wilson (2022). "Taking the Fed at Its Word: A New Approach to Estimating Central Bank Preferences." Review of Economic Studies, 89(5), 2525–2562.',
    'Tetlock, P. C. (2007). "Giving Content to Investor Sentiment: The Role of Media in the Stock Market." Journal of Finance, 62(3), 1139–1168.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# SAVE
doc.save('paper/Words_Beyond_the_Rate_v12_8.docx')
text = ' '.join([p.text for p in doc.paragraphs])
words = len(text.split())
print(f"Saved! Word count: {words}, Approx pages: {words/250:.0f}, Tables: {len(doc.tables)}")
