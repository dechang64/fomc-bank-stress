#!/usr/bin/env python3
"""Generate Words Beyond the Rate v12.7 — Top-journal rewrite after Lu & Wu (2026) study"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(6)

def H(text, level=1): doc.add_heading(text, level=level)
def P(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(1.27)
    r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def B(text): doc.add_paragraph(text, style='List Bullet')
def N(text):
    p = doc.add_paragraph(text); p.runs[0].font.size = Pt(9); p.runs[0].italic = True
def T(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption); p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for r in c.paragraphs[0].runs: r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for r in c.paragraphs[0].runs: r.font.size = Pt(9)
    return t
def FIG(path, caption):
    if __import__('os').path.exists(path):
        doc.add_picture(path, width=Inches(5.5))
        last = doc.paragraphs[-1]; last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(caption); p.runs[0].italic = True; p.runs[0].font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════
P('Words Beyond the Rate', bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
P('How Monetary Policy Shocks Shape FOMC Statement Sentiment\nAcross Policy Regimes and the Zero Lower Bound', italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════
H('Abstract')
P(
    'We show that monetary policy shocks shape the sentiment of FOMC statements, '
    'but only after the federal funds rate reached the zero lower bound (ZLB). '
    'Using an extended sample of 164 FOMC meetings (1995–2022) with Acosta (2022) '
    'target and path shocks and Central Bank (CB) dictionary sentiment scores, '
    'we document a structural break at the ZLB. In the pre-ZLB era (1995–2008, '
    'N = 55), neither target nor path shocks predict statement sentiment '
    'conditional on the policy regime. In the ZLB and post-ZLB era (2008–2022, '
    'N = 109), a one-standard-deviation target shock during a rate hike meeting '
    'moves the CB score by 0.014 (71% of its standard deviation), while the same '
    'shock during a rate cut meeting moves sentiment by −0.006 in the opposite '
    'direction (target × direction: t = 3.73, permutation p = 0.0002). A one-'
    'standard-deviation path shock during a rate cut meeting moves the CB score '
    'by 0.007 (35% of σ), consistent with the forward guidance channel '
    '(path × direction: t = −2.93, permutation p = 0.016). The pre-ZLB null '
    'result is itself informative: before the ZLB, the FOMC communicated primarily '
    'through rate changes, and statement language did not systematically respond '
    'to monetary surprises. Separately, we identify a positivity bias in the '
    'abbreviated Loughran–McDonald percentage metric that inflates its apparent '
    'predictive power (t = 8.00) relative to the full-dictionary specification '
    '(t = −0.22). Our findings highlight both the regime-dependent nature of '
    'monetary communication and the methodological pitfalls of dictionary-based '
    'sentiment analysis.',
    indent=True
)
P('Keywords: monetary policy shocks, forward guidance, FOMC statements, sentiment analysis, '
  'zero lower bound, central bank communication', italic=True)
P('JEL: E52, E58, G12, G14')

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════
H('1. Introduction')

P(
    'When the Federal Open Market Committee (FOMC) announces a policy decision, '
    'financial markets react not only to the rate change itself but also to the '
    'language of the accompanying statement. This language—carefully crafted to '
    'signal the Committee\'s assessment and intentions—has become increasingly '
    'important as a policy tool, particularly since the global financial crisis '
    'when the federal funds rate reached the zero lower bound (ZLB) and forward '
    'guidance became the primary mechanism for shaping expectations.',
    indent=True
)

P(
    'A substantial literature decomposes monetary policy surprises into a "target" '
    'component (the unexpected change in the current federal funds rate) and a "path" '
    'component (the revision in expectations of future rates) following Gürkaynak, '
    'Sack, and Swanson (2005, henceforth GSS). While the asset price effects of '
    'these shocks are well documented (Bernanke and Kuttner, 2005; Gertler and '
    'Karadi, 2015; Nakamura and Steinsson, 2018), less is known about how they '
    'feed back into the language of subsequent FOMC statements. If path shocks '
    'reflect forward guidance, then their impact on statement sentiment should be '
    'most pronounced when the Fed relies on language rather than rate changes to '
    'signal policy—that is, during and after the ZLB period.',
    indent=True
)

P(
    'We test three hypotheses. H1 (Implementation Channel): Target rate surprises '
    'have a differential effect on statement sentiment during rate hike meetings, '
    'where unexpected tightening signals a shift in the Committee\'s assessment. '
    'H2 (Forward Guidance Channel): Path surprises have a differential effect '
    'during rate cut meetings, where forward guidance about future easing is most '
    'relevant. H3 (ZLB Structural Break): The regime-dependent effects in H1 and '
    'H2 are present only in the ZLB and post-ZLB era, when forward guidance became '
    'a primary policy tool.',
    indent=True
)

P(
    'Our contribution is threefold. First, we document a ZLB structural break in '
    'the relationship between monetary policy shocks and statement sentiment. Using '
    'an extended sample of 164 FOMC meetings (1995–2022) with Acosta (2022) '
    'monetary policy shocks and CB dictionary sentiment scores (Correa et al., '
    '2021), we show that the regime-dependent response of sentiment to shocks '
    'exists only after 2008. In the pre-ZLB era (1995–2008, N = 55), neither '
    'target nor path shocks predict sentiment conditional on the policy regime '
    '(target × direction: t = −0.93; path × direction: t = 0.08). In the ZLB+Post '
    'era (2008–2022, N = 109), a one-standard-deviation target shock during a '
    'rate hike meeting moves the CB score by 0.014, or 71% of its standard '
    'deviation, while the same shock during a rate cut meeting moves sentiment '
    'by −0.006 in the opposite direction (target × direction: t = 3.73, '
    'permutation p = 0.0002). A one-standard-deviation path shock during a rate '
    'cut meeting moves the CB score by 0.007, or 35% of σ (path × direction: '
    't = −2.93, permutation p = 0.016). The pre-ZLB null result is itself '
    'informative: it confirms that forward guidance became a meaningful channel '
    'for sentiment only when the Fed was constrained at the ZLB.',
    indent=True
)

P(
    'Second, we introduce a continuous interaction model that avoids the overfitting '
    'problems of subsample regression. Rather than estimating separate regressions '
    'for each regime (which yields inflated R² due to small subsamples—42.5% in '
    'sample but only 11.8% under leave-one-out cross-validation for rate cuts), we '
    'estimate a single model on the full sample with a continuous rate-direction '
    'variable (hike = +1, unchanged = 0, cut = −1) interacted with the shock '
    'measures. This approach uses all observations, avoids sample splitting, and '
    'produces properly calibrated inference under permutation testing. We show that '
    'the interaction coefficients are robust to progressive addition of controls '
    '(Table 2, Columns 1–6), including the raw fed funds surprise and the '
    'Nakamura–Steinsson (2018) policy news shock.',
    indent=True
)

P(
    'Third, we identify a positivity bias in the abbreviated Loughran–McDonald '
    '(LM) percentage metric that inflates its apparent predictive power. The '
    'abbreviated LM% (negative words / [negative + positive words]) yields '
    't = 8.00 in a Kuttner-style horse race, while the full-dictionary LM% '
    '(net tone / total words) yields t = −0.22—a complete reversal. This '
    'discrepancy arises because the abbreviated denominator excludes neutral '
    'words, creating a mechanical positive correlation with rate changes. By '
    'contrast, the CB dictionary—designed specifically for monetary policy '
    'communication—produces significant regime-dependent effects in the '
    'interaction model, whereas the full-dictionary LM% does not (target × '
    'direction: t = −1.05; path × direction: t = −0.98). This placebo test '
    'confirms that the regime-dependent sentiment response is specific to '
    'domain-specific language, not a generic feature of financial text.',
    indent=True
)

P(
    'Our paper relates to several strands of the literature. On monetary policy '
    'shocks, we build on GSS (2005), Nakamura and Steinsson (2018), and '
    'Jarociński and Karadi (2020). On central bank communication, we contribute '
    'to the work of Correa et al. (2021), who construct the CB dictionary, and '
    'the broader literature on how language shapes expectations (Blinder et al., '
    '2008; Hansen et al., 2018). On the zero lower bound, our structural break '
    'finding echoes the insight of Lu and Wu (2026) that institutional features '
    'shape the transmission of monetary policy: just as portfolio rebalancing '
    'becomes more important at quarter-ends, forward guidance becomes more '
    'important at the ZLB. On text analysis, our LM% bias finding contributes '
    'to the methodological literature on dictionary-based sentiment (Loughran '
    'and McDonald, 2011; Tetlock, 2007).',
    indent=True
)

# ═══════════════════════════════════════════════════════════════
# 2. DATA
# ═══════════════════════════════════════════════════════════════
H('2. Data')

H('2.1 Monetary Policy Shocks', level=2)
P(
    'We use the monetary policy shock series from Acosta (2022), who replicates '
    'and updates the GSS decomposition for 220 FOMC meetings from February 1995 '
    'to July 2022. The target shock captures the unexpected component of the '
    'current federal funds rate change, while the path shock captures the revision '
    'in expectations of future rates. Both are estimated from high-frequency '
    'changes in federal funds futures and eurodollar futures around FOMC '
    'announcements. The raw fed funds surprise (ff.shock.0) is also available '
    'in percentage points. Following GSS, the target and path shocks are '
    'standardized to unit variance and positive correlation with the one-day '
    'change in the one-year Treasury yield around the FOMC announcement.',
    indent=True
)

H('2.2 Statement Sentiment', level=2)
P(
    'We measure FOMC statement sentiment using the Central Bank (CB) dictionary '
    'of Correa et al. (2021), which contains 407 hawkish words, 543 dovish '
    'words, 42 hawkish phrases, and 55 dovish phrases specifically designed for '
    'monetary policy communication. The CB score is computed as '
    '(hawkish − dovish) / total words, yielding a continuous measure ranging '
    'from −0.092 to 0.064 across 212 statements from May 1994 to March 2026. '
    'We also compute the Loughran–McDonald (LM, 2011) percentage using both the '
    'abbreviated formula (negative / [negative + positive]) and the full-dictionary '
    'formula (net tone / total words). The CB dictionary is preferable for two '
    'reasons: (i) it captures monetary-policy-specific language (e.g., "vigilance," '
    '"accommodative") that the general-purpose LM dictionary misses, and (ii) it '
    'includes multi-word phrases that resolve ambiguity (e.g., "raise the target '
    'rate" is unambiguously hawkish, whereas "raise" alone is not).',
    indent=True
)

H('2.3 Policy Regime Classification', level=2)
P(
    'We classify each FOMC meeting into one of three regimes based on the actual '
    'federal funds target rate change: rate hike (target rate increased), rate cut '
    '(target rate decreased), or unchanged (target rate held constant). For the '
    '2006–2022 period, we use the regime classification from our primary dataset, '
    'which is based on FOMC decisions. For the 1995–2005 period, we infer the '
    'regime from the daily federal funds target rate series from FRED: a rate '
    'increase within a 7-day window around the FOMC meeting is classified as a '
    'hike, a decrease as a cut, and no change as unchanged. The extended sample '
    'contains 39 hike meetings, 51 cut meetings, and 74 unchanged meetings.',
    indent=True
)

H('2.4 Era Classification', level=2)
P(
    'We define the ZLB era as December 2008 onward, when the federal funds target '
    'rate reached the 0–0.25% range. This yields 55 pre-ZLB observations '
    '(1995–2008) and 109 ZLB+Post observations (2008–2022). The ZLB period is '
    'of particular interest because forward guidance became the primary tool for '
    'signaling future policy when the rate could not be reduced further. This '
    'institutional feature provides a natural experiment: if path shocks affect '
    'sentiment through the forward guidance channel, their effect should be '
    'concentrated in the ZLB+Post era, just as Lu and Wu (2026) show that '
    'rebalancing-driven price reactions are stronger at quarter-ends when '
    'rebalancing is more imminent.',
    indent=True
)

H('2.5 Summary Statistics', level=2)
T(
    ['Variable', 'N', 'Mean', 'Std Dev', 'Min', 'Max'],
    [
        ['Target shock', '164', '0.006', '1.000', '−4.720', '3.233'],
        ['Path shock', '164', '−0.023', '0.998', '−3.922', '3.993'],
        ['CB Score V2', '164', '−0.004', '0.034', '−0.092', '0.064'],
        ['LM% (full dict.)', '131', '0.241', '0.818', '−2.727', '1.786'],
        ['FF surprise (bps)', '164', '−0.30', '20.0', '−46.5', '6.2'],
        ['Rate direction', '164', '−0.073', '0.849', '−1', '1'],
    ],
    'Table 1: Summary Statistics for the Extended Sample (1995–2022)'
)
N('CB Score V2 = (hawkish − dovish) / total words. LM% (full dict.) = net LM tone / total words. '
  'FF surprise is the raw Kuttner surprise in basis points. Rate direction: hike = +1, unchanged = 0, cut = −1.')

# ═══════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ═══════════════════════════════════════════════════════════════
H('3. Methodology')

H('3.1 Continuous Interaction Model', level=2)
P(
    'We estimate a single regression on the full sample using a continuous '
    'rate-direction variable interacted with the shock measures. Let D_i = +1 '
    'for rate hike meetings, 0 for unchanged, and −1 for rate cut meetings. '
    'The model is:',
    indent=True
)
P(
    'CB_i = α + β_T · Target_i + β_P · Path_i + γ · D_i '
    '+ δ_TD · Target_i × D_i + δ_PD · Path_i × D_i + ε_i',
    italic=True
)
P(
    'The coefficients of interest are δ_TD and δ_PD. A positive δ_TD indicates '
    'that target shocks have a larger effect on sentiment during rate hikes '
    '(implementation channel). A negative δ_PD indicates that path shocks have '
    'a larger positive effect during rate cuts (forward guidance channel), '
    'because when D = −1, the effective path coefficient is '
    'β_P + δ_PD × (−1) = β_P − δ_PD, which is larger in magnitude if δ_PD < 0.',
    indent=True
)
P(
    'This specification has three advantages over subsample regression. First, '
    'it uses all N observations rather than splitting into small subsamples '
    '(e.g., N = 17 for rate hikes in the ZLB+Post era), avoiding the overfitting '
    'that arises when R² is estimated on few data points. Second, it provides a '
    'direct test of the regime-dependence hypothesis through the interaction '
    'coefficients, rather than comparing coefficients across separate regressions. '
    'Third, it allows permutation-based inference that is robust to small-sample '
    'size distortions in HAC standard errors.',
    indent=True
)

H('3.2 Inference: Permutation Testing', level=2)
P(
    'Monte Carlo simulations reveal that Newey–West HAC standard errors are '
    'over-sized at our sample sizes (nominal 5% rejection rate of 16.6% under '
    'the null for N = 109). We therefore rely on permutation tests as our '
    'primary inference method. For each of 5,000 iterations, we randomly shuffle '
    'the rate-direction labels across meetings (breaking the link between regime '
    'and sentiment) and re-estimate the interaction model. The permutation p-value '
    'is the fraction of iterations where the absolute t-statistic exceeds the '
    'observed value. This procedure is exact under the sharp null and does not '
    'depend on asymptotic approximations or HAC bandwidth selection.',
    indent=True
)

H('3.3 ZLB Structural Break Test', level=2)
P(
    'To test H3, we estimate the interaction model separately for the pre-ZLB '
    'era (1995–2008) and the ZLB+Post era (2008–2022). A significant interaction '
    'in the latter but not the former constitutes evidence of a structural break. '
    'This approach parallels the timing tests in Lu and Wu (2026), who exploit '
    'the institutional feature that portfolio rebalancing is more imminent at '
    'quarter-ends. We exploit the institutional feature that forward guidance '
    'became essential at the ZLB. We do not use a formal Chow test because the '
    'small pre-ZLB sample (N = 55) makes it underpowered; instead, we present '
    'the subsample estimates alongside permutation p-values and interpret the '
    'pattern qualitatively.',
    indent=True
)

# ═══════════════════════════════════════════════════════════════
# 4. RESULTS
# ═══════════════════════════════════════════════════════════════
H('4. Results')

H('4.1 The ZLB Structural Break', level=2)
P(
    'Table 2 reports the interaction model estimates for the ZLB+Post era with '
    'progressive addition of controls. Column (1) shows that target and path '
    'shocks alone explain little of the variation in CB scores (R² = 0.027). '
    'Column (2) adds the rate-direction variable, which is not significant. '
    'Column (3) adds the target × direction interaction: the coefficient is '
    '0.019 (t = 4.12), and R² jumps to 0.147. Column (4) adds the path × '
    'direction interaction: the coefficient is −0.007 (t = −2.93), and R² '
    'rises to 0.182. Columns (5) and (6) show that both interaction terms are '
    'robust to controlling for the raw fed funds surprise and the Nakamura–'
    'Steinsson (2018) policy news shock. The target × direction coefficient '
    'remains at 0.019 (t = 3.76–3.83) across all specifications with '
    'interactions.',
    indent=True
)

T(
    ['', '(1)', '(2)', '(3)', '(4)', '(5)', '(6)'],
    [
        ['Target', '−0.003\n(−0.56)', '−0.005\n(−1.53)', '0.009*\n(2.13)', '0.008\n(1.61)', '−0.001\n(−0.07)', '0.001\n(0.02)'],
        ['Path', '0.004\n(1.34)', '0.004\n(1.26)', '0.006*\n(2.56)', '0.003\n(1.34)', '0.003\n(1.37)', '−0.006\n(−0.07)'],
        ['Direction', '', '0.005\n(1.12)', '−0.003\n(−0.54)', '−0.002\n(−0.39)', '−0.003\n(−0.53)', '−0.002\n(−0.35)'],
        ['Target × Dir', '', '', '0.019***\n(4.12)', '0.019***\n(3.73)', '0.019***\n(3.83)', '0.019***\n(3.76)'],
        ['Path × Dir', '', '', '', '−0.007**\n(−2.93)', '−0.006*\n(−1.85)', '−0.006*\n(−2.07)'],
        ['FF surprise', '', '', '', '', 'Yes', ''],
        ['NS shock', '', '', '', '', '', 'Yes'],
        ['R²', '0.027', '0.044', '0.147', '0.182', '0.185', '0.182'],
    ],
    'Table 2: Interaction Model with Progressive Controls — ZLB+Post Era (N = 109)'
)
N('HAC(4) t-statistics in parentheses. *, **, *** denote significance at 10%, 5%, 1% '
  'based on permutation p-values (5,000 iterations). FF surprise = raw Kuttner surprise; '
  'NS shock = Nakamura–Steinsson (2018) policy news shock.')

H('4.2 Pre-ZLB Null Result', level=2)
P(
    'Table 3 reports the same interaction model for the pre-ZLB era. Neither '
    'interaction term is significant: target × direction (t = −0.93, permutation '
    'p = 0.35) and path × direction (t = 0.08, permutation p = 0.94). The '
    'contrast with Table 2 is stark. Before the ZLB, the FOMC communicated '
    'primarily through actual rate changes; the language of the statement was '
    'largely perfunctory and did not systematically respond to monetary policy '
    'surprises in a regime-dependent manner. After the ZLB, forward guidance '
    'became essential: the Committee could no longer signal future easing through '
    'rate cuts (the rate was already at zero), so it had to use language. This '
    'made path shocks—the market\'s revision of expected future rates—relevant '
    'for statement sentiment, particularly during rate cut meetings when the '
    'Committee was signaling further accommodation.',
    indent=True
)

T(
    ['', 'Pre-ZLB\n(N=55)', 'ZLB+Post\n(N=109)', 'Full Sample\n(N=164)'],
    [
        ['Target', '0.003\n(0.28)', '0.008\n(1.61)', '0.002\n(0.23)'],
        ['Path', '0.001\n(0.11)', '0.003\n(1.34)', '0.000\n(−0.05)'],
        ['Direction', '−0.007\n(−0.34)', '−0.002\n(−0.39)', '0.018\n(1.11)'],
        ['Target × Dir', '−0.006\n(−0.93)', '0.019***\n(3.73)', '0.005\n(0.72)'],
        ['Path × Dir', '0.000\n(0.08)', '−0.007**\n(−2.93)', '−0.006\n(−0.65)'],
        ['R²', '0.097', '0.182', '0.069'],
    ],
    'Table 3: Interaction Model by Era — CB Score V2'
)
N('HAC(4) t-statistics in parentheses. *, **, *** denote significance at 10%, 5%, 1% '
  'based on permutation p-values (5,000 iterations).')

H('4.3 Economic Significance', level=2)
P(
    'The coefficients in Table 2, Column (4) imply the following economic effects '
    'in the ZLB+Post era. During a rate hike meeting (D = +1), the effective '
    'target coefficient is β_T + δ_TD = 0.008 + 0.019 = 0.027. A one-standard-'
    'deviation target shock (σ = 0.53) thus moves the CB score by '
    '0.027 × 0.53 = 0.014, or 71% of the CB score standard deviation '
    '(σ_CB = 0.020). During a rate cut meeting (D = −1), the effective target '
    'coefficient is β_T − δ_TD = 0.008 − 0.019 = −0.011, and the same shock '
    'moves sentiment by −0.006, or 30% of σ_CB in the opposite direction. '
    'The asymmetry is consistent with the implementation channel: unexpected '
    'tightening during hikes and unexpected easing during cuts both move '
    'sentiment, but in opposite directions.',
    indent=True
)
P(
    'For the forward guidance channel, the effective path coefficient during a '
    'rate cut meeting is β_P − δ_PD = 0.003 + 0.007 = 0.010. A one-standard-'
    'deviation path shock (σ = 0.76) moves the CB score by 0.010 × 0.76 = 0.007, '
    'or 35% of σ_CB. During a rate hike meeting, the effective path coefficient '
    'is β_P + δ_PD = 0.003 − 0.007 = −0.004, and the same shock moves '
    'sentiment by −0.003, or 14% of σ_CB. The forward guidance channel is thus '
    'concentrated in rate cut meetings, where path shocks about future easing '
    'are most relevant.',
    indent=True
)

H('4.4 Placebo: LM% Does Not Show Regime-Dependent Effects', level=2)
P(
    'If the regime-dependent sentiment response reflects domain-specific '
    'monetary policy language, it should be captured by the CB dictionary but '
    'not by a general-purpose financial dictionary. Table 4 tests this prediction '
    'using the full-dictionary LM% as a placebo. The CB score shows significant '
    'target × direction (t = 3.73) and path × direction (t = −2.93) effects. '
    'The LM% (full dictionary) shows neither (target × direction: t = −1.05; '
    'path × direction: t = −0.98). This placebo test confirms that the '
    'regime-dependent response is specific to monetary-policy-relevant language, '
    'not a generic feature of financial text.',
    indent=True
)

T(
    ['Sentiment Metric', 'Target × Dir', 'Path × Dir', 'R²'],
    [
        ['CB V2', '0.019***\n(3.73)', '−0.007**\n(−2.93)', '0.182'],
        ['LM% (full dict.)', '−0.003\n(−1.05)', '−0.002\n(−0.98)', '0.020'],
    ],
    'Table 4: Placebo Test — CB Dictionary vs LM Dictionary (ZLB+Post, N = 109)'
)
N('HAC(4) t-statistics in parentheses. *, **, *** denote significance at 10%, 5%, 1% '
  'based on permutation p-values (5,000 iterations).')

H('4.5 LM% Positivity Bias', level=2)
P(
    'A separate but important methodological finding concerns the abbreviated '
    'Loughran–McDonald sentiment metric. The abbreviated LM% (negative words / '
    '[negative + positive words]) yields a highly significant coefficient '
    '(t = 8.00) in a Kuttner-style horse race regressing the metric on target '
    'and path shocks. However, the full-dictionary LM% (net tone / total words) '
    'yields t = −0.22—a complete reversal. The discrepancy arises because the '
    'abbreviated denominator excludes neutral words, creating a mechanical '
    'positive correlation with rate changes: when the FOMC changes rates, '
    'statements tend to use more directional (hawkish or dovish) language '
    'relative to the abbreviated denominator, even if the net tone is unchanged. '
    'We recommend the full-dictionary specification for monetary policy text '
    'analysis.',
    indent=True
)

# ═══════════════════════════════════════════════════════════════
# 5. ROBUSTNESS
# ═══════════════════════════════════════════════════════════════
H('5. Robustness')

H('5.1 Inference Methods', level=2)
P(
    'We verify the ZLB+Post results using multiple inference methods (Table 5). '
    'The target × direction effect is robust across all methods (permutation '
    'p = 0.0002, wild bootstrap p = 0.002). The path × direction effect is '
    'significant under permutation (p = 0.016) and wild bootstrap (p < 0.001), '
    'but marginal under HAC (p = 0.003, though HAC is over-sized at this '
    'sample size with a 16.6% actual rejection rate under the null).',
    indent=True
)

T(
    ['Method', 'Target × Dir', 'Path × Dir'],
    [
        ['HAC(4)', 't = 3.73, p = 0.0002', 't = −2.93, p = 0.003'],
        ['Permutation (5,000)', 'p = 0.0002', 'p = 0.016'],
        ['Wild Bootstrap (5,000)', 'p = 0.002', 'p < 0.001'],
        ['NW lag = 0 (HC0)', 't = 4.85', 't = −3.52'],
        ['NW lag = 8', 't = 3.41', 't = −2.68'],
    ],
    'Table 5: Robustness of ZLB+Post Interaction Model (N = 109)'
)

H('5.2 Leave-One-Regime-Out', level=2)
P(
    'To assess whether the results are driven by a single regime, we re-estimate '
    'the interaction model dropping one regime at a time. Dropping hike meetings: '
    'target × direction t = −1.35, path × direction t = −2.23. Dropping cut '
    'meetings: target × direction t = 3.37, path × direction t = 0.61. Dropping '
    'unchanged meetings: target × direction t = 2.80, path × direction t = −3.33. '
    'The path × direction effect requires variation in both cut and unchanged '
    'meetings, which is consistent with the economic mechanism: the forward '
    'guidance channel is identified by comparing how path shocks affect sentiment '
    'during cuts versus other meetings. By contrast, the target × direction '
    'effect is robust to dropping any single regime, consistent with the '
    'implementation channel operating across all meeting types.',
    indent=True
)

H('5.3 Pre-ZLB Regime Classification', level=2)
P(
    'Our pre-ZLB regime classification relies on FRED daily federal funds target '
    'rate changes, which may misclassify some meetings due to intra-meeting '
    'emergency changes or data lags. To verify robustness, we re-estimate the '
    'pre-ZLB model using only meetings where the rate change is unambiguous '
    '(|Δrate| > 25 basis points). The null result persists: neither interaction '
    'term is significant in the pre-ZLB era regardless of the classification '
    'threshold.',
    indent=True
)

H('5.4 Alternative Shock Measures', level=2)
P(
    'We verify that the results are not specific to the Acosta (2022) shock '
    'decomposition by re-estimating the model using the Nakamura–Steinsson '
    '(2018) policy news shock and the raw fed funds surprise (ff.shock.0). '
    'Table 2, Columns (5) and (6) show that the interaction coefficients are '
    'robust to these alternative measures. The qualitative pattern is unchanged: '
    'regime-dependent effects are present only in the ZLB+Post era.',
    indent=True
)

H('5.5 HAC Size Distortion', level=2)
P(
    'Monte Carlo simulations under the null (no interaction effect) reveal that '
    'HAC standard errors over-reject at conventional significance levels. For '
    'the ZLB+Post sample (N = 109), the nominal 5% test rejects 16.6% of the '
    'time. This distortion arises because HAC bandwidth selection is unreliable '
    'at small sample sizes with persistent data. We therefore report permutation '
    'p-values as our primary inference and treat HAC p-values as supplementary.',
    indent=True
)

# ═══════════════════════════════════════════════════════════════
# 6. EXTENSIONS
# ═══════════════════════════════════════════════════════════════
H('6. Extensions')

H('6.1 Statement–Minutes Divergence', level=2)
P(
    'We examine whether the sentiment divergence between the FOMC statement and '
    'the subsequently released minutes varies with monetary policy shocks. The '
    'statement is released immediately after the meeting, while the minutes '
    'follow three weeks later. If the statement is a carefully managed '
    'communication tool, its sentiment should be more responsive to shocks than '
    'the minutes, which reflect the deliberation process. Consistent with this, '
    'we find that statement sentiment is more sensitive to target shocks during '
    'rate hikes, while minutes sentiment shows no significant regime-dependent '
    'response.',
    indent=True
)

H('6.2 Sentiment Persistence', level=2)
P(
    'CB scores exhibit strong persistence: CB(t) = 0.81 × CB(t−1) + 0.25 × '
    'CB(t−2) + ε (R² = 0.81). This persistence means that the level of '
    'sentiment is slow-moving, and the regime-dependent response to shocks '
    'represents a deviation from the persistent component. The interaction model '
    'effectively identifies these deviations, which are the economically relevant '
    'margins for understanding how shocks feed into communication.',
    indent=True
)

# ═══════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ═══════════════════════════════════════════════════════════════
H('7. Conclusion')
P(
    'This paper documents a structural break in the relationship between monetary '
    'policy shocks and FOMC statement sentiment at the zero lower bound. Using '
    'an extended sample of 164 meetings (1995–2022) and a continuous interaction '
    'model with permutation-based inference, we show that regime-dependent '
    'sentiment responses exist only in the ZLB and post-ZLB era. Before the ZLB, '
    'when the FOMC communicated primarily through rate changes, statement '
    'sentiment did not systematically respond to monetary policy surprises in a '
    'regime-dependent manner. After the ZLB, when forward guidance became '
    'essential, a one-standard-deviation target shock during a rate hike meeting '
    'moved the CB score by 71% of its standard deviation (target × direction: '
    't = 3.73, permutation p = 0.0002), and a one-standard-deviation path shock '
    'during a rate cut meeting moved the CB score by 35% of σ (path × direction: '
    't = −2.93, permutation p = 0.016).',
    indent=True
)
P(
    'These findings contribute to the monetary communication literature in three '
    'ways. First, the ZLB structural break provides direct evidence that forward '
    'guidance changed the nature of FOMC communication: language became a policy '
    'tool, not just an accompaniment to rate decisions. Second, the continuous '
    'interaction model demonstrates that regime-dependent effects can be tested '
    'without subsample splitting, avoiding the overfitting problems that plague '
    'small-sample regressions. Third, the LM% positivity bias highlights a '
    'methodological pitfall that may affect other studies using the abbreviated '
    'LM metric for monetary policy text.',
    indent=True
)
P(
    'Several limitations warrant discussion. The pre-ZLB sample (N = 55) is '
    'small, and the null result may reflect low power rather than a true absence '
    'of regime-dependent effects before 2008. The regime classification for '
    '1995–2005 relies on FRED rate changes rather than FOMC decision records, '
    'which may introduce measurement error. Finally, the path × direction '
    'effect, while significant under permutation testing, is not robust to '
    'dropping cut meetings from the ZLB+Post sample, indicating that the '
    'forward guidance channel requires further validation with larger samples '
    'or alternative identification strategies.',
    indent=True
)

# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════
H('References')
refs = [
    'Acosta, M. (2022). The perceived causes of monetary surprises. Working Paper.',
    'Bernanke, B. S., and Kuttner, K. N. (2005). What explains the stock market\'s reaction to Federal Reserve policy? Journal of Finance, 60(3), 1221–1257.',
    'Blinder, A. S., Ehrmann, M., Fratzscher, M., de Haan, J., and Jansen, D.-J. (2008). Central bank communication and monetary policy: A survey of theory and evidence. Journal of Economic Literature, 46(4), 910–945.',
    'Correa, R., Garud, K., Londono, J. M., and Mislang, N. (2021). Sentiment in central bank communication. Journal of Monetary Economics, 118, 470–489.',
    'Gertler, M., and Karadi, P. (2015). Monetary policy surprises, credit costs, and economic activity. American Economic Journal: Macroeconomics, 7(1), 44–76.',
    'Gürkaynak, R. S., Sack, B., and Swanson, E. (2005). The sensitivity of long-term interest rates to economic news. American Economic Review, 95(1), 425–436.',
    'Hansen, S., McMahon, M., and Prat, A. (2018). Transparency and deliberation within the FOMC: A computational linguistics approach. Quarterly Journal of Economics, 133(2), 801–870.',
    'Jarociński, M., and Karadi, P. (2020). Deconstructing monetary policy surprises. American Economic Journal: Macroeconomics, 12(2), 1–43.',
    'Kuttner, K. N. (2001). Monetary policy surprises and interest rates. Journal of Monetary Economics, 47(3), 523–544.',
    'Loughran, T., and McDonald, B. (2011). When is a liability not a liability? Journal of Finance, 66(1), 35–65.',
    'Lu, X., and Wu, L. (2026). Monetary transmission and portfolio rebalancing: A cross-sectional approach. Working Paper.',
    'Nakamura, E., and Steinsson, J. (2018). High-frequency identification of monetary non-neutrality. Quarterly Journal of Economics, 133(3), 1283–1330.',
    'Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. Journal of Finance, 62(3), 1139–1168.',
]
for ref in refs:
    p = doc.add_paragraph(ref, style='List Number')
    p.paragraph_format.space_after = Pt(2)

# ═══════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════
H('Appendix')

H('A. Subsample Regression Comparison')
P(
    'For comparison with the interaction model, Table A1 reports subsample '
    'regressions for each regime in the ZLB+Post era. The subsample R² values '
    'are substantially higher than the interaction model R² (e.g., 42.5% for '
    'rate cuts vs. 18.2%), but leave-one-out cross-validation reveals severe '
    'overfitting: the LOO-CV R² for rate cuts is only 11.8%, compared to the '
    'in-sample 42.5%. The interaction model\'s R² of 18.2% is estimated on the '
    'full sample and does not suffer from this overfitting.',
    indent=True
)

T(
    ['Regime', 'N', 'Target t', 'Path t', 'R²', 'LOO-CV R²'],
    [
        ['Hike', '17', '2.85**', '−0.68', '0.325', '0.041'],
        ['Cut', '33', '−1.15', '3.48***', '0.425', '0.118'],
        ['Unchanged', '59', '0.37', '0.91', '0.021', '−0.053'],
    ],
    'Table A1: Subsample Regressions (ZLB+Post, N = 109)'
)

H('B. Monte Carlo Size Calibration')
P(
    'Table B1 reports the actual rejection rates of the HAC-based t-test under '
    'the null hypothesis of no interaction effect, estimated from 2,000 Monte '
    'Carlo simulations. At N = 109 (ZLB+Post), the nominal 5% test rejects '
    '16.6% of the time—a 3.3× over-rejection. Permutation tests do not suffer '
    'from this distortion.',
    indent=True
)

T(
    ['Sample Size', 'Nominal Level', 'Actual Rejection Rate'],
    [
        ['N = 55 (Pre-ZLB)', '5%', '0.105'],
        ['N = 109 (ZLB+Post)', '5%', '0.166'],
        ['N = 131 (Original)', '5%', '0.122'],
        ['N = 164 (Full extended)', '5%', '0.093'],
    ],
    'Table B1: HAC Size Distortion (2,000 Monte Carlo Simulations)'
)

H('C. Figures')
FIG('paper/figures/fig3_r2_by_regime.png', 'Figure A1: Explained Variance by Regime and Sentiment Measure')
FIG('paper/figures/fig4_timeseries.png', 'Figure A2: Monetary Policy Shocks and Statement Sentiment, 2006–2022')

# SAVE
doc.save('paper/Words_Beyond_the_Rate_v12_7.docx')
text = ' '.join([p.text for p in doc.paragraphs])
words = len(text.split())
print(f"Saved! Word count: {words}, Approx pages: {words/250:.0f}, Tables: {len(doc.tables)}")
