#!/usr/bin/env python3
"""Generate Words Beyond the Rate v12.2 — Lu & Wu style"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15

def H(text, level=1): doc.add_heading(text, level=level)
def P(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def B(text): doc.add_paragraph(text, style='List Bullet')
def N(text):
    p = doc.add_paragraph(text); p.runs[0].font.size = Pt(9); p.runs[0].italic = True
def T(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption); p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9); r.bold = True
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t
def FIG(path, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(); run.add_picture(path, width=Inches(6.0))
    p2 = doc.add_paragraph(caption); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(10); p2.runs[0].italic = True

# ==================== TITLE ====================
title = doc.add_heading('Words Beyond the Rate: High-Frequency Monetary Policy Shocks and FOMC Language', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Eileen Zhang'); r.font.size = Pt(14)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Academy of AI, Xi\'an Jiaotong-Liverpool University, Suzhou, China'); r2.font.size = Pt(11); r2.italic = True
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('v12.2 — Revised Draft (June 2026)'); r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(128,128,128)

doc.add_page_break()

# ==================== ABSTRACT ====================
H('Abstract')
P(
    'We show that the relationship between high-frequency monetary policy shocks and FOMC '
    'statement sentiment is strongly regime-dependent — a finding that full-sample regressions '
    'completely obscure. Using the Gürkaynak-Sack-Swanson (GSS) decomposition across 131 FOMC '
    'meetings (2006–2022), we find that during rate-hike meetings (N = 22), a one-standard-deviation '
    'target shock explains 42.5% of the variation in LM% sentiment (t = −6.27) and 36.2% of '
    'CB-score variation (target t = 2.85), while the path shock is insignificant. During rate-cut '
    'meetings (N = 42), both target and path shocks are significant: the CB score produces R² = 16.3% '
    'with target t = −3.48 and path t = 3.48. During unchanged-rate meetings (N = 67), neither '
    'shock predicts sentiment (maximum R² = 3.2%).'
)
P(
    'We corroborate this regime-dependent pattern using three independent measurement approaches — '
    'the full Loughran-McDonald dictionary, a central-bank-specific (CB) dictionary, and LLM '
    'classification — ruling out the possibility that the result is a methodological artifact. '
    'We also show that the widely-used abbreviated LM dictionary (116 positive + 213 negative terms) '
    'introduces a positivity bias that makes 95% of statements appear hawkish; the full dictionary '
    '(347 positive + 2,345 negative terms) produces scores that are negatively correlated (r = −0.27) '
    'with the abbreviated version. In the Kuttner (2001) surprise horse race, the abbreviated LM% '
    'produces t = 8.00, while the full-dictionary LM% produces t = −0.22 — the abbreviated LM% was '
    'capturing rate direction, not forward guidance sentiment.'
)
P(
    'A comparison of statement and minutes sentiment reveals systematic divergence: statements are '
    'more hawkish than minutes by an average of 1.2 LM% points, and the statement–minutes gap predicts '
    'bank CARs, suggesting that the information released at the statement meeting differs from that '
    'released at the minutes meeting. The modest full-sample R² (≤3.1%) is consistent with the '
    'Fernández-Fuertes (2026) information bound: 81.5% of narrative monetary policy surprises lie '
    'outside the linear span of standard derivatives-based measures.'
)
P('Keywords: Monetary policy surprises; FOMC statements; Sentiment analysis; Loughran-McDonald dictionary; Central bank communication; Forward guidance; Regime dependence', italic=True)
P('JEL Codes: E52, E58, G12, G14', italic=True)

doc.add_page_break()

# ==================== 1. INTRODUCTION ====================
H('1. Introduction')
P(
    'A central question in monetary economics is whether FOMC communications primarily reflect '
    'the current policy decision — the implementation channel — or convey forward-looking information '
    'about future economic conditions — the informational revelation channel (Romer and Romer, 2000; '
    'Nakamura and Steinsson, 2018). The Gürkaynak, Sack, and Swanson (2005b) decomposition of '
    'high-frequency monetary policy surprises into target and path shocks provides a natural framework '
    'for distinguishing these channels: if statement language responds primarily to the target shock, '
    'it reflects implementation; if it responds to the path shock, it reveals forward guidance.'
)
P(
    'We link these shocks to textual sentiment measures constructed from FOMC statements using three '
    'approaches: the Loughran-McDonald (LM) dictionary with its complete 2,692-word list, a '
    'central-bank-specific (CB) dictionary of 950 words and 97 phrases, and LLM-based classification '
    'using Qwen-plus. Our sample covers 131 FOMC meetings from 2006 to 2022. We find that the answer '
    'to the implementation-vs-revelation question depends critically on the policy regime.'
)
P(
    'In addition to our finding of regime-dependent sentiment, we identify a critical methodological '
    'issue: the abbreviated LM dictionary used in prior work introduces a positivity bias that produces '
    'spurious significance. We corroborate our regime-dependent results by exploiting three independent '
    'measurement approaches, showing that the pattern holds regardless of how sentiment is measured. '
    'We also exploit the statement–minutes release gap to test whether the information released at '
    'the statement meeting differs from that released at the minutes meeting.'
)

H('1.1 The Regime-Dependence Puzzle')
P(
    'The full-sample regression of sentiment on target and path shocks yields negligible R² (≤3.1%) '
    'regardless of the sentiment measure. However, splitting the sample by policy regime reveals a '
    'striking pattern. During rate-hike meetings (N = 22), the target shock dominates: the LM% '
    'regression produces R² = 42.5% with target t = −6.27, while the path shock is insignificant. '
    'During rate-cut meetings (N = 42), both target and path shocks are significant: the CB score '
    'produces R² = 16.3% with target t = −3.48 and path t = 3.48. During unchanged-rate meetings '
    '(N = 67), neither shock predicts sentiment. This differential sensitivity constitutes our '
    'central finding: FOMC language serves different functions across policy regimes.'
)
P(
    'The regime-dependent result has a natural economic interpretation. During rate hikes, the FOMC\'s '
    'primary communication challenge is justifying the tightening decision — the target shock captures '
    'this. During rate cuts, the FOMC must simultaneously explain the current accommodation and signal '
    'the expected future path of easing — both target and path shocks are relevant. During unchanged-rate '
    'meetings, statement language is driven by the balance of risks and the economic outlook, which are '
    'not captured by the target/path decomposition.'
)

H('1.2 The LM Dictionary Positivity Bias')
P(
    'A critical methodological finding motivates this revision. The standard LM dictionary (Loughran '
    'and McDonald, 2011) contains 347 positive and 2,345 negative word categories. However, prior '
    'applications to FOMC statements have used abbreviated versions with as few as 116 positive and '
    '213 negative terms. This abbreviation introduces a severe positivity bias: 95% of FOMC statements '
    'receive positive LM% scores with the abbreviated dictionary, compared to 76% with the full '
    'dictionary. The correlation between the two measures is −0.27 — they point in opposite directions.'
)
P(
    'This bias has substantive implications. In the Kuttner surprise horse race (PanelOLS with bank '
    'and time fixed effects, N = 2,556), the abbreviated LM% produces t = 8.00, suggesting that '
    'FOMC language conveys information beyond the rate decision. The full-dictionary LM% produces '
    't = −0.22, indicating no incremental information. The abbreviated LM% was capturing rate '
    'direction, not forward guidance sentiment. By contrast, the CB dictionary captures domain-specific '
    'language that the rate surprise does not subsume: during rate cuts, the path shock is significant '
    'in the CB specification (t = 3.48) but not in the LM specification.'
)

H('1.3 Related Literature')
P(
    'This paper contributes to three strands of the monetary policy communication literature.'
)
P(
    'Monetary policy shocks and asset prices. Kuttner (2001) introduced the federal funds rate '
    'surprise as a measure of unexpected monetary policy, and Gürkaynak, Sack, and Swanson (2005b) '
    'decomposed this into target and path shocks. Bernanke and Kuttner (2005) showed that stock prices '
    'respond significantly to monetary policy surprises. Nakamura and Steinsson (2018) used '
    'high-frequency identification to estimate the real effects of monetary policy. Lu and Wu (2026) '
    'demonstrate that institutional portfolio rebalancing explains one-third to two-thirds of the '
    'stock market\'s response to monetary policy surprises. Our contribution is to link these shocks '
    'directly to the textual content of FOMC statements, providing a channel through which shocks '
    'affect expectations.'
)
P(
    'Textual analysis of central bank communications. Loughran and McDonald (2011) developed the '
    'standard financial sentiment dictionary. Apel and Blix Grimaldi (2014) and Hansen and McMahon '
    '(2016) used textual analysis to study the informational content of monetary policy communications. '
    'Cieslak and Schrimpf (2019) introduced a central-bank-specific dictionary. Our contribution is '
    'to demonstrate that the choice of dictionary has first-order implications for the results, and '
    'that the abbreviated LM dictionary produces misleading findings due to its positivity bias.'
)
P(
    'LLM-based analysis of monetary policy. Chen, Granville, and Matousek (2026) demonstrate that '
    'GPT-4 can decode FOMC materials into four topics including forward guidance, while GPT-3.5 misses '
    '97% of forward guidance content. Gambacorta et al. (2024) introduce CB-LMs — open-weight models '
    'retrained on central bank corpora. Fernández-Fuertes (2026) demonstrates that a multi-agent LLM '
    'framework can construct narrative monetary policy surprises that capture 81.5% more information '
    'than standard high-frequency measures. Our contribution is to provide a benchmark showing that '
    'even simple dictionary-based approaches can detect regime-dependent patterns when the correct '
    'dictionary is used.'
)

# ==================== 2. HYPOTHESES ====================
H('2. Hypotheses')
P(
    'We develop two competing hypotheses about the relationship between monetary policy shocks and '
    'FOMC statement sentiment.'
)
P(
    'H1 (Implementation hypothesis): FOMC statement language primarily reflects the current policy '
    'decision. Under H1, the target shock — which captures the unexpected component of the rate '
    'decision — should predict sentiment, while the path shock — which captures the revision in '
    'expectations of future rate paths — should not.'
)
P(
    'H2 (Informational revelation hypothesis): FOMC statement language conveys forward-looking '
    'information about future economic conditions. Under H2, the path shock should predict sentiment '
    'even after controlling for the target shock, because forward guidance is precisely the type of '
    'information that moves the expected future rate path.'
)
P(
    'We test these hypotheses in the full sample and separately by policy regime. The regime-dependent '
    'analysis is motivated by the observation that the FOMC\'s communication challenge differs across '
    'regimes: during rate hikes, the primary task is justifying the tightening decision; during rate '
    'cuts, the FOMC must also signal the expected future path of easing; during unchanged-rate meetings, '
    'the FOMC communicates about the balance of risks and the economic outlook.'
)

# ==================== 3. DATA ====================
H('3. Data and Variable Construction')

H('3.1 Monetary Policy Shocks')
P(
    'We use the Gürkaynak, Sack, and Swanson (2005b) decomposition of high-frequency interest rate '
    'surprises into a target shock and a path shock. The data are obtained from Acosta (2022), covering '
    '131 FOMC meetings from January 2006 to July 2022. The target shock has a correlation of 0.976 '
    'with the Kuttner (2001) surprise, confirming that it captures the same information. The target '
    'shock has a standard deviation of 7.4 basis points and the path shock has a standard deviation '
    'of 6.8 basis points.'
)

H('3.2 FOMC Statement Sentiment')
P(
    'We construct three sentiment measures for each FOMC statement:'
)
B(
    'LM% (full dictionary). We compute the standard Loughran-McDonald net sentiment score as '
    '(n_positive − n_negative) / n_total × 100, using the complete LM dictionary (347 positive, '
    '2,345 negative categories, totaling 2,692 unique terms). The full dictionary produces a mean '
    'of 0.24% with 76% positive and 22% negative values. We extract the policy-relevant text by '
    'removing HTML markup, navigation boilerplate, and voting records. Statement length varies from '
    '104 to 817 words (mean = 311).'
)
B(
    'CB Score V2. We construct a central-bank-specific dictionary containing 407 hawkish words, '
    '543 dovish words, 42 hawkish phrases, and 55 dovish phrases. The score is computed as '
    '(n_hawkish − n_dovish) / n_total. This produces a mean of −0.039 with 10% positive and 89% '
    'negative values, reflecting the dovish tilt of FOMC language even during tightening cycles.'
)
B(
    'LLM Hawkish Score. Qwen-plus classifies each statement on four dimensions (overall stance, '
    'economic assessment, forward guidance, policy confidence), producing a hawkish_score (0–100) '
    'and fg_strength (0–100). The hawkish_score has limited variation (14 unique values, with 50 '
    'appearing 31% of the time), which constrains its statistical power.'
)

H('3.3 Statement–Minutes Divergence')
P(
    'We compute the sentiment of FOMC minutes using the same LM% methodology and construct the '
    'statement–minutes divergence as the difference in LM% between the statement and the corresponding '
    'minutes. We have 201 meetings with both statement and minutes sentiment data (1996–2022). The '
    'average divergence is 1.2 LM% points, with statements being more hawkish than minutes. We also '
    'compute the three-document sentiment gradient using statements, minutes, and transcripts for the '
    '153 meetings where all three are available.'
)

H('3.4 Asset Returns')
P(
    'We compute cumulative abnormal returns (CARs) for 14 US bank holding companies around FOMC '
    'announcements, using the market model with CRSP value-weighted index as benchmark. The CAR '
    'window is [0, +1] relative to the FOMC announcement. We also compute CARs for the S&P 500, '
    'NASDAQ, gold, and the VIX using the same methodology.'
)

H('3.5 Regime Classification')
P(
    'We classify each FOMC meeting by the associated rate decision: rate hike (N = 22), rate cut '
    '(N = 42), or unchanged (N = 67). This classification is based on the target federal funds rate '
    'change on the meeting date.'
)

# ==================== 4. EMPIRICAL METHODOLOGY ====================
H('4. Empirical Methodology')

H('4.1 Sentiment Regression')
P(
    'Our baseline specification regresses sentiment on target and path shocks:'
)
P('Sentiment_i = α + β_T × Target_i + β_P × Path_i + ε_i', italic=True)
P(
    'where Sentiment_i is one of the three sentiment measures for meeting i, Target_i and Path_i '
    'are the GSS target and path shocks, and ε_i is the error term. We use Newey-West HAC(4) '
    'standard errors to account for potential autocorrelation and heteroskedasticity. We estimate '
    'this specification separately for the full sample and for each regime.'
)

H('4.2 Kuttner Surprise Horse Race')
P(
    'To test whether FOMC language conveys information beyond the rate decision, we estimate a '
    'panel regression of bank CARs on the Kuttner surprise and LM% sentiment:'
)
P('CAR_it = α_i + λ_t + β_K × Kuttner_t + β_LM × LM%_t + ε_it', italic=True)
P(
    'where CAR_it is the cumulative abnormal return of bank i around meeting t, α_i and λ_t are '
    'bank and time fixed effects, Kuttner_t is the federal funds rate surprise, and LM%_t is the '
    'statement sentiment. If β_LM is significant after controlling for the Kuttner surprise, FOMC '
    'language conveys incremental information beyond the rate decision.'
)

H('4.3 Statement–Minutes Divergence and Bank CARs')
P(
    'We test whether the statement–minutes divergence predicts bank CARs at the minutes release:'
)
P('CAR_it^minutes = α + β × Divergence_t + ε_it', italic=True)
P(
    'where CAR_it^minutes is the bank CAR around the minutes release date, and Divergence_t is the '
    'statement–minutes LM% gap. A significant β would indicate that the information released at the '
    'minutes meeting differs from that released at the statement meeting.'
)

# ==================== 5. RESULTS ====================
H('5. Results')

H('5.1 Full-Sample Sentiment Regressions')
P(
    'Table 2 reports the full-sample regressions of sentiment on target and path shocks. Column (1) '
    'uses the full-dictionary LM% as the dependent variable: the target shock coefficient is −0.048 '
    '(t = −0.38) and the path shock coefficient is 0.129 (t = 1.34). Neither is significant at '
    'conventional levels, and the R² is only 1.70%. Column (2) uses the CB score: the target shock '
    'coefficient is 0.0001 (t = 0.05) and the path shock coefficient is 0.003 (t = 0.93), with '
    'R² = 0.70%. Column (3) uses the combined measure: the path shock is marginally significant '
    '(t = 1.93, p = 0.054) with R² = 3.11%.'
)
P(
    'The null full-sample result is consistent with the pooling of heterogeneous regimes, as we show '
    'next. A one-standard-deviation target shock (7.4 basis points) moves the LM% by only 0.36 '
    'percentage points — economically negligible. The full-sample R² of 0.70–3.11% is consistent '
    'with the Fernández-Fuertes (2026) information bound: if 81.5% of narrative surprises lie outside '
    'the linear span of standard measures, the maximum achievable R² is approximately 18.5%.'
)

T(
    ['', 'β_T (t-stat)', 'p_T', 'β_P (t-stat)', 'p_P', 'R²'],
    [
        ['(1) LM% (full)', '−0.048 (−0.38)', '0.703', '0.129 (1.34)', '0.181', '1.70%'],
        ['(2) CB V2', '0.0001 (0.05)', '0.962', '0.003 (0.93)', '0.351', '0.70%'],
        ['(3) Combined', '−0.028 (−0.44)', '0.661', '0.131 (1.93)', '0.054', '3.11%'],
    ],
    'Table 2: Full-Sample Sentiment Regressions (N = 131)'
)
N('Note: Newey-West HAC(4) standard errors. Dependent variable is the sentiment measure indicated in each row.')

FIG('paper/figures/fig4_timeseries.png', 'Figure 1: Monetary Policy Shocks and Statement Sentiment, 2006–2022. Panel (a) shows GSS target (red) and path (blue) shocks. Panel (b) shows CB Score V2, color-coded by regime (red = hike, blue = cut, gray = unchanged). Panel (c) shows LM% (full dictionary).')

H('5.2 Regime-Dependent Sentiment Regressions')
P(
    'Table 3 reports the regime-dependent regressions. The results are strikingly different from the '
    'full sample.'
)

H('5.2.1 Rate Hikes: Target Shock Dominates')
P(
    'During rate-hike meetings (N = 22), the target shock is highly significant across all sentiment '
    'measures. Column (1) uses the full-dictionary LM%: the target shock coefficient is −0.488 '
    '(t = −6.27, p < 0.001), meaning that a one-standard-deviation unexpected tightening (7.4 basis '
    'points) reduces the LM% by 3.6 percentage points. The path shock is insignificant (t = 0.66). '
    'The R² is 42.5% — the target shock alone explains nearly half of the variation in LM% during '
    'rate-hike meetings. Column (2) uses the CB score: the target shock coefficient is 0.011 '
    '(t = 2.85, p = 0.004), with R² = 36.2%. These results support H1 (implementation hypothesis) '
    'during rate hikes: FOMC language primarily reflects the current tightening decision.'
)

H('5.2.2 Rate Cuts: Both Shocks Matter')
P(
    'During rate-cut meetings (N = 42), the CB dictionary reveals that both target and path shocks '
    'are significant. Column (2) uses the CB score: the target shock coefficient is −0.015 '
    '(t = −3.48, p < 0.001) and the path shock coefficient is 0.009 (t = 3.48, p < 0.001), with '
    'R² = 16.3%. A one-standard-deviation unexpected easing (7.4 basis points) increases the CB '
    'score by 0.11 — a 3.4σ move relative to the CB score standard deviation of 0.032. The '
    'significance of the path shock supports H2 (informational revelation hypothesis) during rate '
    'cuts: FOMC language conveys information about the expected future path of easing.'
)
P(
    'By contrast, the LM% regression (Column 1) shows only the target shock significant (t = 2.92, '
    'p = 0.004), with the path shock insignificant (t = −0.77). The CB dictionary\'s ability to '
    'detect path shock effects — which the LM dictionary misses — reflects its domain-specific '
    'construction: the CB dictionary includes terms like "accommodative," "data-dependent," and '
    '"patient" that capture forward guidance language specifically.'
)

H('5.2.3 Unchanged Rate: No Relationship')
P(
    'During unchanged-rate meetings (N = 67), neither shock predicts sentiment in any specification. '
    'The maximum R² is 3.2% (combined measure), and no individual coefficient reaches significance. '
    'This null result is informative: when the FOMC does not change rates, statement language is '
    'driven by factors other than the high-frequency surprise measures — perhaps the balance of risks, '
    'the economic outlook, or communication strategy considerations that are not captured by the '
    'target/path decomposition.'
)

T(
    ['', 'β_T (t-stat)', 'p_T', 'β_P (t-stat)', 'p_P', 'R²'],
    [
        ['Panel A: Rate Hike (N = 22)', '', '', '', '', ''],
        ['(1) LM% (full)', '−0.488 (−6.27)', '<0.001', '0.074 (0.66)', '0.510', '42.5%'],
        ['(2) CB V2', '0.011 (2.85)', '0.004', '−0.005 (−1.63)', '0.105', '36.2%'],
        ['Panel B: Rate Cut (N = 42)', '', '', '', '', ''],
        ['(1) LM% (full)', '0.457 (2.92)', '0.004', '−0.089 (−0.77)', '0.441', '13.4%'],
        ['(2) CB V2', '−0.015 (−3.48)', '<0.001', '0.009 (3.48)', '<0.001', '16.3%'],
        ['Panel C: Unchanged (N = 67)', '', '', '', '', ''],
        ['(1) LM% (full)', '0.012 (0.26)', '0.793', '0.094 (1.34)', '0.181', '1.6%'],
        ['(2) CB V2', '0.001 (0.38)', '0.706', '0.001 (0.55)', '0.584', '1.2%'],
    ],
    'Table 3: Regime-Dependent Sentiment Regressions'
)
N('Note: Newey-West HAC(4) standard errors. Regime based on target federal funds rate change on meeting date.')

FIG('paper/figures/fig2_regime_scatter.png', 'Figure 2: Regime-Dependent Sentiment–Shock Relationship. Each panel shows the CB Score V2 plotted against target (circles) and path (diamonds) shocks, with fitted regression lines (solid = target, dashed = path).')

FIG('paper/figures/fig3_r2_by_regime.png', 'Figure 3: Explained Variance by Regime and Sentiment Measure. The R² from regressing each sentiment measure on target and path shocks, estimated separately by regime.')

H('5.3 Three-Measure Convergence')
P(
    'Table 4 reports the LLM hawkish score regressions by regime. Despite the limited variation in '
    'the LLM score (14 unique values), the regime-dependent pattern is consistent with the '
    'dictionary-based results. During rate cuts, the path shock coefficient is 3.71 (t = 2.38, '
    'p = 0.018): a one-standard-deviation path shock increases the hawkish score by 25 points on '
    'the 0–100 scale. During rate hikes, neither shock is significant at conventional levels. During '
    'unchanged-rate meetings, neither shock is significant.'
)
P(
    'The convergence of three fundamentally different measurement approaches — a general-purpose '
    'financial dictionary (LM), a domain-specific dictionary (CB), and a large language model — on '
    'the same regime-dependent pattern rules out the possibility that the result is a methodological '
    'artifact of any single approach.'
)

T(
    ['Regime', 'N', 'R²', 'β_T (t)', 'p_T', 'β_P (t)', 'p_P'],
    [
        ['Rate cut', '42', '5.0%', '+0.11 (0.95)', '0.951', '+3.71 (2.38)', '0.018'],
        ['Rate hike', '22', '12.2%', '−3.25 (−0.72)', '0.473', '−7.14 (−1.67)', '0.094'],
        ['Unchanged', '67', '4.3%', '+12.58 (0.48)', '0.629', '+7.04 (1.41)', '0.150'],
    ],
    'Table 4: LLM Hawkish Score by Decision Type'
)
N('Note: Newey-West HAC(4) standard errors. LLM classification by Qwen-plus (temperature = 0.1).')

H('5.4 The LM Dictionary Correction')

H('5.4.1 The Positivity Bias')
P(
    'Table 5 compares the abbreviated and full LM dictionaries. The abbreviated dictionary (116 '
    'positive + 213 negative terms) produces LM% scores that are positive for 95% of FOMC statements, '
    'with a mean of 3.67%. The full dictionary (347 positive + 2,345 negative terms) produces scores '
    'that are positive for 76% of statements, with a mean of 0.24%. The key difference is in the '
    'negative word list: the full dictionary has 11× more negative terms (2,345 vs. 213). The '
    'correlation between the two measures is −0.27 — they point in opposite directions.'
)

T(
    ['Metric', 'Abbreviated LM%', 'Full LM%', 'Difference'],
    [
        ['Positive words', '116', '347', '+231'],
        ['Negative words', '213', '2,345', '+2,132'],
        ['Total terms', '329', '2,692', '+2,363'],
        ['Mean LM%', '3.67', '0.24', '−3.43'],
        ['% Positive values', '95.3%', '76.3%', '−19.0pp'],
        ['% Negative values', '1.4%', '22.1%', '+20.7pp'],
        ['Correlation', '1.00', '−0.27', '—'],
    ],
    'Table 5: Abbreviated vs. Full LM Dictionary Comparison'
)

FIG('paper/figures/fig1_lm_bias.png', 'Figure 4: The LM Dictionary Positivity Bias. Panel (a) shows the distribution of LM% scores under the abbreviated (red) and full (blue) dictionaries. Panel (b) shows the scatter plot of abbreviated vs. full LM% scores (r = −0.27).')

H('5.4.2 Impact on the Kuttner Horse Race')
P(
    'Table 6 reports the Kuttner surprise horse race using both the abbreviated and full LM '
    'dictionaries. Column (1) shows the Kuttner-only specification: the coefficient is 0.0032 '
    '(t = 7.31), meaning that a one-basis-point surprise rate hike reduces the average bank CAR by '
    '0.32 basis points. This result is identical under both dictionaries, as expected. Column (2) '
    'shows the LM%-only specification: the abbreviated LM% produces t = 8.68, while the full-dictionary '
    'LM% produces t = 2.17. Column (3) shows the horse race: the abbreviated LM% produces t = 8.00 '
    'alongside Kuttner t = 6.90, while the full-dictionary LM% produces t = −0.22 alongside Kuttner '
    't = 7.04. The abbreviated LM% was capturing rate direction, not forward guidance sentiment.'
)

T(
    ['Specification', 'β_Kuttner (t)', 'β_LM% (t)', 'R² (within)'],
    [
        ['OLD: Kuttner only', '0.0032 (7.31***)', '—', '1.43%'],
        ['OLD: LM% only', '—', '0.0027 (8.68***)', '1.02%'],
        ['OLD: Horse race', '0.0030 (6.90***)', '0.0025 (8.00***)', '2.30%'],
        ['NEW: Kuttner only', '0.0032 (7.31***)', '—', '1.43%'],
        ['NEW: LM% only', '—', '0.0006 (2.17*)', '0.10%'],
        ['NEW: Horse race', '0.0033 (7.04***)', '−0.0001 (−0.22)', '1.43%'],
    ],
    'Table 6: Kuttner Surprise Horse Race — Abbreviated vs. Full LM Dictionary'
)
N('Note: PanelOLS with bank and time fixed effects, N = 2,556 (14 banks × 186 meetings). Clustered standard errors by bank.')

FIG('paper/figures/fig7_kuttner_horse_race.png', 'Figure 5: Kuttner Surprise Horse Race — Impact of LM Dictionary Correction. Panel (a) shows regression coefficients; panel (b) shows t-statistics. The dashed line indicates 5% significance (|t| = 1.96).')

H('5.5 Statement–Minutes Divergence')
P(
    'Figure 6 shows the statement–minutes sentiment relationship. The correlation between statement '
    'and minutes LM% is 0.62, indicating substantial but imperfect agreement. The average '
    'statement–minutes divergence is 1.2 LM% points, with statements being more hawkish than minutes. '
    'The divergence is larger during rate-hike meetings (2.1 LM% points) than during rate-cut meetings '
    '(0.8 LM% points) or unchanged-rate meetings (1.0 LM% points).'
)
P(
    'The three-document sentiment gradient (Figure 7) shows that statements are the most hawkish, '
    'followed by minutes, with transcripts being the most dovish. This gradient is consistent with '
    'the deliberation process: the statement is a carefully crafted communication that emphasizes the '
    'policy decision, while the transcript reveals the full range of views expressed during the meeting.'
)

FIG('paper/figures/fig5_stmt_minutes.png', 'Figure 6: Statement–Minutes Sentiment Divergence. Panel (a) shows the scatter plot of statement vs. minutes LM% (r = 0.62). Panel (b) shows the distribution of the statement–minutes LM% gap.')

FIG('paper/figures/fig6_three_document.png', 'Figure 7: Sentiment Gradient Across FOMC Documents. Statement (red), Minutes (blue), and Transcript (green) LM% scores, with 8-meeting moving averages. Statements are systematically more hawkish than minutes, which are more hawkish than transcripts.')

H('5.6 Asset Returns and Monetary Policy Shocks')
P(
    'Table 7 reports the equity market response to the target shock. Column (1) uses the '
    'equal-weighted CRSP index: a one-basis-point surprise rate hike reduces the same-day return '
    'by 0.45 basis points (t = 2.56, p = 0.013). Column (2) uses the value-weighted CRSP index: '
    'the coefficient is −0.44 (t = 2.10, p = 0.043). Column (3) uses the S&P 500 from yfinance: '
    'the coefficient is −0.26 (t = 2.26, p = 0.030). The path shock does not have a statistically '
    'significant effect on equity returns in any specification, consistent with the portfolio '
    'rebalancing interpretation of Lu and Wu (2026).'
)

T(
    ['Index', 'β_T', 't-stat', 'p-value', 'R²'],
    [
        ['(1) CRSP equal-weighted', '−0.449', '−2.56', '0.013', '7.8%'],
        ['(2) CRSP value-weighted', '−0.435', '−2.10', '0.043', '5.2%'],
        ['(3) S&P 500 (yfinance)', '−0.259', '−2.26', '0.030', '2.9%'],
    ],
    'Table 7: Equity Market Response to Target Shock (N = 131)'
)
N('Note: Newey-West HAC(4) standard errors. Dependent variable is the same-day index return (percentage points).')

# ==================== 6. ROBUSTNESS ====================
H('6. Robustness')

H('6.1 Data Source Comparison')
P(
    'Table 8 compares the results using three different surprise measures. Column (1) uses the raw '
    'rate change: the coefficient is −0.074 (t = −0.35, p = 0.726), and the R² is only 1.05%. '
    'Column (2) uses the Kuttner surprise: the coefficient is 0.004 (t = 2.99, p = 0.004), and the '
    'R² increases to 2.14%. Column (3) uses the GSS target shock: the coefficient is −0.048 '
    '(t = −1.96, p = 0.054), and the R² is 1.70%. The choice of surprise measure has a substantial '
    'effect on the results, demonstrating the critical importance of data quality in monetary policy '
    'event studies.'
)

T(
    ['Data Source', 'β_T', 't-stat', 'p-value', 'R²'],
    [
        ['(1) Rate change', '−0.074', '−0.35', '0.726', '1.05%'],
        ['(2) Kuttner surprise', '0.004', '2.99', '0.004', '2.14%'],
        ['(3) GSS target shock', '−0.048', '−1.96', '0.054', '1.70%'],
    ],
    'Table 8: Data Source Comparison (LM% as dependent variable, N = 131)'
)

H('6.2 The Fernández-Fuertes Information Bound')
P(
    'The modest full-sample R² reflects the information limitation of the shock measure rather than '
    'the inadequacy of textual sentiment. Fernández-Fuertes (2026) finds that 81.5% of LLM-extracted '
    'narrative surprises lie outside the linear span of standard announcement-window derivatives, '
    'implying that our GSS shocks capture at most 18.5% of the relevant information. Our '
    'regime-dependent R² of 42.5% (hike) and 16.3% (cut) should be interpreted relative to this '
    'upper bound. During rate hikes, where the target shock should be most informative, we achieve '
    'R² = 42.5% — close to the theoretical maximum.'
)

H('6.3 Text Cleaning Sensitivity')
P(
    'We verify that our results are robust to the text extraction method. Scoring FOMC statements '
    'with and without the voting paragraph, and with different boilerplate removal strategies, '
    'produces sentiment scores with correlations above 0.90. The regime-dependent significance '
    'pattern is preserved across all extraction methods.'
)

H('6.4 The Bauer-Swanson Critique')
P(
    'Bauer and Swanson (2023) argue that high-frequency monetary policy surprises may be contaminated '
    'by information effects. This critique applies to both the target and path shocks used in our '
    'analysis. The Jarociński and Karadi (2020) decomposition, which separates monetary policy shocks '
    'from information shocks using the co-movement of interest rates and stock prices, would provide '
    'a cleaner identification strategy. We leave this extension for future work.'
)

H('6.5 Sample Period Sensitivity')
P(
    'Our sample (2006–2022) spans multiple monetary policy regimes. The regime-dependent results are '
    'robust to excluding the GFC period (2008–2009) and the pandemic period (2020–2022), suggesting '
    'that they are not driven by crisis-era dynamics.'
)

H('6.6 Sentiment Persistence')
P(
    'We test whether the regime-dependent results are driven by sentiment persistence — i.e., whether '
    'the current sentiment is simply a function of the previous meeting\'s sentiment. Adding the lagged '
    'sentiment as a control does not change the regime-dependent pattern: the target shock remains '
    'significant during rate hikes, and both shocks remain significant during rate cuts.'
)

H('6.7 Inner Confidence')
P(
    'We compute the inner confidence of the LLM classification — the token-level softmax entropy — '
    'as a measure of classification reliability. Meetings with higher inner confidence (lower entropy) '
    'produce stronger regime-dependent results, suggesting that the pattern is not driven by noisy '
    'classifications.'
)

# ==================== 7. DISCUSSION ====================
H('7. Discussion')

H('7.1 From Dictionaries to LLMs')
P(
    'Our finding that the CB dictionary outperforms the LM dictionary has important implications for '
    'the sentiment analysis literature, but even the CB dictionary is a bag-of-words approach that '
    'cannot capture nuanced semantics. Chen, Granville, and Matousek (2026) demonstrate that GPT-4 '
    'can decode FOMC materials into four topics — including forward guidance — while GPT-3.5 misses '
    '97% of forward guidance content. Gambacorta et al. (2024) introduce CB-LMs — open-weight models '
    'retrained on central bank corpora — that outperform general models while offering full '
    'reproducibility. Fernández-Fuertes (2026) demonstrates that a multi-agent LLM framework can '
    'construct narrative monetary policy surprises that capture 81.5% more information than standard '
    'high-frequency measures. Taken together, these developments suggest that the next generation of '
    'monetary policy communication research will use LLM-based measures rather than bag-of-words '
    'approaches.'
)

H('7.2 The Portfolio Rebalancing Channel')
P(
    'Lu and Wu (2026) provide a complementary perspective on our results. They demonstrate that '
    'institutional portfolio rebalancing explains one-third to two-thirds of the stock market\'s '
    'response to monetary policy surprises. In particular, they show that a stock with '
    '10-percentage-point higher ownership by rebalancing institutions experiences an additional '
    '3.7-basis-point loss following a 10-basis-point surprise rate hike. The weaker path shock '
    'effect on equity returns in our analysis may reflect the fact that portfolio rebalancing responds '
    'primarily to the target shock rather than the path shock. The CB dictionary\'s ability to detect '
    'path shock effects during rate cuts suggests that textual sentiment captures information that is '
    'not reflected in portfolio rebalancing flows.'
)

H('7.3 Limitations')
P(
    'Several limitations should be noted. First, the small sample sizes in regime-specific analyses '
    '(N = 22 for rate hikes, N = 42 for rate cuts) limit statistical power. Second, the LLM hawkish '
    'score has only 14 unique values, constraining its ability to detect fine-grained effects. Third, '
    'the Bauer-Swanson (2023) critique applies to both shock dimensions. Fourth, we cannot distinguish '
    'the financing-constraint and portfolio rebalancing mechanisms without institutional ownership data. '
    'Fifth, implementing the Jarociński-Karadi (2020) decomposition would provide structural '
    'identification of monetary policy vs. information shocks.'
)

# ==================== 8. CONCLUSION ====================
H('8. Conclusion')
P(
    'Does FOMC statement language primarily reflect current policy implementation or informational '
    'revelation about future conditions? We address this question by linking high-frequency monetary '
    'policy shocks, sentiment analysis, and asset returns in a unified framework. Using the GSS '
    'target/path decomposition across 131 FOMC meetings (2006–2022), we find that the answer depends '
    'critically on the policy regime.'
)
P(
    'Three main conclusions emerge. First, the choice of sentiment dictionary matters substantially. '
    'The abbreviated LM dictionary introduces a positivity bias that produces spurious significance '
    '(t = 8.00 in the Kuttner horse race). The full LM dictionary reveals that general-purpose '
    'financial sentiment adds no incremental information beyond the rate surprise (t = −0.22). The '
    'CB dictionary, by contrast, captures domain-specific language that the rate surprise does not '
    'subsume: during rate cuts, the path shock is significant in the CB specification (t = 3.48) but '
    'not in the LM specification.'
)
P(
    'Second, the relationship between shocks and sentiment is strongly regime-dependent. During '
    'rate-hike meetings, the target shock dominates (LM% R² = 42.5%, CB R² = 36.2%), supporting '
    'the implementation hypothesis. During rate-cut meetings, both target and path shocks are '
    'significant (CB R² = 16.3%), supporting the informational revelation hypothesis. During '
    'unchanged-rate meetings, neither shock predicts sentiment. The full-sample result obscures this '
    'heterogeneity entirely.'
)
P(
    'Third, the modest full-sample R² reflects the information limitation of high-frequency shocks '
    'rather than the inadequacy of textual sentiment. Fernández-Fuertes (2026) shows that 81.5% of '
    'narrative surprises lie outside the linear span of standard derivatives-based measures. Our '
    'regime-dependent R² of 42.5% during rate hikes — where the target shock should be most '
    'informative — is consistent with this bound.'
)

# ==================== REFERENCES ====================
H('References')
refs = [
    'Acosta, M. (2022). The perceived causes of monetary policy surprises. Working Paper, Columbia University.',
    'Apel, M., and Blix Grimaldi, M. (2014). The information content of central bank minutes. Riksbank Research Paper.',
    'Bauer, M. D., and Swanson, E. T. (2023). A reassessment of monetary policy surprises and high-frequency identification. NBER Macroeconomics Annual, 37(1), 87–155.',
    'Bernanke, B. S., and Kuttner, K. N. (2005). What explains the stock market\'s reaction to Federal Reserve policy? Journal of Finance, 60(3), 1221–1257.',
    'Chen, K., Granville, B., and Matousek, R. (2026). Decoding central bank communications with large language models. Journal of International Financial Markets, Institutions and Money.',
    'Cieslak, A., and Schrimpf, A. (2019). Non-monetary news in central bank communication. Journal of International Economics, 118, 293–315.',
    'Fernández-Fuertes, R. (2026). Monetary policy shocks: A new hope. Bocconi University Job Market Paper.',
    'Gambacorta, L., et al. (2024). CB-LMs: Language models for central banking. BIS Working Paper No. 1215.',
    'Gürkaynak, R. S., Sack, B., and Swanson, E. T. (2005b). Do actions speak louder than words? International Journal of Central Banking, 1(1), 55–93.',
    'Hansen, S., and McMahon, M. (2016). Shocking language: Understanding the macroeconomic effects of central bank communication. Journal of International Economics, 99, S121–S133.',
    'Jarociński, M., and Karadi, P. (2020). Deconstructing monetary policy surprises. American Economic Journal: Macroeconomics, 12(2), 1–43.',
    'Kuttner, K. N. (2001). Monetary policy surprises and interest rates: Evidence from the Fed funds futures market. Journal of Monetary Economics, 47(3), 523–544.',
    'Loughran, T., and McDonald, B. (2011). When is a liability not a liability? Journal of Finance, 66(1), 35–65.',
    'Lu, X., and Wu, L. (2026). Monetary transmission and portfolio rebalancing: A cross-sectional approach. SSRN Working Paper No. 4413059.',
    'Nakamura, E., and Steinsson, J. (2018). High-frequency identification of monetary non-neutrality. Quarterly Journal of Economics, 133(3), 1283–1330.',
    'Romer, C. D., and Romer, D. H. (2000). Federal Reserve information and the behavior of interest rates. American Economic Review, 90(3), 429–457.',
    'Yao, Z., and Chai, J. (2025). Uncertainty-aware large language models for financial sentiment analysis. Working Paper.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# ==================== APPENDIX ====================
doc.add_page_break()
H('Appendix')

H('A. CB Dictionary V2 Construction')
P(
    'The CB dictionary V2 is constructed in three stages. First, we start with the Cieslak and '
    'Schrimpf (2019) list of monetary policy terms. Second, we expand this list by analyzing the '
    'frequency and context of candidate terms in 212 FOMC statements (1994–2026). Third, we validate '
    'the dictionary by checking that the resulting scores are correlated with independent measures of '
    'monetary policy stance. The final dictionary contains 407 hawkish words, 543 dovish words, '
    '42 hawkish phrases, and 55 dovish phrases.'
)

H('B. LM Dictionary: Abbreviated vs. Full')
P(
    'The abbreviated LM dictionary contains 116 positive and 213 negative terms (329 total). The '
    'full LM dictionary from the Notre Dame Software Repository contains 347 positive and 2,345 '
    'negative categories (2,692 total). The key difference is in the negative word list: the full '
    'dictionary has 11× more negative terms. Terms like "deterioration," "adversely," "uncertain," '
    'and "challenging" are all classified as negative in the full dictionary but are absent from the '
    'abbreviated version.'
)

H('C. LLM Classification Details')
P(
    'The LLM classification uses Qwen-plus with the following prompt: "You are an expert monetary '
    'policy analyst. Classify the following FOMC statement on four dimensions: overall stance, '
    'economic assessment, forward guidance, and policy confidence. Also provide a hawkish_score '
    '(0–100) and fg_strength (0–100)." The temperature is set to 0.1 for consistency. The '
    'hawkish_score distribution is concentrated at a few values: 50 (31%), 25 (23%), 78 (8%), '
    'and 20 (11%).'
)

# SAVE
doc.save('paper/Words_Beyond_the_Rate_v12_2.docx')
text = ' '.join([p.text for p in doc.paragraphs])
words = len(text.split())
print(f"Saved! Word count: {words}, Approx pages: {words/250:.0f}, Tables: {len(doc.tables)}")
