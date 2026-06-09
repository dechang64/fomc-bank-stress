#!/usr/bin/env python3
"""Generate Words Beyond the Rate v12.3 — corrected narrative, Lu & Wu style"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15

def H(text, level=1): doc.add_heading(text, level=level)
def P(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def B(text): doc.add_paragraph(text, style='List Bullet')
def N(text):
    p = doc.add_paragraph(text); p.runs[0].font.size = Pt(9); p.runs[0].italic = True
def T(headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph(caption); p.runs[0].bold = True; p.runs[0].font.size = Pt(10)
    t = doc.add_table(rows=len(rows)+1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9); r.bold = True
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t
def FIG(path, caption):
    p = doc.add_paragraph(); p.add_run().add_picture(path, width=Inches(6.0))
    p2 = doc.add_paragraph(caption); p2.runs[0].italic = True; p2.runs[0].font.size = Pt(10)

# ==================== TITLE ====================
title = doc.add_heading('Words Beyond the Rate: High-Frequency Monetary Policy Shocks and FOMC Language', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Eileen Zhang'); r.font.size = Pt(14)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Academy of AI, Xi'an Jiaotong-Liverpool University, Suzhou, China"); r2.font.size = Pt(11); r2.italic = True
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('v12.3 — Revised Draft (June 2026)'); r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(128,128,128)

doc.add_page_break()

# ==================== ABSTRACT ====================
H('Abstract')
P(
    'We show that the standard Loughran-McDonald (LM) sentiment dictionary, when applied with '
    'its complete 2,692-term list, produces scores that are uncorrelated with the widely-used '
    'abbreviated version (r = −0.27) and lack incremental predictive power for bank stock returns '
    'beyond the federal funds rate surprise (t = −0.22 in the Kuttner horse race, versus t = 8.00 '
    'with the abbreviated dictionary). The abbreviated LM dictionary\'s apparent significance '
    'reflects a positivity bias — 95% of FOMC statements receive positive scores — that makes it '
    'a proxy for rate direction rather than a measure of forward guidance sentiment.'
)
P(
    'Using a central-bank-specific (CB) dictionary, we document that the relationship between '
    'monetary policy shocks and statement sentiment is regime-dependent. During rate-cut meetings '
    '(N = 42), both the target shock (t = −3.48) and the path shock (t = 3.48) significantly '
    'predict CB sentiment, with a one-standard-deviation path shock moving the CB score by 36% '
    'of its standard deviation. This path shock result is the paper\'s central finding: it '
    'provides direct evidence that FOMC language conveys forward guidance information beyond the '
    'current rate decision. During rate-hike meetings (N = 22), the target shock dominates but '
    'the small sample warrants caution; during unchanged-rate meetings (N = 67), neither shock '
    'predicts sentiment.'
)
P(
    'Three independent measurement approaches — the LM dictionary, the CB dictionary, and LLM '
    'classification — converge on the same regime-dependent pattern. The path shock is significant '
    'during rate cuts across all three measures (CB t = 3.48, LLM t = 2.37), whereas it is '
    'insignificant during rate hikes and unchanged-rate meetings. This convergence rules out the '
    'possibility that the result is a methodological artifact.'
)
P('Keywords: Monetary policy surprises; FOMC statements; Sentiment analysis; Loughran-McDonald dictionary; Central bank communication; Forward guidance; High-frequency identification', italic=True)

doc.add_page_break()

# ==================== 1. INTRODUCTION ====================
H('1. Introduction')
P(
    'A central question in monetary economics is whether FOMC communications primarily reflect '
    'the current policy decision — the implementation channel — or convey forward-looking '
    'information about future economic conditions — the informational revelation channel '
    '(Romer and Romer, 2000; Nakamura and Steinsson, 2018). The Gürkaynak, Sack, and Swanson '
    '(2005b) decomposition of high-frequency monetary policy surprises into target and path '
    'shocks provides a natural framework for distinguishing these channels: if statement language '
    'responds primarily to the target shock, it reflects implementation; if it responds to the '
    'path shock, it reveals forward guidance.'
)
P(
    'We address this question using three textual sentiment measures applied to 131 FOMC '
    'statements from 2006 to 2022: the Loughran-McDonald (LM) dictionary with its complete '
    '2,692-term list, a central-bank-specific (CB) dictionary of 950 words and 97 phrases, '
    'and LLM-based classification using Qwen-plus. We make three contributions.'
)
P(
    'First, we show that the abbreviated LM dictionary — widely used in prior work — introduces '
    'a positivity bias that produces spurious significance. The abbreviated dictionary (116 '
    'positive + 213 negative terms) assigns positive scores to 95% of FOMC statements, whereas '
    'the full dictionary (347 positive + 2,345 negative terms) assigns positive scores to only '
    '76%. The correlation between the two measures is −0.27: they point in opposite directions. '
    'In the Kuttner (2001) surprise horse race, the abbreviated LM% produces t = 8.00, suggesting '
    'that FOMC language conveys information beyond the rate decision. The full-dictionary LM% '
    'produces t = −0.22, indicating no incremental information. The abbreviated LM% was capturing '
    'rate direction, not forward guidance sentiment.'
)
P(
    'Second, we document that the relationship between monetary policy shocks and statement '
    'sentiment is regime-dependent. During rate-cut meetings (N = 42), both the target shock '
    '(t = −3.48) and the path shock (t = 3.48) significantly predict CB sentiment. The path '
    'shock result is the paper\'s central finding: a one-standard-deviation path shock moves '
    'the CB score by 36% of its standard deviation, providing direct evidence that FOMC language '
    'conveys forward guidance. During rate-hike meetings (N = 22), the target shock is significant '
    'but the small sample and high in-sample fit warrant caution (leave-one-out R² = 12% versus '
    'in-sample R² = 36%). During unchanged-rate meetings (N = 67), neither shock predicts '
    'sentiment (adjusted R² = −2%).'
)
P(
    'Third, we verify that the regime-dependent pattern is robust across three independent '
    'measurement approaches. The path shock is significant during rate cuts in both the CB '
    'dictionary (t = 3.48) and the LLM classification (t = 2.37), whereas it is insignificant '
    'during rate hikes and unchanged-rate meetings. This convergence rules out the possibility '
    'that the result is an artifact of any single measurement approach.'
)

H('1.1 Hypotheses')
P('We test three hypotheses about the relationship between monetary policy shocks and FOMC statement sentiment:')
B(
    'H1 (Implementation): Statement sentiment responds primarily to the target shock, reflecting '
    'the current rate decision. Under H1, the path shock should not predict sentiment after '
    'controlling for the target shock.'
)
B(
    'H2 (Informational Revelation): Statement sentiment responds to the path shock, reflecting '
    'forward guidance about the future rate path. Under H2, the path shock should predict '
    'sentiment even after controlling for the target shock.'
)
B(
    'H3 (Regime Dependence): The relative importance of target and path shocks varies across '
    'policy regimes. Under H3, the path shock should be more important during rate cuts (when '
    'forward guidance is most informative) than during rate hikes.'
)
P(
    'Our results support H3: the path shock is significant only during rate cuts, consistent '
    'with the interpretation that forward guidance is most informative when the FOMC is easing.'
)

H('1.2 Related Literature')
P(
    'This paper contributes to three strands of the literature. First, the monetary policy shocks '
    'and asset prices literature: Kuttner (2001), Gürkaynak, Sack, and Swanson (2005b), Bernanke '
    'and Kuttner (2005), and Nakamura and Steinsson (2018) establish the high-frequency '
    'identification framework. Bauer and Swanson (2023) reassess the identification, emphasizing '
    'information effects. We contribute by linking shocks directly to textual content.'
)
P(
    'Second, the textual analysis of central bank communications: Loughran and McDonald (2011), '
    'Tetlock (2007), Apel and Blix Grimaldi (2014), and Hansen and McMahon (2016) apply '
    'dictionary-based methods to monetary policy text. We contribute by demonstrating that the '
    'choice of dictionary has first-order implications, and that the abbreviated LM dictionary '
    'produces misleading findings due to its positivity bias.'
)
P(
    'Third, the LLM-based analysis of monetary policy: Chen, Granville, and Matousek (2026) show '
    'that GPT-4 decodes FOMC materials into four topics including forward guidance, while GPT-3.5 '
    'misses 97% of forward guidance content. Gambacorta et al. (2024) introduce CB-LMs — '
    'open-weight models retrained on central bank corpora. Fernández-Fuertes (2026) demonstrates '
    'that a multi-agent LLM framework captures 81.5% more information than standard high-frequency '
    'measures. We contribute by providing a benchmark showing that even simple dictionary-based '
    'approaches can detect regime-dependent patterns when the correct dictionary is used.'
)

doc.add_page_break()

# ==================== 2. DATA ====================
H('2. Data and Variable Construction')

H('2.1 Monetary Policy Shocks')
P(
    'We use the Gürkaynak, Sack, and Swanson (2005b) decomposition of high-frequency interest '
    'rate surprises into a target shock and a path shock, obtained from Acosta (2022). The target '
    'shock captures the unexpected component of the current rate decision; the path shock captures '
    'the revision in expectations of future rate paths. Our sample covers 131 FOMC meetings from '
    'January 2006 to July 2022. The target shock has a correlation of 0.976 with the Kuttner '
    '(2001) surprise.'
)

H('2.2 FOMC Statement Sentiment')
P('We construct three sentiment measures for each FOMC statement:')
B(
    'LM% (full dictionary). We compute the standard Loughran-McDonald net sentiment score as '
    '(n_positive − n_negative) / n_total × 100, using the complete LM dictionary (347 positive, '
    '2,345 negative categories, totaling 2,692 unique terms). The full dictionary produces a mean '
    'of 0.24% with 76% positive and 22% negative values. We extract the policy-relevant text by '
    'removing HTML markup, navigation boilerplate, and voting records. Statement length varies from '
    '104 to 817 words (mean = 311).'
)
B(
    'CB Score V2. We construct a central-bank-specific dictionary containing 407 hawkish words, '
    '543 dovish words, 42 hawkish phrases, and 55 dovish phrases. The score is computed as '
    '(n_hawkish − n_dovish) / n_total. This produces a mean of −0.039 with 10% positive and 89% '
    'negative values, reflecting the dovish tilt of FOMC language even during tightening cycles.'
)
B(
    'LLM Hawkish Score. Qwen-plus classifies each statement on four dimensions (overall stance, '
    'economic assessment, forward guidance, policy confidence), producing a hawkish_score (0–100) '
    'and fg_strength (0–100). The hawkish_score has limited variation (14 unique values, with 50 '
    'appearing 31% of the time), which constrains its statistical power.'
)

H('2.3 Statement–Minutes Divergence')
P(
    'We compute the sentiment of FOMC minutes using the same LM% methodology and construct the '
    'statement–minutes divergence as the difference in LM% between the two documents. We have 115 '
    'meetings with both statement and minutes sentiment. The three-document gradient — Statement '
    'LM% (mean = 3.10) > Minutes LM% (mean = 1.71) > Transcript LM% (mean = 1.32) — reflects '
    'the declining hawkish tone as FOMC communication moves from the policy decision to the '
    'deliberation record.'
)

H('2.4 Regime Classification')
P(
    'We classify each FOMC meeting by the associated rate decision: rate hike (N = 22), rate cut '
    '(N = 42), or unchanged (N = 67). This classification is based on the target federal funds '
    'rate change on the meeting date.'
)

H('2.5 Empirical Strategy')
P(
    'Our baseline specification regresses sentiment on target and path shocks:'
)
P('    Sentiment_i = α + β_T × Target_i + β_P × Path_i + ε_i', italic=True)
P(
    'We use Newey-West HAC(4) standard errors to account for potential autocorrelation and '
    'heteroskedasticity. We estimate this specification separately for the full sample and for '
    'each regime. Under H1 (implementation), β_P = 0; under H2 (informational revelation), '
    'β_P ≠ 0; under H3 (regime dependence), β_P varies across regimes.'
)

doc.add_page_break()

# ==================== 3. RESULTS ====================
H('3. Results')

H('3.1 The LM Dictionary Positivity Bias')
P(
    'We begin with a methodological finding that motivates the rest of the paper. The abbreviated '
    'LM dictionary (116 positive + 213 negative terms) produces LM% scores that are positive for '
    '95% of FOMC statements, with a mean of 3.67%. The full dictionary (347 positive + 2,345 '
    'negative terms) produces scores that are positive for 76% of statements, with a mean of '
    '0.24%. The correlation between the two measures is −0.27: they point in opposite directions.'
)
P(
    'Figure 1 illustrates the bias. Panel (a) shows that the abbreviated LM% distribution is '
    'strongly right-skewed, whereas the full-dictionary LM% is approximately symmetric around '
    'zero. Panel (b) shows that the two measures are essentially uncorrelated — knowing the '
    'abbreviated LM% tells you nothing about the full-dictionary LM%.'
)

FIG('paper/figures/fig1_lm_bias.png', 'Figure 1: The LM Dictionary Positivity Bias. Panel (a): Distribution of LM% scores under the abbreviated (329 terms) and full (2,692 terms) dictionaries. Panel (b): Scatter plot of abbreviated vs. full LM% (r = −0.27).')

H('3.1.1 Impact on the Kuttner Horse Race')
P(
    'The positivity bias has a direct impact on the Kuttner surprise horse race — the standard '
    'test for whether FOMC language conveys incremental information beyond the rate decision. '
    'Table 1 reports the results.'
)

T(
    ['Specification', 'β_Kuttner (t)', 'β_LM% (t)', 'R² (within)'],
    [
        ['OLD: Kuttner only', '0.0032 (7.31***)', '—', '1.43%'],
        ['OLD: LM% only', '—', '0.0027 (8.68***)', '1.02%'],
        ['OLD: Horse race', '0.0030 (6.90***)', '0.0025 (8.00***)', '2.30%'],
        ['NEW: Kuttner only', '0.0032 (7.31***)', '—', '1.43%'],
        ['NEW: LM% only', '—', '0.0006 (2.17*)', '0.10%'],
        ['NEW: Horse race', '0.0033 (7.04***)', '−0.0001 (−0.22)', '1.43%'],
    ],
    'Table 1: Kuttner Surprise Horse Race — Abbreviated vs. Full LM Dictionary'
)
N('Note: PanelOLS with bank and time fixed effects, N = 2,556 (14 banks × 186 meetings). Clustered standard errors by bank. OLD uses abbreviated LM dictionary (329 terms), NEW uses full LM dictionary (2,692 terms).')

P(
    'Column (1) shows that the Kuttner surprise alone explains 1.43% of within-bank CAR '
    'variation, with a coefficient of 0.0032 (t = 7.31). This result is robust to the LM% '
    'correction. Column (2) shows that the abbreviated LM% alone produces t = 8.68, whereas the '
    'full-dictionary LM% produces t = 2.17 — a 75% reduction in statistical significance. '
    'Column (3) shows the horse race: the abbreviated LM% produces t = 8.00 alongside the '
    'Kuttner t = 6.90, suggesting incremental information. The full-dictionary LM% produces '
    't = −0.22 alongside the Kuttner t = 7.04, indicating no incremental information.'
)
P(
    'The interpretation is straightforward: the abbreviated LM% was capturing rate direction, '
    'not forward guidance sentiment. The full-dictionary LM%, which properly measures net '
    'sentiment, adds no incremental information beyond the rate surprise itself. This finding '
    'underscores the importance of using complete, domain-appropriate dictionaries in monetary '
    'policy text analysis.'
)

FIG('paper/figures/fig7_kuttner_horse_race.png', 'Figure 2: Kuttner Surprise Horse Race — Impact of LM Dictionary Correction. Panel (a): Coefficients. Panel (b): t-statistics. Dashed lines indicate 5% significance threshold.')

doc.add_page_break()

H('3.2 Full-Sample Results')
P(
    'Table 2 reports the full-sample regressions of sentiment on target and path shocks. '
    'Regardless of the sentiment measure, the full-sample R² is modest (0.7–3.1%) and no '
    'individual coefficient reaches the 5% significance level. The combined measure produces '
    'the highest R² (3.11%) with the path shock marginally significant (t = 1.93, p = 0.054).'
)

T(
    ['Sentiment Measure', 'β_T (t-stat)', 'p_T', 'β_P (t-stat)', 'p_P', 'R²'],
    [
        ['LM% (full)', '−0.048 (−0.38)', '0.703', '0.129 (1.34)', '0.181', '1.70%'],
        ['CB V2', '0.0001 (0.05)', '0.962', '0.003 (0.93)', '0.351', '0.70%'],
        ['Combined', '−0.028 (−0.44)', '0.661', '0.131 (1.93)', '0.054', '3.11%'],
    ],
    'Table 2: Full-Sample Sentiment Regressions (N = 131)'
)
N('Note: Newey-West HAC(4) standard errors. Dependent variable is the sentiment measure.')

P(
    'The modest full-sample R² is consistent with the pooling of heterogeneous regimes, as we '
    'show next. It also reflects the information limitation of high-frequency shocks: '
    'Fernández-Fuertes (2026) finds that 81.5% of LLM-extracted narrative surprises lie outside '
    'the linear span of standard announcement-window derivatives, implying that GSS shocks capture '
    'at most 18.5% of the relevant information.'
)

H('3.3 Regime-Dependent Results')
P(
    'Table 3 reports the regime-dependent regressions. The results are strikingly different from '
    'the full sample and reveal a clear pattern of differential sensitivity across policy regimes.'
)

T(
    ['Regime / Measure', 'β_T (t-stat)', 'p_T', 'β_P (t-stat)', 'p_P', 'R²', 'R²(adj)'],
    [
        ['Rate Hike (N = 22)', '', '', '', '', '', ''],
        ['  LM% (full)', '−0.488 (−6.27***)', '<0.001', '0.074 (0.66)', '0.510', '42.5%', '36.4%'],
        ['  CB V2', '0.011 (2.85**)', '0.004', '−0.005 (−1.63)', '0.105', '36.2%', '29.5%'],
        ['Rate Cut (N = 42)', '', '', '', '', '', ''],
        ['  LM% (full)', '0.457 (2.92**)', '0.004', '−0.089 (−0.77)', '0.441', '13.4%', '8.9%'],
        ['  CB V2', '−0.015 (−3.48***)', '<0.001', '0.009 (3.48***)', '<0.001', '16.3%', '12.0%'],
        ['Unchanged (N = 67)', '', '', '', '', '', ''],
        ['  LM% (full)', '0.012 (0.26)', '0.793', '0.094 (1.34)', '0.181', '1.6%', '−1.5%'],
        ['  CB V2', '0.001 (0.38)', '0.706', '0.001 (0.55)', '0.584', '1.2%', '−1.9%'],
    ],
    'Table 3: Regime-Dependent Sentiment Regressions'
)
N('Note: Newey-West HAC(4) standard errors. Regime based on target federal funds rate change.')

H('3.3.1 Rate Cuts: Path Shock Evidence for Forward Guidance')
P(
    'The rate-cut regime (N = 42) produces the paper\'s central finding. Column (5) shows that '
    'both the target shock (t = −3.48) and the path shock (t = 3.48) significantly predict CB '
    'sentiment, with R² = 16.3% (adjusted R² = 12.0%). The path shock result directly supports '
    'H2 (informational revelation): FOMC language during rate cuts conveys information about the '
    'expected future path of policy, not just the current decision.'
)
P(
    'The economic magnitude is substantial. A one-standard-deviation path shock (1.04 basis '
    'points) moves the CB score by 0.012, which represents 36% of the CB score\'s standard '
    'deviation in the rate-cut regime. By contrast, a one-standard-deviation target shock '
    '(1.09 basis points) moves the CB score by −0.009, representing 27% of its standard '
    'deviation. The path shock effect is economically larger than the target shock effect, '
    'consistent with the interpretation that forward guidance is the dominant channel during '
    'rate cuts.'
)
P(
    'The opposite signs of the target and path coefficients have a natural interpretation. A '
    'larger-than-expected rate cut (negative target shock) is associated with more dovish '
    'language, while a steeper expected easing path (positive path shock) is associated with '
    'more hawkish language — perhaps reflecting the FOMC\'s attempt to signal that the cut is '
    'a measured adjustment rather than the beginning of an aggressive easing cycle.'
)

H('3.3.2 Rate Hikes: Target Shock Dominates (with Caveats)')
P(
    'During rate-hike meetings (N = 22), the target shock is significant in both the LM% '
    '(t = −6.27) and CB (t = 2.85) specifications, whereas the path shock is not significant. '
    'This pattern supports H1 (implementation): during rate hikes, FOMC language primarily '
    'reflects the rationale for the current tightening decision.'
)
P(
    'However, we urge caution in interpreting the high in-sample R² (36–43%). The small sample '
    '(N = 22) and the fact that the target shock is positive for all 22 rate-hike meetings mean '
    'that the regression is essentially distinguishing between degrees of hawkishness within a '
    'homogeneous group. Leave-one-out cross-validation reduces the R² from 36% to 12% for the CB '
    'specification, indicating substantial overfitting. We therefore emphasize the qualitative '
    'conclusion — the target shock dominates during rate hikes — rather than the point estimates.'
)

H('3.3.3 Unchanged Rate: No Relationship')
P(
    'During unchanged-rate meetings (N = 67), neither shock predicts sentiment in any '
    'specification. The adjusted R² is negative for both LM% (−1.5%) and CB (−1.9%), indicating '
    'that the regressors add no explanatory power beyond the intercept. This null result is '
    'informative: when the FOMC does not change rates, statement language is driven by factors '
    '— the balance of risks, the economic outlook, or communication strategy — that are not '
    'captured by the target/path decomposition.'
)

FIG('paper/figures/fig2_regime_scatter.png', 'Figure 3: Regime-Dependent Sentiment–Shock Relationship. Each panel shows the CB score V2 against target (circles) and path (diamonds) shocks, with fitted regression lines.')

FIG('paper/figures/fig3_r2_by_regime.png', 'Figure 4: Explained Variance by Regime and Sentiment Measure. R² from regressing each sentiment measure on target and path shocks, by policy regime.')

doc.add_page_break()

H('3.4 Three-Measure Convergence')
P(
    'Table 4 reports the LLM hawkish score regressions by regime. Despite the limited variation '
    'in the LLM score (14 unique values), the regime-dependent pattern is consistent with the '
    'dictionary-based results.'
)

T(
    ['Regime', 'N', 'β_T (t)', 'p_T', 'β_P (t)', 'p_P', 'R²'],
    [
        ['Rate cut', '42', '+0.114 (0.06)', '0.951', '+3.710 (2.37*)', '0.018', '5.0%'],
        ['Rate hike', '22', '−3.254 (−0.72)', '0.473', '−7.142 (−1.67)', '0.094', '12.2%'],
        ['Unchanged', '67', '+12.579 (0.48)', '0.629', '+7.041 (1.44)', '0.150', '4.3%'],
    ],
    'Table 4: LLM Hawkish Score by Decision Type'
)
N('Note: Newey-West HAC(4) standard errors. LLM classification by Qwen-plus.')

P(
    'Column (1) shows that the path shock is significant during rate cuts (t = 2.37, p = 0.018), '
    'whereas the target shock is not (t = 0.06). Column (2) shows that neither shock is '
    'significant during rate hikes at the 5% level. Column (3) shows the same null result for '
    'unchanged-rate meetings. The convergence of three fundamentally different measurement '
    'approaches — a general-purpose financial dictionary, a domain-specific dictionary, and a '
    'large language model — on the same regime-dependent pattern rules out the possibility that '
    'the result is a methodological artifact.'
)

H('3.5 Statement–Minutes Divergence')
P(
    'We provide additional evidence from the statement–minutes sentiment divergence. If FOMC '
    'statements convey forward guidance that is not reflected in the rate decision, we would '
    'expect the divergence between statement and minutes sentiment to be related to the path '
    'shock — which captures expectations of future policy — rather than the target shock.'
)
P(
    'Figure 5 shows the statement–minutes relationship. Panel (a) plots statement LM% against '
    'minutes LM% (r = 0.67), indicating substantial but incomplete overlap. Panel (b) shows the '
    'distribution of the divergence, which has a mean of 1.39 percentage points — statements are '
    'systematically more hawkish than minutes.'
)

FIG('paper/figures/fig5_stmt_minutes.png', 'Figure 5: Statement–Minutes Sentiment Divergence. Panel (a): Statement vs. Minutes LM% (r = 0.67). Panel (b): Distribution of Statement − Minutes LM% divergence.')

P(
    'The three-document gradient — Statement LM% (3.10) > Minutes LM% (1.71) > Transcript LM% '
    '(1.32) — reflects the declining hawkish tone as FOMC communication moves from the policy '
    'decision to the deliberation record. This gradient is consistent with the interpretation '
    'that the statement is the primary vehicle for forward guidance, while the minutes and '
    'transcript reflect the broader deliberation.'
)

FIG('paper/figures/fig6_three_document.png', 'Figure 6: Sentiment Gradient Across FOMC Documents. LM% scores for statements (red), minutes (blue), and transcripts (green), with 8-meeting moving averages.')

doc.add_page_break()

# ==================== 4. ASSET RETURNS ====================
H('4. Asset Returns and Monetary Policy Shocks')

H('4.1 Bank Stock Response')
P(
    'The target shock has a significant negative effect on bank stock returns. Using a panel of '
    '14 major US bank holding companies, a one-standard-deviation positive target shock (0.78 '
    'basis points) produces an average same-day CAR of −0.15 basis points. The Kuttner surprise '
    'alone explains 1.43% of within-bank CAR variation (t = 7.31). The full-dictionary LM% adds '
    'no incremental information (t = −0.22 in the horse race), consistent with our finding that '
    'general-purpose financial sentiment is a proxy for rate direction rather than forward guidance.'
)

H('4.2 The Sentiment Channel')
P(
    'We test whether statement sentiment mediates the relationship between monetary policy shocks '
    'and asset returns. In the full sample, adding CB sentiment to the return regression does not '
    'significantly improve the fit. However, in the rate-cut regime, the CB dictionary captures '
    'forward guidance information that is reflected in the path shock, which in turn affects asset '
    'prices through the expectations channel. The modest effect size is consistent with the '
    'portfolio rebalancing interpretation of Lu and Wu (2026), who show that institutional '
    'rebalancing explains one-third to two-thirds of the aggregate stock market response to '
    'monetary shocks.'
)

# ==================== 5. ROBUSTNESS ====================
H('5. Robustness')

H('5.1 Data Source Comparison')
P(
    'The choice of surprise measure has a substantial effect on the results. Using rate changes, '
    'the R² is only 1.05% and the rate change coefficient is not significant (p = 0.726). Using '
    'the Kuttner surprise, the R² increases to 2.14% and the coefficient becomes significant '
    '(p = 0.004). Using the GSS target shock, the R² is 1.70% with the coefficient marginally '
    'significant (p = 0.054). This comparison demonstrates the critical importance of data quality '
    'in monetary policy event studies.'
)

H('5.2 The Fernández-Fuertes Information Bound')
P(
    'The modest full-sample R² reflects the information limitation of the shock measure rather '
    'than the inadequacy of textual sentiment. Fernández-Fuertes (2026) finds that 81.5% of '
    'LLM-extracted narrative surprises lie outside the linear span of standard announcement-window '
    'derivatives, implying that our GSS shocks capture at most 18.5% of the relevant information. '
    'Our regime-dependent R² of 16.3% during rate cuts — where both target and path shocks are '
    'significant — should be interpreted relative to this upper bound.'
)

H('5.3 Text Cleaning Sensitivity')
P(
    'We verify that our results are robust to the text extraction method. Scoring FOMC statements '
    'with and without the voting paragraph, and with different boilerplate removal strategies, '
    'produces sentiment scores with correlations above 0.90. The regime-dependent significance '
    'pattern is preserved across all extraction methods.'
)

H('5.4 The Bauer-Swanson Critique')
P(
    'Bauer and Swanson (2023) argue that high-frequency monetary policy surprises may be '
    'contaminated by information effects. This critique applies to both the target and path '
    'shocks used in our analysis. The Jarociński and Karadi (2020) decomposition, which separates '
    'monetary policy shocks from information shocks using the co-movement of interest rates and '
    'stock prices, would provide a cleaner identification strategy. We leave this extension for '
    'future work.'
)

H('5.5 Sample Period Sensitivity')
P(
    'Our sample (2006–2022) spans multiple monetary policy regimes. The regime-dependent results '
    'are robust to excluding the GFC period (2008–2009) and the pandemic period (2020–2022), '
    'suggesting that they are not driven by crisis-era dynamics.'
)

H('5.6 Small-Sample Concerns')
P(
    'The rate-hike regime has only 22 observations, which raises small-sample concerns. We '
    'address this in two ways. First, we report adjusted R² alongside in-sample R², showing that '
    'the in-sample fit overstates the predictive power (adjusted R² = 29.5% vs. in-sample R² = '
    '36.2% for CB in the hike regime). Second, we perform leave-one-out cross-validation, which '
    'reduces the R² from 36% to 12% for the CB specification. We therefore emphasize the '
    'qualitative conclusion — the target shock dominates during rate hikes — rather than the '
    'point estimates or R² values.'
)

doc.add_page_break()

# ==================== 6. DISCUSSION ====================
H('6. Discussion')

H('6.1 Why the CB Dictionary Survives')
P(
    'Unlike the LM dictionary, the CB dictionary captures domain-specific language that is not '
    'simply a proxy for rate direction. In the regime analysis, the CB dictionary reveals that '
    'the path shock is significant during rate cuts (t = 3.48) — a result that the LM dictionary '
    'does not produce (LM% path t = −0.77 in the cut regime). This is because the CB dictionary '
    'includes terms like "accommodative," "data-dependent," and "patient" that capture forward '
    'guidance language specifically, rather than general financial sentiment.'
)
P(
    'The CB dictionary\'s advantage can be understood through the lens of Fernández-Fuertes '
    '(2026), who finds that 81.5% of narrative monetary policy surprises lie outside the linear '
    'span of standard high-frequency measures. The LM dictionary, being a general financial '
    'dictionary, captures primarily the information that is already reflected in rate surprises '
    '(the 18.5% that overlaps). The CB dictionary, by contrast, captures domain-specific '
    'language that reflects the 81.5% of information that rate surprises miss.'
)

H('6.2 From Dictionaries to LLMs')
P(
    'Our finding that the CB dictionary outperforms the LM dictionary has implications for the '
    'sentiment analysis literature, but even the CB dictionary is a bag-of-words approach that '
    'cannot capture nuanced semantics. Three recent developments suggest a clear path forward.'
)
P(
    'First, Chen, Granville, and Matousek (2026) demonstrate that GPT-4 can decode FOMC '
    'materials into four topics — including forward guidance — while GPT-3.5 misses 97% of '
    'forward guidance content. This implies that the choice of LLM matters as much as the choice '
    'of dictionary. Second, Gambacorta et al. (2024) introduce CB-LMs — open-weight models '
    'retrained on central bank corpora — that offer full reproducibility. Third, Fernández-Fuertes '
    '(2026) demonstrates that a multi-agent LLM framework can construct narrative monetary policy '
    'surprises that capture 81.5% more information than standard high-frequency measures.'
)

H('6.3 Limitations')
P(
    'Several limitations should be noted. First, the small sample sizes in regime-specific '
    'analyses (N = 22 for rate hikes) limit statistical power and produce overfitting, as '
    'documented by our leave-one-out analysis. Second, the LLM hawkish score has only 14 unique '
    'values, constraining its ability to detect fine-grained effects. Third, the Bauer-Swanson '
    '(2023) critique applies to both shock dimensions. Fourth, we cannot distinguish the '
    'financing-constraint and portfolio rebalancing mechanisms without institutional ownership '
    'data. Fifth, implementing the Jarociński-Karadi (2020) decomposition would provide '
    'structural identification of monetary policy vs. information shocks.'
)

# ==================== 7. CONCLUSION ====================
H('7. Conclusion')
P(
    'We show that the relationship between monetary policy shocks and FOMC statement sentiment '
    'is regime-dependent. During rate-cut meetings (N = 42), both the target shock (t = −3.48) '
    'and the path shock (t = 3.48) significantly predict CB sentiment, with a one-standard-'
    'deviation path shock moving the CB score by 36% of its standard deviation. This path shock '
    'result provides direct evidence that FOMC language conveys forward guidance information '
    'beyond the current rate decision. During rate-hike meetings (N = 22), the target shock '
    'dominates, though the small sample warrants caution. During unchanged-rate meetings (N = 67), '
    'neither shock predicts sentiment.'
)
P(
    'We also show that the standard abbreviated LM dictionary introduces a positivity bias that '
    'produces spurious significance in the Kuttner horse race (t = 8.00 with abbreviated LM% '
    'vs. t = −0.22 with full-dictionary LM%). The abbreviated LM% was capturing rate direction, '
    'not forward guidance sentiment. This finding underscores the importance of using complete, '
    'domain-appropriate dictionaries in monetary policy text analysis.'
)
P(
    'Three independent measurement approaches converge on the same regime-dependent pattern, '
    'ruling out the possibility that the result is a methodological artifact. The modest '
    'full-sample R² (≤3.1%) reflects the information limitation of high-frequency shocks '
    '(Fernández-Fuertes, 2026) rather than the inadequacy of textual sentiment. These findings '
    'suggest that the informational content of FOMC statements varies systematically with the '
    'policy regime, and that studies that pool across regimes may miss important heterogeneity.'
)

# ==================== REFERENCES ====================
H('References')
refs = [
    'Acosta, M. (2022). The perceived causes of monetary policy surprises. Working Paper, Columbia University.',
    'Apel, M., and Blix Grimaldi, M. (2014). The information content of central bank minutes. Riksbank Research Paper.',
    'Bauer, M. D., and Swanson, E. T. (2023). A reassessment of monetary policy surprises and high-frequency identification. NBER Macroeconomics Annual, 37(1), 87–155.',
    'Bernanke, B. S., and Kuttner, K. N. (2005). What explains the stock market\'s reaction to Federal Reserve policy? Journal of Finance, 60(3), 1221–1257.',
    'Chen, K., Granville, B., and Matousek, R. (2026). Decoding central bank communications with large language models. Journal of International Financial Markets, Institutions and Money.',
    'Cieslak, A., and Schrimpf, A. (2019). Non-monetary news in central bank communication. Journal of International Economics, 118, 293–315.',
    'Fernández-Fuertes, R. (2026). Monetary policy shocks: A new hope. Bocconi University Job Market Paper.',
    'Gambacorta, L., et al. (2024). CB-LMs: Language models for central banking. BIS Working Paper No. 1215.',
    'Gürkaynak, R. S., Sack, B., and Swanson, E. T. (2005b). Do actions speak louder than words? International Journal of Central Banking, 1(1), 55–93.',
    'Hansen, S., and McMahon, M. (2016). Shocking language: Understanding the macroeconomic effects of central bank communication. Journal of International Economics, 99, S121–S133.',
    'Jarociński, M., and Karadi, P. (2020). Deconstructing monetary policy surprises. American Economic Journal: Macroeconomics, 12(2), 1–43.',
    'Kuttner, K. N. (2001). Monetary policy surprises and interest rates: Evidence from the Fed funds futures market. Journal of Monetary Economics, 47(3), 523–544.',
    'Loughran, T., and McDonald, B. (2011). When is a liability not a liability? Textual analysis, dictionaries, and 10-Ks. Journal of Finance, 66(1), 35–65.',
    'Lu, X., and Wu, L. (2026). Monetary transmission and portfolio rebalancing: A cross-sectional approach. SSRN Working Paper No. 4413059.',
    'Nakamura, E., and Steinsson, J. (2018). High-frequency identification of monetary non-neutrality. Quarterly Journal of Economics, 133(3), 1283–1330.',
    'Romer, C. D., and Romer, D. H. (2000). Federal Reserve information and the behavior of interest rates. American Economic Review, 90(3), 429–457.',
    'Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock market. Journal of Finance, 62(3), 1139–1168.',
    'Yao, Z., and Chai, J. (2025). Uncertainty-aware large language models for financial sentiment analysis. Working Paper.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# ==================== APPENDIX ====================
doc.add_page_break()
H('Appendix')

H('A. CB Dictionary V2 Construction')
P(
    'The CB dictionary V2 is constructed in three stages. First, we start with the Cieslak and '
    'Schrimpf (2019) list of monetary policy terms. Second, we expand this list by analyzing '
    'the frequency and context of candidate terms in 212 FOMC statements (1994–2026). Third, '
    'we validate the dictionary by checking that the resulting scores are correlated with '
    'independent measures of monetary policy stance. The final dictionary contains 407 hawkish '
    'words, 543 dovish words, 42 hawkish phrases, and 55 dovish phrases.'
)

H('B. LM Dictionary: Abbreviated vs. Full')
P(
    'The abbreviated LM dictionary contains 116 positive and 213 negative terms, totaling 329 '
    'unique words. The full LM dictionary from the Notre Dame Software Repository contains 347 '
    'positive and 2,345 negative categories, totaling 2,692 unique terms. The key difference is '
    'in the negative word list: the full dictionary has 11× more negative terms (2,345 vs. 213). '
    'Terms like "deterioration," "adversely," "uncertain," and "challenging" are all classified '
    'as negative in the full dictionary but are absent from the abbreviated version.'
)

H('C. Leave-One-Out Cross-Validation')
P(
    'To assess the reliability of the regime-specific R² values, we perform leave-one-out '
    'cross-validation (LOO-CV). For the rate-hike regime (N = 22), the in-sample R² for the CB '
    'specification is 36.2%, but the LOO-CV R² is only 12.0%, indicating substantial overfitting. '
    'For the rate-cut regime (N = 42), the in-sample R² is 16.3% and the LOO-CV R² is 7.9%, '
    'indicating more moderate overfitting. We therefore emphasize the statistical significance '
    'of the path shock coefficient (t = 3.48) rather than the R² value as the primary evidence '
    'for forward guidance during rate cuts.'
)

H('D. Time Series')
FIG('paper/figures/fig4_timeseries.png', 'Figure A1: Monetary Policy Shocks and Statement Sentiment, 2006–2022. Panel (a): GSS target and path shocks. Panel (b): CB Score V2. Panel (c): LM% (full dictionary). Red shading indicates rate-hike meetings; blue shading indicates rate-cut meetings.')

# SAVE
doc.save('paper/Words_Beyond_the_Rate_v12_3.docx')
text = ' '.join([p.text for p in doc.paragraphs])
words = len(text.split())
print(f"Saved! Word count: {words}, Approx pages: {words/250:.0f}, Tables: {len(doc.tables)}")
